import os
from docx import Document

# Base folder inside your project
BASE = os.path.join(os.path.dirname(__file__), 'templates', 'docx')
os.makedirs(BASE, exist_ok=True)

# Helper to add a label + placeholder paragraph
def add_labeled_paragraph(doc, label, placeholder):
    if label:
        doc.add_paragraph(label)
    doc.add_paragraph(placeholder)

# 1) Invite Letter
doc = Document()
doc.add_heading('Invite to Performance Improvement Plan Meeting', level=1)

# Header details
add_labeled_paragraph(doc, 'Date:', '{{today}}')
add_labeled_paragraph(doc, 'To:', '{{employee_name}} ({{employee_role}})')
add_labeled_paragraph(doc, 'Service:', '{{employee_service}}')
doc.add_paragraph('')

# Greeting
doc.add_paragraph('Dear {{employee_name}},')
doc.add_paragraph('')

# Body
doc.add_paragraph(
    'You are invited to a meeting to discuss your Performance Improvement Plan (PIP). '
    'The details are as follows:'
)
add_labeled_paragraph(doc, 'Meeting date:', '{{meeting_date}}')
add_labeled_paragraph(doc, 'Meeting time:', '{{meeting_time}}')
add_labeled_paragraph(doc, 'Location/venue:', '{{meeting_location}}')
doc.add_paragraph('')

# Context summary (from PIP + AI)
add_labeled_paragraph(doc, 'Summary of concerns:', '{{concerns_summary}}')
add_labeled_paragraph(doc, 'AI summary (optional):', '{{ai_summary}}')
add_labeled_paragraph(doc, 'AI suggested focus areas:', '{{ai_action_suggestions_bullets}}')
doc.add_paragraph('')

# Closing
doc.add_paragraph(
    'If you require any adjustments or support in advance of the meeting, please let me know.'
)
doc.add_paragraph('')
doc.add_paragraph('Sincerely,')
doc.add_paragraph('{{manager_name}}')
doc.add_paragraph('{{manager_role}}')
doc.save(os.path.join(BASE, 'invite_letter_template.docx'))


# 2) Plan Template
doc = Document()
doc.add_heading('Performance Improvement Plan', level=1)

# Header details
add_labeled_paragraph(doc, 'Employee:', '{{employee_name}} ({{employee_role}})')
add_labeled_paragraph(doc, 'Service:', '{{employee_service}}')
add_labeled_paragraph(doc, 'PIP Start Date:', '{{pip_start_date}}')
add_labeled_paragraph(doc, 'PIP Review Date:', '{{pip_review_date}}')
add_labeled_paragraph(doc, 'PIP End Date (if set):', '{{pip_end_date}}')
doc.add_paragraph('')

# Concerns & evidence
add_labeled_paragraph(doc, 'Concerns:', '{{concerns_summary}}')
add_labeled_paragraph(doc, 'AI summary (optional):', '{{ai_summary}}')
doc.add_paragraph('')

# Action plan (accepted actions preferred; falls back to all suggestions)
doc.add_paragraph('Action Plan:')
doc.add_paragraph('{{plan_actions_bullets}}')  # will be newline-separated items
doc.add_paragraph('')

# (Optional) Existing action items already recorded in the system
doc.add_paragraph('Current Action Items on Record:')
doc.add_paragraph('{{current_action_items_bullets}}')  # populate from DB if you like
doc.add_paragraph('')

# Next steps nudges
add_labeled_paragraph(doc, 'Next up (AI):', '{{ai_next_up_bullets}}')
doc.add_paragraph('')

# Sign-off
add_labeled_paragraph(doc, 'Manager:', '{{manager_name}}')
add_labeled_paragraph(doc, 'Role:', '{{manager_role}}')
doc.save(os.path.join(BASE, 'plan_template.docx'))


# 3) Outcome Letter
doc = Document()
doc.add_heading('Outcome of Performance Improvement Plan', level=1)

# Header details
add_labeled_paragraph(doc, 'Date:', '{{today}}')
add_labeled_paragraph(doc, 'Employee:', '{{employee_name}} ({{employee_role}})')
add_labeled_paragraph(doc, 'Service:', '{{employee_service}}')
doc.add_paragraph('')

# Greeting
doc.add_paragraph('Dear {{employee_name}},')
doc.add_paragraph('')

# Outcome statement
doc.add_paragraph('Following the review of your Performance Improvement Plan, the outcome is:')
doc.add_paragraph('{{outcome_status}}')  # e.g., Successful / Extended / Unsuccessful
doc.add_paragraph('')

# Notes and rationale
add_labeled_paragraph(doc, 'Outcome notes:', '{{outcome_notes}}')
add_labeled_paragraph(doc, 'Summary of concerns (for record):', '{{concerns_summary}}')
doc.add_paragraph('')

# Optional AI reflection / next steps
add_labeled_paragraph(doc, 'AI summary (optional):', '{{ai_summary}}')
add_labeled_paragraph(doc, 'Recommended next steps (AI):', '{{ai_next_up_bullets}}')
doc.add_paragraph('')

# Closing
doc.add_paragraph('Thank you for your engagement with the process.')
doc.add_paragraph('')
doc.add_paragraph('Sincerely,')
doc.add_paragraph('{{manager_name}}')
doc.add_paragraph('{{manager_role}}')
doc.save(os.path.join(BASE, 'outcome_letter_template.docx'))

print("Templates written to:", BASE)
