from __future__ import annotations

import os
import re
from io import BytesIO
from pathlib import Path

import bleach
import mammoth
from docx import Document
from flask import abort
from html2docx import html2docx

from pip_app.services.time_utils import now_local

BASE_DIR = str(Path(__file__).resolve().parents[2])
TEMPLATE_DIR = Path(BASE_DIR) / "templates" / "docx"

LEGACY_TO_NEW_KEYS = {}

ALLOWED_TAGS = bleach.sanitizer.ALLOWED_TAGS.union({
    "p", "h1", "h2", "h3", "h4", "h5", "h6",
    "strong", "em", "u", "span", "div", "br", "hr",
    "ul", "ol", "li",
    "table", "thead", "tbody", "tr", "th", "td",
    "blockquote"
})
ALLOWED_ATTRS = {"*": ["class"]}


def build_doc_rel_dir(pip_id: int, doc_type: str, version: int) -> str:
    return os.path.join("pip_docs", str(pip_id), doc_type, f"v{version}")


def build_placeholder_mapping(pip_rec):
    employee = getattr(pip_rec, "employee", None)

    employee_name = ""
    if employee:
        employee_name = f"{employee.first_name or ''} {employee.last_name or ''}".strip()

    action_items = getattr(pip_rec, "action_items", []) or []
    action_lines = []
    for idx, item in enumerate(action_items, start=1):
        desc = (getattr(item, "description", "") or "").strip()
        status = (getattr(item, "status", "") or "").strip()
        if desc and status:
            action_lines.append(f"{idx}. {desc} ({status})")
        elif desc:
            action_lines.append(f"{idx}. {desc}")

    return {
        "[[EMPLOYEE_NAME]]": employee_name,
        "[[FIRST_NAME]]": getattr(employee, "first_name", "") or "",
        "[[LAST_NAME]]": getattr(employee, "last_name", "") or "",
        "[[JOB_TITLE]]": getattr(employee, "job_title", "") or "",
        "[[LINE_MANAGER]]": getattr(employee, "line_manager", "") or "",
        "[[SERVICE]]": getattr(employee, "service", "") or "",
        "[[EMAIL]]": getattr(employee, "email", "") or "",
        "[[START_DATE]]": pip_rec.start_date.strftime("%d/%m/%Y") if getattr(pip_rec, "start_date", None) else "",
        "[[REVIEW_DATE]]": pip_rec.review_date.strftime("%d/%m/%Y") if getattr(pip_rec, "review_date", None) else "",
        "[[CAPABILITY_MEETING_DATE]]": (
            pip_rec.capability_meeting_date.strftime("%d/%m/%Y")
            if getattr(pip_rec, "capability_meeting_date", None) else ""
        ),
        "[[CAPABILITY_MEETING_TIME]]": getattr(pip_rec, "capability_meeting_time", "") or "",
        "[[CAPABILITY_MEETING_VENUE]]": getattr(pip_rec, "capability_meeting_venue", "") or "",
        "[[CONCERNS]]": getattr(pip_rec, "concerns", "") or "",
        "[[CONCERN_CATEGORY]]": getattr(pip_rec, "concern_category", "") or "",
        "[[SEVERITY]]": getattr(pip_rec, "severity", "") or "",
        "[[FREQUENCY]]": getattr(pip_rec, "frequency", "") or "",
        "[[TAGS]]": getattr(pip_rec, "tags", "") or "",
        "[[MEETING_NOTES]]": getattr(pip_rec, "meeting_notes", "") or "",
        "[[ACTION_PLAN_ITEMS]]": "\n".join(action_lines),
        "[[OUTCOME_CHOICE]]": getattr(pip_rec, "outcome_choice", "") or "",
        "[[CREATED_BY]]": getattr(pip_rec, "created_by", "") or "",
        "[[GENERATED_DATE]]": now_local().strftime("%d %B %Y"),
    }


def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p


def replace_placeholders_docx(doc: Document, context: dict):
    base_ctx = context or {}
    mapping = {k: ("" if v is None else str(v)) for k, v in base_ctx.items()}

    for legacy_key, new_key in LEGACY_TO_NEW_KEYS.items():
        if legacy_key in base_ctx and new_key not in mapping:
            v = base_ctx.get(legacy_key)
            mapping[new_key] = "" if v is None else str(v)

    now_uk = now_local()
    mapping.setdefault("[[GENERATED_DATE]]", now_uk.strftime("%d %B %Y"))
    mapping.setdefault("[[DOC_VERSION]]", now_uk.strftime("v%Y.%m.%d"))

    placeholder_pattern = re.compile(r"\[\[[A-Z0-9_]+\]\]")

    for p in _iter_all_paragraphs(doc):
        if not p.runs:
            continue

        original_text = "".join(run.text for run in p.runs)
        if not original_text:
            continue

        new_text = original_text
        for k, v in mapping.items():
            if k in new_text:
                new_text = new_text.replace(k, v)

        if "[[" in new_text:
            new_text = placeholder_pattern.sub("", new_text)

        if new_text == original_text:
            continue

        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""


def strip_outcome_conditionals(doc: Document, keep: str):
    valid = {"SUCCESSFUL", "EXTENSION", "UNSUCCESSFUL"}
    choice = (keep or "").upper().strip()
    if choice not in valid:
        raise ValueError(f"Invalid outcome choice: {keep}")

    start_tokens = {f"[[IF_{tag}]]": tag for tag in valid}
    in_block = None
    keep_block = False
    to_delete = []

    paragraphs = list(doc.paragraphs)

    def contains_token(p, token):
        return any(token in r.text for r in p.runs) or token in p.text

    for p in paragraphs:
        for token, tag in start_tokens.items():
            if contains_token(p, token):
                in_block = tag
                keep_block = (tag == choice)
                for r in p.runs:
                    r.text = r.text.replace(token, "")
                if not keep_block:
                    to_delete.append(p)
                break

        if in_block:
            if not keep_block and p not in to_delete:
                to_delete.append(p)
            end_token = f"[[/IF_{in_block}]]"
            if contains_token(p, end_token):
                for r in p.runs:
                    r.text = r.text.replace(end_token, "")
                if not keep_block and p not in to_delete:
                    to_delete.append(p)
                in_block = None
                keep_block = False

    for p in to_delete:
        p._element.getparent().remove(p._element)


def render_docx(template_filename: str, context: dict, outcome_choice: str | None = None) -> BytesIO:
    template_path = TEMPLATE_DIR / template_filename
    if not template_path.exists():
        abort(404, f"Template not found: {template_path.name}")
    doc = Document(str(template_path))
    replace_placeholders_docx(doc, context or {})
    if outcome_choice:
        strip_outcome_conditionals(doc, outcome_choice)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def sanitize_html(html: str) -> str:
    return bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS, strip=True)


def generate_docx_bytes(template_path: str, mapping: dict, outcome_choice: str = None) -> bytes:
    doc = Document(template_path)
    replace_placeholders_docx(doc, mapping)
    if outcome_choice:
        strip_outcome_conditionals(doc, outcome_choice)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def docx_to_html(docx_bytes: bytes) -> str:
    with BytesIO(docx_bytes) as f:
        return mammoth.convert_to_html(f).value


def html_to_docx_bytes(html: str) -> bytes:
    clean_html = sanitize_html(html or "").strip()
    if not clean_html:
        return b""

    result = html2docx(clean_html, title="PIP Document")

    if isinstance(result, bytes):
        return result

    if hasattr(result, "getvalue"):
        data = result.getvalue()
        return data if data else b""

    if hasattr(result, "save"):
        out = BytesIO()
        result.save(out)
        data = out.getvalue()
        return data if data else b""

    raise ValueError(f"Unsupported html2docx return type: {type(result)}")