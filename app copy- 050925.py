import os
import csv, io
import zipfile
import tempfile
import re
import bleach

from io import BytesIO
from functools import wraps
from datetime import datetime, timedelta, timezone
from pathlib import Path
from zoneinfo import ZoneInfo

from sqlalchemy.sql import func
from sqlalchemy.exc import SQLAlchemyError

from flask import (
    Flask, session, render_template, redirect, url_for,
    request, flash, send_file, jsonify, abort
)
from flask_migrate import Migrate
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, login_user, login_required,
    logout_user, current_user
)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect
from flask_wtf.csrf import generate_csrf

from flask import render_template, request, redirect, url_for, flash, send_from_directory, abort
from flask_login import login_required, current_user
from models import db, PIPRecord, DocumentFile, TimelineEvent  # <-- ensure DocumentFile is in models.py
import os
from werkzeug.utils import secure_filename
from io import BytesIO
import mammoth
from html2docx import html2docx
from datetime import datetime




from dotenv import load_dotenv
load_dotenv()

# Use python-docx for our new templates system
from docx import Document

# If you have openpyxl installed, we can read .xlsx; otherwise CSV only.
try:
    import openpyxl  # noqa: F401
    XLSX_ENABLED = True
except Exception:
    XLSX_ENABLED = False

ALLOWED_EXTS = {"csv", "xlsx"} if XLSX_ENABLED else {"csv"}

# Map Employee model fields to import
EMPLOYEE_FIELDS = [
    "first_name", "last_name", "email", "job_title", "line_manager",
    "service", "team_id", "start_date"
]

# Minimal required fields
REQUIRED_FIELDS = ["first_name", "last_name"]

# ---- Curated concern tags (can move to DB later) ----
CURATED_TAGS = {
    "Timekeeping": ["lateness", "missed clock-in", "extended breaks", "early finish", "timekeeping policy"],
    "Attendance": ["unauthorised absence", "short notice absence", "patterns of absence", "fit note", "return to work"],
    "Quality of Work": ["accuracy", "attention to detail", "rework", "documentation", "SOP adherence"],
    "Productivity": ["missed deadlines", "slow throughput", "low output", "prioritisation", "time management"],
    "Conduct": ["inappropriate language", "unprofessional behaviour", "non-compliance", "policy breach", "conflict"],
    "Communication": ["tone", "late replies", "stakeholder updates", "handover quality", "listening"],
    "Teamwork/Collaboration": ["handover gaps", "knowledge sharing", "supporting peers", "collaboration tools"],
    "Compliance/Process": ["data entry errors", "checklist missed", "audit finding", "process deviation"],
    "Customer Service": ["response times", "complaint handling", "service standards", "follow-up"],
    "Health & Safety": ["PPE", "risk assessment", "manual handling", "incident reporting"]
}

def _merge_curated_and_recent(category: str, recent_tags: list[str], cap: int = 30):
    """Merge curated list for category with recent DB tags, de-duplicated, curated first."""
    out, seen = [], set()
    cat_list = CURATED_TAGS.get(category, [])
    for t in cat_list:
        if t.lower() not in seen:
            out.append(t); seen.add(t.lower())
    for t in recent_tags:
        if not t: continue
        key = t.lower().strip()
        if key and key not in seen:
            out.append(t.strip()); seen.add(key)
        if len(out) >= cap: break
    return out

# ---- Curated Action Plan Templates (by category + severity) ----
# You can tune these freely later. Keep items short and concrete.
ACTION_TEMPLATES = {
    "Timekeeping": {
        "Low": [
            "Agree start-time target and grace window",
            "Daily check-in for 2 weeks at start of shift",
            "Keep punctuality log; review weekly",
        ],
        "Moderate": [
            "Formal punctuality target with variance log",
            "Escalate if 2+ breaches in a week",
            "Buddy assigned for morning routine",
        ],
        "High": [
            "Issue written reminder citing policy",
            "Daily manager sign-off for 3 weeks",
            "Escalation to formal stage if breaches continue",
        ],
        "_default": [
            "Agree punctuality expectations",
            "Daily check-in for first 2 weeks",
            "Weekly review of log",
        ],
    },
    "Performance": {
        "Low": [
            "Break down tasks into weekly milestones",
            "Mid-week check-in with progress update",
            "Share example of quality standard",
        ],
        "Moderate": [
            "Set SMART targets per task",
            "Stand-up updates Mon/Wed/Fri",
            "Peer review before handoff",
        ],
        "High": [
            "Written performance targets with deadlines",
            "Daily status update for 10 working days",
            "Escalate to formal PIP stage if no improvement",
        ],
        "_default": [
            "Agree 2–3 SMART targets",
            "Weekly progress review",
            "Identify training/module to close gap",
        ],
    },
    "Conduct": {
        "_default": [
            "Reference conduct policy and expectations",
            "Agree behaviour standards; confirm by email",
            "Book values refresher session",
        ]
    },
    "Attendance": {
        "_default": [
            "Follow reporting procedure for absence",
            "Return-to-work meeting after each absence",
            "Pattern review after 4 weeks",
        ]
    },
    "Communication": {
        "_default": [
            "Acknowledge messages within agreed SLA",
            "Use agreed update template for stakeholders",
            "Add handover note at end of shift",
        ]
    },
    "Quality of Work": {
        "_default": [
            "Introduce checklist for critical steps",
            "Peer review for first 4 weeks",
            "Log defects and agree prevention steps",
        ]
    },
    "Productivity": {
        "_default": [
            "Time-block key tasks; share plan daily",
            "Weekly throughput targets",
            "Remove low-value tasks with manager",
        ]
    },
}

def _pick_actions_from_templates(category: str, severity: str) -> list[str]:
    cat = (category or "").strip()
    sev = (severity or "").strip()
    block = ACTION_TEMPLATES.get(cat) or {}
    if not block:
        # fall back to a generic pack
        block = {"_default": ["Agree clear targets", "Weekly review", "Training / buddy support as needed"]}
    # prefer exact severity, else _default, else any first list
    if sev in block and block[sev]:
        return block[sev]
    if block.get("_default"):
        return block["_default"]
    # final fallback: first list found
    for v in block.values():
        if isinstance(v, list) and v:
            return v
    return []

def _read_csv_bytes(file_bytes: bytes):
    text = file_bytes.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    rows = [dict(r) for r in reader]
    headers = reader.fieldnames or []
    return headers, rows

def _read_xlsx_bytes(file_bytes: bytes):
    # Requires openpyxl; already guarded by XLSX_ENABLED
    workbook = openpyxl.load_workbook(io.BytesIO(file_bytes))
    sheet = workbook.active
    headers = [c.value for c in next(sheet.iter_rows(min_row=1, max_row=1))]
    rows = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        rows.append({headers[i]: (row[i] if i < len(headers) else None) for i in range(len(headers))})
    return headers, rows

def _normalize_header(h):
    # Loose normalisation to help mapping (e.g., "First Name" -> "first_name")
    return (h or "").strip().lower().replace(" ", "_")

def _try_parse_date(value):
    if not value:
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(str(value).strip(), fmt).date()
        except Exception:
            continue
    return None

def _parse_iso_date(s):
    try:
        return datetime.strptime((s or "").strip(), "%Y-%m-%d").date()
    except Exception:
        return None


def _open_pips_scoped_query():
    """
    Returns a base query for *open* PIPs, automatically scoped to the current user's team
    when admin_level == 0. Use this to build stats safely.
    """
    base = PIPRecord.query.filter(PIPRecord.status == 'Open')
    if current_user.admin_level == 0:
        base = base.join(Employee).filter(Employee.team_id == current_user.team_id)
    return base

def _counts_by_field(field_expr):
    """
    Utility to return {label: count} for a given PIPRecord field expression,
    using the same scoped base (open PIPs only, team-limited for level 0).
    """
    q = _open_pips_scoped_query().with_entities(field_expr, func.count(PIPRecord.id))\
                                 .group_by(field_expr)
    rows = q.all()
    # Normalize None → "Unspecified" for display
    out = {}
    for label, cnt in rows:
        out[(label or "Unspecified")] = int(cnt or 0)
    return out
def _next_version_for(pip_id, doc_type):
    last = (DocumentFile.query
            .filter_by(pip_id=pip_id, doc_type=doc_type)
            .order_by(DocumentFile.version.desc())
            .first())
    return 1 if not last else last.version + 1

def _save_file(bytes_data: bytes, rel_dir: str, filename: str) -> str:
    dir_path = os.path.join(app.config['UPLOAD_FOLDER'], rel_dir)
    os.makedirs(dir_path, exist_ok=True)
    fpath = os.path.join(dir_path, secure_filename(filename))
    with open(fpath, 'wb') as f:
        f.write(bytes_data)
    return os.path.join(rel_dir, secure_filename(filename))

def generate_docx_bytes(template_path: str, mapping: dict, outcome_choice: str = None) -> bytes:
    from docx import Document
    doc = Document(template_path)
    replace_placeholders_docx(doc, mapping)  # your existing helper
    if outcome_choice:
        strip_outcome_conditionals(doc, outcome_choice)  # your existing helper
    out = BytesIO()
    doc.save(out)
    return out.getvalue()

def docx_to_html(docx_bytes: bytes) -> str:
    with BytesIO(docx_bytes) as f:
        return mammoth.convert_to_html(f).value

def html_to_docx_bytes(html: str) -> bytes:
    out = BytesIO()
    html2docx(html, out)
    return out.getvalue()

def build_doc_rel_dir(pip_id: int, doc_type: str, version: int) -> str:
    return os.path.join(f"PIP_{pip_id}", doc_type, f"v{version}")

def build_placeholder_mapping(pip_rec: PIPRecord) -> dict:
    emp = getattr(pip_rec, "employee", None)

    # Convenient accessors
    first = getattr(emp, "first_name", "") or ""
    last  = getattr(emp, "last_name", "") or ""
    full  = f"{first} {last}".strip()
    job   = getattr(emp, "job_title", "") or getattr(emp, "role", "") or ""
    dept  = getattr(emp, "service", "") or ""
    mgr   = getattr(pip_rec, "created_by", "") or getattr(emp, "line_manager", "") or getattr(current_user, "username", "") or ""

    # Compose meeting datetime string
    md = pip_rec.capability_meeting_date.strftime("%d %B %Y") if getattr(pip_rec, "capability_meeting_date", None) else ""
    mt = pip_rec.capability_meeting_time or ""
    meeting_dt = f"{md} {mt}".strip()

    # Base fields shared across templates
    mapping = {
        "[[GENERATED_DATE]]": datetime.now(ZoneInfo("Europe/London")).strftime("%d %B %Y"),
        "[[LETTER_DATE]]":    datetime.now(ZoneInfo("Europe/London")).strftime("%d %B %Y"),
        "[[DOC_VERSION]]":    datetime.now().strftime("v%Y.%m.%d"),

        # Employee / manager
        "[[EMPLOYEE_NAME]]":        full,
        "[[EMPLOYEE_FIRST_NAME]]":  first,
        "[[JOB_TITLE]]":            job,
        "[[DEPARTMENT]]":           dept,
        "[[MANAGER_NAME]]":         mgr,
        "[[MANAGER_TITLE]]":        getattr(emp, "line_manager", "") or "",

        # Meeting details (invite)
        "[[MEETING_LOCATION]]":     getattr(pip_rec, "capability_meeting_venue", "") or "",
        "[[MEETING_DATETIME]]":     meeting_dt,

        # PIP dates
        "[[PIP_START_DATE]]":       pip_rec.start_date.strftime("%d %B %Y") if getattr(pip_rec, "start_date", None) else "",
        "[[PIP_END_DATE]]":         pip_rec.review_date.strftime("%d %B %Y") if getattr(pip_rec, "review_date", None) else "",

        # Narrative fields
        "[[CONCERNS_SUMMARY]]":     getattr(pip_rec, "concerns", "") or "",
        "[[EVIDENCE_SUMMARY]]":     getattr(pip_rec, "evidence", "") or "",
        "[[SUPPORT_LIST]]":         getattr(pip_rec, "support_list", "") or "",

        # Meta / contacts
        "[[REVIEW_PERIOD_WEEKS]]":  str(getattr(pip_rec, "review_weeks", "") or ""),
        "[[HR_CONTACT_NAME]]":      getattr(pip_rec, "hr_contact_name", "") or "",
        "[[HR_CONTACT_EMAIL]]":     getattr(pip_rec, "hr_contact_email", "") or "",

        # Plan-specific extras (safe to leave blank for other docs)
        "[[REVIEWER_NAME]]":        getattr(pip_rec, "reviewer_name", "") or "",
        "[[POLICY_REFERENCE]]":     getattr(pip_rec, "policy_reference", "") or "",
        "[[CONCERN_CATEGORIES]]":   getattr(pip_rec, "concern_category", "") or "",
        "[[CONCERN_TAGS]]":         getattr(pip_rec, "tags", "") or "",
        "[[SEVERITY_RATING]]":      getattr(pip_rec, "severity", "") or "",
        "[[FREQUENCY_RATING]]":     getattr(pip_rec, "frequency", "") or "",
        "[[CONFIDENTIAL_NOTES]]":   getattr(pip_rec, "confidential_notes", "") or "",
        "[[REVIEW_DATES_LIST]]":    getattr(pip_rec, "review_dates_str", "") or "",
        "[[REPORTING_METHOD]]":     getattr(pip_rec, "reporting_method", "") or "",
        "[[ADJUSTMENTS]]":          getattr(pip_rec, "adjustments", "") or "",
        "[[TRAINING_PLAN]]":        getattr(pip_rec, "training_plan", "") or "",

        # Outcome-specific extras (safe blanks elsewhere)
        "[[ONGOING_SUPPORT]]":      getattr(pip_rec, "ongoing_support", "") or "",
        "[[EXTENSION_WEEKS]]":      str(getattr(pip_rec, "extension_weeks", "") or ""),
        "[[APPENDIX_DETAIL]]":      getattr(pip_rec, "outcome_appendix", "") or "",
    }

    return mapping


ALLOWED_TAGS = bleach.sanitizer.ALLOWED_TAGS.union({
    "p","h1","h2","h3","h4","h5","h6",
    "strong","em","u","span","div","br","hr",
    "ul","ol","li",
    "table","thead","tbody","tr","th","td",
    "blockquote"
})
ALLOWED_ATTRS = {
    "*": ["class"],   # removed "style"
}
def sanitize_html(html: str) -> str:
    return bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS, strip=True)



def log_timeline_event(pip_id: int, event_type: str, notes: str):
    try:
        ev = TimelineEvent(
            pip_record_id=pip_id,          # <-- fix
            event_type=event_type,
            notes=notes,
            updated_by=getattr(current_user, "username", None) or "system",
        )
        db.session.add(ev)
        db.session.commit()
    except Exception as e:
        app.logger.exception(f"TimelineEvent failed: {e}")


# --- helper: fetch active probation draft for banner ---
def get_active_probation_draft_for_user(user_id: int):
    return DraftProbation.query.filter_by(user_id=user_id, is_dismissed=False).first()

# ---------- Models & Forms ----------
from models import (
    db, User, Employee, PIPRecord, PIPActionItem, TimelineEvent,
    ProbationRecord, ProbationReview, ProbationPlan, DraftPIP, DraftProbation
)
from forms import (
    PIPForm, EmployeeForm, LoginForm, ProbationRecordForm,
    ProbationReviewForm, ProbationPlanForm, UserForm
)

# ---------- OpenAI ----------
from openai import OpenAI
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

# ---------- Flask Init ----------
import os
from flask import Flask
from flask_wtf.csrf import CSRFProtect

app = Flask(__name__)

# Secrets & base paths
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-only-secret')
BASE_DIR = os.path.abspath(os.path.dirname(__file__))

# Database
DB_PATH = os.path.join(BASE_DIR, 'pip_crm.db')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + DB_PATH
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# File uploads for generated documents
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads', 'documents')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# CSRF
app.config['WTF_CSRF_ENABLED'] = True          # default True; make explicit
app.config['WTF_CSRF_TIME_LIMIT'] = None        # tokens never expire while debugging
csrf = CSRFProtect(app)

db.init_app(app)
migrate = Migrate(app, db)

LoginManager.login_view = 'login'
LoginManager.login_message_category = 'info'


# ----- Context Processor -----
@app.context_processor
def inject_module():
    return dict(active_module=session.get('active_module'))

@app.context_processor
def inject_csrf_token():
    return dict(csrf_token=generate_csrf)

# ----- Login -----
login_manager = LoginManager()
login_manager.login_view = 'login'
login_manager.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

# ----- Superuser decorator -----
def superuser_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_superuser():
            abort(403)  # Forbidden
        return f(*args, **kwargs)
    return decorated_function

# ----- Set active module in session -----
@app.before_request
def set_active_module():
    path = request.path.lower()
    if path.startswith('/pip/') or \
       path.startswith('/employee/') or \
       path in ['/dashboard', '/pip_list', '/employee/list', '/employee/add', '/pip/select-employee']:
        session['active_module'] = 'PIP'
    elif path.startswith('/probation/'):
        session['active_module'] = 'Probation'
    elif path == '/':
        session.pop('active_module', None)
# ===== Document Helpers (NEW unified system for [[ALL_CAPS]] placeholders) =====
TEMPLATE_DIR = Path(__file__).resolve().parent / "templates" / "docx"
# Put these files in the folder above:
#   PIP_Invite_Letter_Template_v2025-08-28.docx
#   PIP_Action_Plan_Template_v2025-08-28.docx
#   PIP_Outcome_Letter_Template_v2025-08-28.docx

# Optional: support old keys alongside new ones if you still populate legacy fields
LEGACY_TO_NEW_KEYS = {
    # "employee_name": "[[EMPLOYEE_NAME]]",
    # "{EmployeeName}": "[[EMPLOYEE_NAME]]",
    # "{{employee.first_name}}": "[[EMPLOYEE_FIRST_NAME]]",
}

def _iter_all_paragraphs(doc: Document):
    # Body paragraphs
    for p in doc.paragraphs:
        yield p
    # Tables (safe to include)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    # Headers/footers for each section
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p

def _replace_in_runs(paragraph, mapping: dict):
    """Safer replacement that keeps run-level formatting if tokens aren't split across runs."""
    for run in paragraph.runs:
        text = run.text
        for k, v in mapping.items():
            if k in text:
                text = text.replace(k, v)
        run.text = text

def replace_placeholders_docx(doc: Document, context: dict):
    """
    Replace [[KEY]] placeholders throughout the document.
    Also maps any legacy keys to the new scheme if present in context.
    """
    # Build a single mapping of [[KEY]] → value (strings)
    mapping = {k: str(v) for k, v in (context or {}).items()}

    # Map any legacy keys in the context to the new keys
    for legacy_key, new_key in LEGACY_TO_NEW_KEYS.items():
        if legacy_key in (context or {}):
            mapping[new_key] = str(context[legacy_key])

    # Auto-fill some sensible defaults if missing
    now_uk = datetime.now(ZoneInfo("Europe/London"))
    mapping.setdefault("[[GENERATED_DATE]]", now_uk.strftime("%d %B %Y"))
    mapping.setdefault("[[DOC_VERSION]]", now_uk.strftime("v%Y.%m.%d"))

    # Replace across all paragraphs (runs to preserve formatting)
    for p in _iter_all_paragraphs(doc):
        _replace_in_runs(p, mapping)

def strip_outcome_conditionals(doc: Document, keep: str):
    """
    Keep exactly one block among [[IF_SUCCESSFUL]]...[[/IF_SUCCESSFUL]],
    [[IF_EXTENSION]]...[[/IF_EXTENSION]], [[IF_UNSUCCESSFUL]]...[[/IF_UNSUCCESSFUL]].
    Removes the others entirely. `keep` must be one of: "SUCCESSFUL", "EXTENSION", "UNSUCCESSFUL".
    """
    valid = {"SUCCESSFUL", "EXTENSION", "UNSUCCESSFUL"}
    choice = (keep or "").upper().strip()
    if choice not in valid:
        raise ValueError(f"Invalid outcome choice: {keep}")

    start_tokens = {f"[[IF_{tag}]]": tag for tag in valid}

    # We’ll scan once and mark paragraphs to delete as needed
    in_block = None
    keep_block = False
    to_delete = []

    paragraphs = list(doc.paragraphs)  # only body for conditionals

    def contains_token(p, token):
        return any(token in r.text for r in p.runs) or token in p.text

    for p in paragraphs:
        # Check for block start
        for token, tag in start_tokens.items():
            if contains_token(p, token):
                in_block = tag
                keep_block = (tag == choice)
                # Remove the token itself
                for r in p.runs:
                    r.text = r.text.replace(token, "")
                if not keep_block:
                    to_delete.append(p)
                break

        if in_block:
            # Inside a conditional block
            if not keep_block and p not in to_delete:
                to_delete.append(p)

            # Check for end of block
            end_token = f"[[/IF_{in_block}]]"
            if contains_token(p, end_token):
                for r in p.runs:
                    r.text = r.text.replace(end_token, "")
                if not keep_block and p not in to_delete:
                    to_delete.append(p)
                in_block = None
                keep_block = False

    # Physically remove any paragraphs marked for deletion
    for p in to_delete:
        p._element.getparent().remove(p._element)

def render_docx(template_filename: str, context: dict, outcome_choice: str | None = None) -> BytesIO:
    """
    Loads a .docx template, applies replacements, optional outcome filtering, and returns an in-memory file.
    """
    template_path = TEMPLATE_DIR / template_filename
    if not template_path.exists():
        abort(404, f"Template not found: {template_path.name}")

    doc = Document(str(template_path))
    replace_placeholders_docx(doc, context or {})
    if outcome_choice:
        strip_outcome_conditionals(doc, outcome_choice)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
# ===== End Document Helpers =====

# --- Wizard helper: compute default review date (4 weeks) ---
def _auto_review_date(start_date_str: str | None) -> str | None:
    if not start_date_str:
        return None
    try:
        dt = datetime.strptime(start_date_str.strip(), "%Y-%m-%d").date()
        return (dt + timedelta(days=28)).isoformat()  # 'YYYY-MM-DD'
    except Exception:
        return None




# ----- Routes: Home / Login / Admin -----
@app.route('/')
@login_required
def home():
    return render_template('select_module.html', hide_sidebar=True, layout='fullscreen')

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if request.method == 'POST' and form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and check_password_hash(user.password_hash, form.password.data):
            login_user(user, remember=form.remember.data)
            return redirect(url_for('home'))
        flash('Invalid username or password', 'danger')
    return render_template('login.html', form=form, hide_sidebar=True, current_year=datetime.now().year)

@app.route('/logout', methods=['POST', 'GET'])
def logout():
    try:
        logout_user()          # Flask-Login: clears user session + remember cookie
    finally:
        session.clear()        # also drop any app-specific session data
    flash("You have been logged out.", "info")
    # Redirect target: choose ONE
    return redirect(url_for('login'))                 # option A: login screen
    # return redirect(url_for('module_selection'))    # option B: module hub

@app.route('/admin/dashboard')
@login_required
def admin_dashboard():
    if not current_user.is_superuser():
        flash('You do not have permission to access the admin dashboard.', 'danger')
        return redirect(url_for('home'))
    return render_template('admin_dashboard.html')

# ----- Taxonomy (curated tags + categories) -----
@app.route('/taxonomy/predefined_tags', methods=['GET'])
@login_required
def taxonomy_predefined_tags():
    cat = (request.args.get('category') or '').strip()
    tags = CURATED_TAGS.get(cat, [])
    return jsonify({"category": cat, "tags": tags})

@app.route('/taxonomy/categories', methods=['GET'])
@login_required
def taxonomy_categories():
    categories = list(CURATED_TAGS.keys())
    return jsonify({"categories": categories})

@app.route('/taxonomy/tags_suggest', methods=['GET'])
@login_required
def taxonomy_tags_suggest():
    """
    Suggest tags from curated set (by ?category=) + recent PIP tags (comma lists), optional filter ?q=
    """
    q = (request.args.get('q') or '').strip().lower()
    category = (request.args.get('category') or '').strip()

    # Recent from DB
    try:
        recent_rows = db.session.query(PIPRecord.tags).order_by(PIPRecord.id.desc()).limit(200).all()
    except Exception:
        recent_rows = []

    recent = []
    for (tag_str,) in recent_rows:
        if not tag_str: continue
        for t in (tag_str.split(',') if isinstance(tag_str, str) else []):
            t = (t or '').strip()
            if t: recent.append(t)

    merged = _merge_curated_and_recent(category, recent, cap=40)
    if q:
        merged = [t for t in merged if q in t.lower()]

    return jsonify({"tags": merged[:30]})

# ----- Export Data -----
@app.route('/admin/export')
@login_required
@superuser_required
def export_data():
    zip_buffer = BytesIO()

    with tempfile.TemporaryDirectory() as tmpdir:
        def write_csv(filename, fieldnames, rows):
            filepath = os.path.join(tmpdir, filename)
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(rows)
            export_zip.write(filepath, arcname=filename)

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as export_zip:
            # Employees
            employees = Employee.query.all()
            write_csv(
                'employees.csv',
                ['id', 'first_name', 'last_name', 'job_title', 'line_manager', 'service', 'start_date', 'team_id', 'email'],
                [{
                    'id': e.id,
                    'first_name': getattr(e, 'first_name', ''),
                    'last_name': getattr(e, 'last_name', ''),
                    'job_title': e.job_title,
                    'line_manager': e.line_manager,
                    'service': e.service,
                    'start_date': e.start_date.strftime('%Y-%m-%d') if e.start_date else '',
                    'team_id': e.team_id,
                    'email': e.email
                } for e in employees]
            )

            # PIP Records
            pips = PIPRecord.query.all()
            write_csv(
                'pip_records.csv',
                ['id', 'employee_id', 'concerns', 'concern_category', 'severity', 'frequency', 'tags',
                 'start_date', 'review_date', 'status', 'created_by'],
                [{
                    'id': p.id,
                    'employee_id': p.employee_id,
                    'concerns': p.concerns,
                    'concern_category': getattr(p, 'concern_category', ''),
                    'severity': getattr(p, 'severity', ''),
                    'frequency': getattr(p, 'frequency', ''),
                    'tags': getattr(p, 'tags', ''),
                    'start_date': p.start_date.strftime('%Y-%m-%d') if p.start_date else '',
                    'review_date': p.review_date.strftime('%Y-%m-%d') if p.review_date else '',
                    'status': p.status,
                    'created_by': getattr(p, 'created_by', '')
                } for p in pips]
            )

            # Timeline Events
            events = TimelineEvent.query.all()
            write_csv(
                'timeline_events.csv',
                ['id', 'pip_record_id', 'employee_id', 'event_type', 'notes', 'updated_by', 'timestamp'],
                [{
                    'id': t.id,
                    'pip_record_id': t.pip_record_id if t.pip_record_id else '',
                    'employee_id': t.pip_record.employee_id if t.pip_record else '',
                    'event_type': getattr(t, 'event_type', ''),
                    'notes': getattr(t, 'notes', ''),
                    'updated_by': getattr(t, 'updated_by', ''),
                    'timestamp': t.timestamp.strftime('%Y-%m-%d %H:%M:%S') if t.timestamp else ''
                } for t in events]
            )

            # Users
            users = User.query.all()
            write_csv(
                'users.csv',
                ['id', 'username', 'email', 'admin_level', 'team_id'],
                [{
                    'id': u.id,
                    'username': u.username,
                    'email': u.email,
                    'admin_level': u.admin_level,
                    'team_id': u.team_id
                } for u in users]
            )

            # Probation Records
            probations = ProbationRecord.query.all()
            write_csv(
                'probation_records.csv',
                ['id', 'employee_id', 'status', 'start_date', 'expected_end_date', 'notes'],
                [{
                    'id': p.id,
                    'employee_id': p.employee_id,
                    'status': p.status,
                    'start_date': p.start_date.strftime('%Y-%m-%d') if p.start_date else '',
                    'expected_end_date': p.expected_end_date.strftime('%Y-%m-%d') if p.expected_end_date else '',
                    'notes': p.notes
                } for p in probations]
            )

            # Probation Reviews
            reviews = ProbationReview.query.all()
            write_csv(
                'probation_reviews.csv',
                ['id', 'probation_id', 'review_date', 'reviewer', 'summary', 'concerns_flag'],
                [{
                    'id': r.id,
                    'probation_id': r.probation_id,
                    'review_date': r.review_date.strftime('%Y-%m-%d') if r.review_date else '',
                    'reviewer': r.reviewer,
                    'summary': r.summary,
                    'concerns_flag': r.concerns_flag
                } for r in reviews]
            )

            # Probation Plans
            plans = ProbationPlan.query.all()
            write_csv(
                'probation_plans.csv',
                ['id', 'probation_id', 'objectives', 'outcome', 'deadline'],
                [{
                    'id': p.id,
                    'probation_id': p.probation_id,
                    'objectives': getattr(p, 'objectives', ''),
                    'outcome': getattr(p, 'outcome', ''),
                    'deadline': p.deadline.strftime('%Y-%m-%d') if getattr(p, 'deadline', None) else ''
                } for p in plans]
            )

    zip_buffer.seek(0)
    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M')
    return send_file(zip_buffer, mimetype='application/zip', as_attachment=True, download_name=f'export_{timestamp}.zip')

# ----- User Management -----
@app.route('/admin/users')
@login_required
def manage_users():
    if not current_user.is_superuser():
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    users = User.query.all()
    return render_template('admin_users.html', users=users)

@app.route('/admin/users/edit/<int:user_id>', methods=['GET', 'POST'])
@login_required
def edit_user(user_id):
    if not current_user.is_superuser():
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))

    user = User.query.get_or_404(user_id)
    form = UserForm(obj=user)

    if form.validate_on_submit():
        user.username = form.username.data
        user.email = form.email.data
        user.admin_level = form.admin_level.data
        user.team_id = form.team_id.data
        db.session.commit()
        flash('User updated successfully.', 'success')
        return redirect(url_for('manage_users'))

    return render_template('edit_user.html', form=form, user=user)

@app.route('/admin/users/create', methods=['GET', 'POST'])
@login_required
def create_user():
    if not current_user.is_superuser():
        flash("Access denied: Superuser only.", "danger")
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        admin_level = int(request.form.get('admin_level', 0))
        team_id = request.form.get('team_id') or None

        if not username or not email or not password:
            flash("All fields except team ID are required.", "danger")
            return redirect(request.url)

        if User.query.filter_by(username=username).first():
            flash("Username already exists.", "danger")
            return redirect(request.url)

        if User.query.filter_by(email=email).first():
            flash("Email already in use.", "danger")
            return redirect(request.url)

        hashed_pw = generate_password_hash(password)
        new_user = User(username=username, email=email, password_hash=hashed_pw, admin_level=admin_level, team_id=team_id)
        db.session.add(new_user)
        db.session.commit()
        flash("User created successfully.", "success")
        return redirect(url_for('manage_users'))

    return render_template('admin_create_user.html')

@app.route('/admin/users/delete/<int:user_id>', methods=['POST'])
@login_required
def delete_user(user_id):
    if not current_user.is_superuser():
        flash("Access denied: Superuser only.", "danger")
        return redirect(url_for('dashboard'))

    user = User.query.get_or_404(user_id)

    if user.id == current_user.id:
        flash("You cannot delete your own account while logged in.", "warning")
        return redirect(url_for('manage_users'))

    db.session.delete(user)
    db.session.commit()
    flash("User deleted successfully.", "success")
    return redirect(url_for('manage_users'))

@app.route('/admin/backup')
@login_required
def backup_database():
    if not current_user.is_superuser():
        flash('Access denied: Superuser only.', 'danger')
        return redirect(url_for('home'))

    db_path = os.path.join(os.getcwd(), 'pip_crm.db')
    if os.path.exists(db_path):
        return send_file(db_path, as_attachment=True)
    else:
        flash('Database file not found.', 'danger')
        return redirect(url_for('admin_dashboard'))
# ----- Employee & PIP management -----
@app.route('/employee/<int:employee_id>')
@login_required
def employee_detail(employee_id):
    employee = Employee.query.options(db.joinedload(Employee.pips)).get_or_404(employee_id)
    if current_user.admin_level == 0 and employee.team_id != current_user.team_id:
        flash('Access denied')
        return redirect(url_for('dashboard'))
    return render_template('employee_detail.html', employee=employee)

@app.route('/pip/<int:id>')
@login_required
def pip_detail(id):
    pip = PIPRecord.query.get_or_404(id)
    employee = pip.employee
    return render_template('pip_detail.html', pip=pip, employee=employee)

@app.route('/pip/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_pip(id):
    pip = PIPRecord.query.get_or_404(id)
    employee = pip.employee

    if request.method == 'POST':
        form = PIPForm()
        form.process(request.form)
    else:
        form = PIPForm(obj=pip)
        for _ in range(len(pip.action_items) - len(form.actions.entries)):
            form.actions.append_entry()
        for idx, ai in enumerate(pip.action_items):
            form.actions.entries[idx].form.description.data = ai.description
            form.actions.entries[idx].form.status.data = ai.status

    advice_text = None

    # AI Advice branch
    if request.method == 'POST' and 'generate_advice' in request.form:
        prompt = (
            f"You are a performance coach.\n"
            f"Employee: {employee.first_name} {employee.last_name}\n"
            f"Job Title: {employee.job_title}\n"
            f"Concerns: {form.concerns.data or '[none]'}\n"
            "Action Items:\n"
        )
        for ai_field in form.actions.entries:
            desc = ai_field.form.description.data or '[no description]'
            stat = ai_field.form.status.data or '[no status]'
            prompt += f"- {desc} [{stat}]\n"
        prompt += f"Meeting Notes: {form.meeting_notes.data or '[none]'}\n"
        prompt += "Provide 3 bulleted actionable tips for the manager to support this employee."

        resp = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        advice_text = resp.choices[0].message.content.strip()

        return render_template('edit_pip.html', form=form, pip=pip, employee=employee, advice_text=advice_text)

    # Save branch
    if form.validate_on_submit():
        pip.concerns = form.concerns.data
        pip.start_date = form.start_date.data
        pip.review_date = form.review_date.data
        pip.status = form.status.data
        pip.meeting_notes = form.meeting_notes.data
        pip.capability_meeting_date = form.capability_meeting_date.data
        pip.capability_meeting_time = form.capability_meeting_time.data
        pip.capability_meeting_venue = form.capability_meeting_venue.data

        pip.action_items.clear()
        for ai_field in form.actions.entries:
            pip.action_items.append(
                PIPActionItem(
                    description=ai_field.form.description.data,
                    status=ai_field.form.status.data
                )
            )

        db.session.commit()
        flash('PIP updated successfully.', 'success')
        return redirect(url_for('pip_detail', id=pip.id))

    # Final render
    return render_template('edit_pip.html', form=form, pip=pip, employee=employee, advice_text=advice_text)

@app.route('/pip/<int:id>/generate/advice', methods=['POST'])
@login_required
def generate_ai_advice(id):
    pip = PIPRecord.query.get_or_404(id)
    employee = pip.employee

    prompt = (
        "You are an experienced HR Line Manager based in the UK. "
        "Using the information below, provide three clear and practical suggestions for how a line manager "
        "can support the employee's performance improvement. Your advice should reflect HR best practice in the UK.\n\n"
        f"Employee Name: {employee.first_name} {employee.last_name}\n"
        f"Job Title: {employee.job_title}\n"
        f"Concerns: {pip.concerns or '[none]'}\n"
        "Action Items:\n"
    )
    for item in pip.action_items:
        prompt += f"- {item.description or '[no description]'} [{item.status or '[no status]'}]\n"
    prompt += f"Meeting Notes: {pip.meeting_notes or '[none]'}\n\n"
    prompt += "Provide your advice as a bullet-pointed list."

    resp = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
    )
    pip.ai_advice = resp.choices[0].message.content.strip()
    pip.ai_advice_generated_at = datetime.now(timezone.utc)

    event = TimelineEvent(
        pip_record_id=pip.id,
        event_type="AI Advice Generated",
        notes="Advice generated using OpenAI",
        updated_by=current_user.username
    )
    db.session.add(event)

    db.session.commit()
    flash('AI advice generated.', 'success')
    return redirect(url_for('pip_detail', id=pip.id))

@app.route('/pip/create/<int:employee_id>', methods=['GET', 'POST'])
@login_required
def create_pip(employee_id):
    employee = Employee.query.get_or_404(employee_id)
    if current_user.admin_level == 0 and employee.team_id != current_user.team_id:
        flash('Access denied.')
        return redirect(url_for('dashboard'))

    form = PIPForm()

    # Append empty action entry on GET to prevent Jinja error
    if request.method == 'GET' and len(form.actions.entries) == 0:
        form.actions.append_entry()

    # Recalculate min_entries for dynamic action field JS handling
    if request.method == 'POST':
        action_fields = [k for k in request.form if 'actions-' in k and '-description' in k]
        form.actions.min_entries = len(set(k.split('-')[1] for k in action_fields))

    if form.validate_on_submit():
        pip = PIPRecord(
            employee_id=employee.id,
            concerns=form.concerns.data,
            start_date=form.start_date.data,
            review_date=form.review_date.data,
            meeting_notes=form.meeting_notes.data
        )
        db.session.add(pip)
        db.session.flush()

        for action_form in form.actions.entries:
            item = PIPActionItem(
                pip_record_id=pip.id,
                description=action_form.form.description.data,
                status=action_form.form.status.data
            )
            db.session.add(item)

        db.session.commit()
        flash('New PIP created.')
        return redirect(url_for('employee_detail', employee_id=employee.id))

    return render_template('create_pip.html', form=form, employee=employee)

@app.route('/employee/edit/<int:employee_id>', methods=['GET', 'POST'])
@login_required
def edit_employee(employee_id):
    employee = Employee.query.get_or_404(employee_id)

    if current_user.admin_level == 0 and employee.team_id != current_user.team_id:
        flash('Access denied.')
        return redirect(url_for('dashboard'))

    form = EmployeeForm(obj=employee)

    if form.validate_on_submit():
        employee.first_name = form.first_name.data
        employee.last_name = form.last_name.data
        employee.job_title = form.job_title.data
        employee.line_manager = form.line_manager.data
        employee.service = form.service.data
        employee.start_date = form.start_date.data
        employee.team_id = form.team_id.data
        employee.email = form.email.data

        db.session.commit()
        flash('Employee details updated.', 'success')
        return redirect(url_for('employee_detail', employee_id=employee.id))

    return render_template('edit_employee.html', form=form, employee=employee)

@app.route('/pip/list')
@login_required
def pip_list():
    pips = PIPRecord.query.join(Employee).all()
    return render_template('pip_list.html', pips=pips)

@app.route('/pip/select-employee', methods=['GET', 'POST'])
@login_required
def select_employee_for_pip():
    employees = Employee.query.order_by(Employee.last_name).all()
    if request.method == 'POST':
        return redirect(url_for('create_pip', employee_id=request.form.get('employee_id')))
    return render_template('pip_select_employee.html', employees=employees)


@app.route('/taxonomy/action_templates', methods=['GET'])
@login_required
def taxonomy_action_templates():
    """
    Returns curated action plan items for Step 5 based on category & severity.
    Query: ?category=Timekeeping&severity=High
    """
    category = (request.args.get('category') or '').strip()
    severity = (request.args.get('severity') or '').strip()
    items = _pick_actions_from_templates(category, severity)
    return jsonify({"category": category, "severity": severity, "items": items})


# -------- Wizard entry ----------
class DummyForm(FlaskForm):
    pass

def _max_wizard_step(data: dict) -> int:
    """
    Gate progression. Only unlock Step 6 (Review) when Action Items exist.
    """
    s = 1
    if data.get('employee_id'): s = 2
    if all(data.get(k) for k in ('concerns', 'concern_category', 'severity', 'frequency')): s = 3
    if all(data.get(k) for k in ('start_date', 'review_date')): s = 4
    if all(data.get(k) for k in ('capability_meeting_date', 'capability_meeting_time', 'capability_meeting_venue')): s = 5
    # Only allow review (6) if at least one action item is present
    items = data.get('action_plan_items') or []
    if s == 5 and isinstance(items, list) and any((x or '').strip() for x in items):
        s = 6
    return s

@app.route('/pip/create-wizard', methods=['GET', 'POST'])
@login_required
def create_pip_wizard():
    if 'wizard_step' not in session:
        session['wizard_step'] = 1
        session['pip_data'] = {}

    step = session['wizard_step']
    data = session.get('pip_data', {}) or {}
    wizard_errors = {}
    draft = None  # Enhance later: load a real draft by draft_id

    # --- Clickable stepper: handle ?goto=N (allow jumps only up to the max allowed) ---
    if request.method == 'GET':
        try:
            goto = int(request.args.get('goto', 0))
        except (TypeError, ValueError):
            goto = 0
        if goto:
            max_allowed = _max_wizard_step(data)
            if 1 <= goto <= max_allowed:
                session['wizard_step'] = goto
                step = goto

        # If we are viewing step 3 and we already have a start_date but no review_date,
        # pre-populate review_date = start_date + 28 days and set flags for the banner.
        if step == 3 and data.get('start_date') and not data.get('review_date'):
            auto_val = _auto_review_date(data.get('start_date'))
            if auto_val:
                data['review_date'] = auto_val
                data['auto_review_populated'] = True
                data['auto_review_date'] = auto_val
                session['pip_data'] = data

    if request.method == 'POST':
        if step == 1:
            employee_id = request.form.get('employee_id')
            draft_name = request.form.get('draft_name', '').strip()
            if not employee_id:
                wizard_errors['employee_id'] = "Please select an employee."
            else:
                data['employee_id'] = int(employee_id)
            data['draft_name'] = draft_name

        elif step == 2:
            concerns = request.form.get('concerns', '').strip()
            category = request.form.get('concern_category', '').strip()
            severity = request.form.get('severity', '').strip()
            frequency = request.form.get('frequency', '').strip()
            tags = request.form.get('concern_tags', '').strip()
            draft_name = request.form.get('draft_name', '').strip()

            if not concerns:
                wizard_errors['concerns'] = "Concerns cannot be empty."
            if not category:
                wizard_errors['concern_category'] = "Please choose a concern category."
            if not severity:
                wizard_errors['severity'] = "Please select severity."
            if not frequency:
                wizard_errors['frequency'] = "Please select frequency."

            data.update({
                'concerns': concerns,
                'concern_category': category,
                'severity': severity,
                'frequency': frequency,
                'concern_tags': tags,
                'draft_name': draft_name,
            })

        elif step == 3:
            start_date = (request.form.get('start_date') or '').strip()
            review_date = (request.form.get('review_date') or '').strip()
            draft_name = request.form.get('draft_name', '').strip()
            # Optional: capture weeks control if present in the template
            review_weeks = (request.form.get('review_weeks') or '').strip()

            if not start_date:
                wizard_errors['start_date'] = "Start date is required."

            # If review date is missing but we have a start date, auto-calc +28 days
            auto_flag = False
            if start_date and not review_date:
                auto_val = _auto_review_date(start_date)
                if auto_val:
                    review_date = auto_val
                    auto_flag = True

            if not wizard_errors:
                data['start_date'] = start_date
                data['review_date'] = review_date
                # store the chosen weeks if provided; default to 4
                if review_weeks.isdigit():
                    data['review_weeks'] = int(review_weeks)
                elif 'review_weeks' not in data:
                    data['review_weeks'] = 4
                # flags for Step 3 banner
                data['auto_review_populated'] = bool(auto_flag)
                data['auto_review_date'] = review_date if auto_flag else None

            data['draft_name'] = draft_name

        elif step == 4:
            data['capability_meeting_date'] = request.form.get('capability_meeting_date')
            data['capability_meeting_time'] = request.form.get('capability_meeting_time')
            data['capability_meeting_venue'] = request.form.get('capability_meeting_venue')
            data['draft_name'] = request.form.get('draft_name', '').strip()

        elif step == 5:
            # Step 5 validates and stores items; DB commit happens on Step 6
            action_items = request.form.getlist('action_plan_items[]')
            valid_items = [item.strip() for item in action_items if item.strip()]
            if not valid_items:
                wizard_errors['action_plan_items'] = "Add at least one action plan item."
            else:
                data['action_plan_items'] = valid_items
                session['pip_data'] = data
                session['wizard_step'] = 6
                return redirect(url_for('create_pip_wizard'))

        elif step == 6:
            try:
                items = data.get('action_plan_items') or []
                if not any((x or '').strip() for x in items):
                    flash("Please add at least one action plan item.", "warning")
                    session['wizard_step'] = 5
                    return redirect(url_for('create_pip_wizard'))

                pip = PIPRecord(
                    employee_id=int(data['employee_id']),
                    concerns=data['concerns'],
                    concern_category=data.get('concern_category'),
                    severity=data.get('severity'),
                    frequency=data.get('frequency'),
                    tags=data.get('concern_tags'),
                    start_date=datetime.strptime(data['start_date'], '%Y-%m-%d').date(),
                    review_date=datetime.strptime(data['review_date'], '%Y-%m-%d').date(),
                    capability_meeting_date=datetime.strptime(data['capability_meeting_date'], '%Y-%m-%d')
                        if data.get('capability_meeting_date') else None,
                    capability_meeting_time=data.get('capability_meeting_time'),
                    capability_meeting_venue=data.get('capability_meeting_venue'),
                    created_by=current_user.username
                )

                db.session.add(pip)
                db.session.commit()

                for item_text in items:
                    action = PIPActionItem(pip_record_id=pip.id, description=item_text)
                    db.session.add(action)

                timeline = TimelineEvent(
                    pip_record_id=pip.id,
                    event_type="PIP Created",
                    notes="PIP created via multi-step wizard",
                    updated_by=current_user.username
                )
                db.session.add(timeline)

                db.session.commit()

                session.pop('wizard_step', None)
                session.pop('pip_data', None)

                flash("PIP created successfully!", "success")
                return redirect(url_for('dashboard'))

            except Exception as e:
                print(f"[ERROR] Failed to save PIP: {e}")
                flash(f"Error creating PIP: {str(e)}", "danger")

        # Persist to session after POST handling
        session['pip_data'] = data

        # Move to next step if no errors and not at final commit step (Step 6 handled above)
        if not wizard_errors and step < 5:
            session['wizard_step'] = step + 1
            return redirect(url_for('create_pip_wizard'))

    employees = Employee.query.all() if step == 1 else []

    # Compute allowed max step for template (clickable stepper)
    max_allowed = _max_wizard_step(data)

    return render_template(
        'create_pip_wizard.html',
        step=step,
        draft=draft,
        data=data,
        wizard_errors=wizard_errors,
        employees=employees,
        max_allowed_step=max_allowed,
        # expose flags for the Step 3 banner
        auto_review_populated=bool(data.get('auto_review_populated')),
        auto_review_date=data.get('auto_review_date')
    )
# ----- AJAX validation for wizard -----
@app.route('/validate_wizard_step', methods=['POST'])
@login_required
def validate_wizard_step():
    step = int(request.form.get("step", 1))
    errors = {}

    if step == 1:
        employee_id = request.form.get("employee_id")
        if not employee_id:
            errors["employee_id"] = "Please select an employee."

    elif step == 2:
        concerns = request.form.get("concerns", "").strip()
        category = request.form.get("concern_category", "").strip()
        tags = request.form.get("concern_tags", "").strip()
        severity = request.form.get("severity", "").strip()
        frequency = request.form.get("frequency", "").strip()

        if not concerns:
            errors["concerns"] = "Please describe the concern."
        if not category:
            errors["concern_category"] = "Please select a concern category."
        if not severity:
            errors["severity"] = "Please select severity."
        if not frequency:
            errors["frequency"] = "Please select frequency."

    elif step == 3:
        start_date = request.form.get("start_date", "").strip()
        review_date = request.form.get("review_date", "").strip()
        if not start_date:
            errors["start_date"] = "Please enter a start date."
    # We don’t require review_date here; server will auto-populate it if missing.


    elif step == 4:
        meeting_date = request.form.get("meeting_date", "").strip()
        meeting_time = request.form.get("meeting_time", "").strip()
        if not meeting_date:
            errors["meeting_date"] = "Please enter a meeting date."
        if not meeting_time:
            errors["meeting_time"] = "Please enter a meeting time."

    elif step == 5:
        actions = request.form.getlist("action_plan_items[]")
        actions = [a.strip() for a in actions if a.strip()]
        if not actions:
            errors["action_plan_items"] = "Please add at least one action item."

    return jsonify({"success": not errors, "errors": errors})


# Create draft from template → open editor
@app.route("/pip/<int:pip_id>/doc/create/<string:doc_type>", methods=["POST"])
@login_required
def create_pip_doc_draft(pip_id, doc_type):
    pip_rec = PIPRecord.query.get_or_404(pip_id)

    mapping = build_placeholder_mapping(pip_rec)

    template_map = {
        "invite": os.path.join(BASE_DIR, "templates", "docx", "PIP_Invite_Letter_Template_v2025-08-28.docx"),
        "plan":   os.path.join(BASE_DIR, "templates", "docx", "PIP_Action_Plan_Template_v2025-08-28.docx"),
        "outcome":os.path.join(BASE_DIR, "templates", "docx", "PIP_Outcome_Letter_Template_v2025-08-28.docx"),
    }
    if doc_type not in template_map:
        abort(404)

    outcome_choice = getattr(pip_rec, "outcome_choice", None)
    docx_bytes = generate_docx_bytes(template_map[doc_type], mapping, outcome_choice=outcome_choice)
    html = docx_to_html(docx_bytes)

    version = _next_version_for(pip_id, doc_type)
    rel_dir = build_doc_rel_dir(pip_id, doc_type, version)
    rel_docx_path = _save_file(docx_bytes, rel_dir, f"{doc_type}_v{version}.docx")

    doc = DocumentFile(
        pip_id=pip_id,
        doc_type=doc_type,
        version=version,
        status="draft",
        docx_path=rel_docx_path,
        html_snapshot=html,
        created_by=current_user.username,
    )
    db.session.add(doc)
    db.session.commit()

    log_timeline_event(
    pip_id=pip_id,
    event_type="Document Draft Created",
    notes=f"{doc.doc_type.capitalize()} v{doc.version} created from template.",
)


    flash(f"{doc_type.capitalize()} draft v{version} created.", "success")
    return redirect(url_for("edit_pip_doc", pip_id=pip_id, doc_id=doc.id))


# Edit (HTML) → Save back to DOCX
@app.route("/pip/<int:pip_id>/doc/<int:doc_id>/edit", methods=["GET", "POST"])
@login_required
def edit_pip_doc(pip_id, doc_id):
    pip_rec = PIPRecord.query.get_or_404(pip_id)
    doc = DocumentFile.query.filter_by(id=doc_id, pip_id=pip_id).first_or_404()

    if request.method == "POST":
        html = request.form.get("html", "")
        if not html:
            flash("No content received.", "warning")
            return redirect(request.url)

        new_docx = html_to_docx_bytes(html)
        rel_dir = build_doc_rel_dir(pip_id, doc.doc_type, doc.version)
        rel_docx_path = _save_file(new_docx, rel_dir, f"{doc.doc_type}_v{doc.version}_edited.docx")

        doc.html_snapshot = html
        doc.docx_path = rel_docx_path
        db.session.commit()
        html = sanitize_html(html)


        log_timeline_event(
        pip_id=pip_id,
        event_type="Document Draft Updated",
        notes=f"{doc.doc_type.capitalize()} v{doc.version} draft updated.",
        )


        flash("Draft updated.", "success")
        return redirect(request.url)

    return render_template("doc_editor.html", pip_rec=pip_rec, doc=doc, html_content=doc.html_snapshot)


# Finalise (lock)
@app.route("/pip/<int:pip_id>/doc/<int:doc_id>/finalise", methods=["POST"])
@login_required
def finalise_pip_doc(pip_id, doc_id):
    doc = DocumentFile.query.filter_by(id=doc_id, pip_id=pip_id).first_or_404()
    if doc.status == "final":
        flash("Document already final.", "info")
        return redirect(url_for("pip_documents", pip_id=pip_id))
    doc.status = "final"
    db.session.commit()
    log_timeline_event(
    pip_id=pip_id,
    event_type="Document Finalised",
    notes=f"{doc.doc_type.capitalize()} v{doc.version} finalised.",
)

    flash(f"{doc.doc_type.capitalize()} v{doc.version} finalised.", "success")
    return redirect(url_for("pip_documents", pip_id=pip_id))


# History view
@app.route("/pip/<int:pip_id>/documents", methods=["GET"])
@login_required
def pip_documents(pip_id):
    pip_rec = PIPRecord.query.get_or_404(pip_id)
    docs = (DocumentFile.query
            .filter_by(pip_id=pip_id)
            .order_by(DocumentFile.created_at.desc())
            .all())
    return render_template("pip_documents.html", pip_rec=pip_rec, docs=docs)


# Secure download
@app.route("/download/doc/<int:doc_id>")
@login_required
def download_doc(doc_id):
    doc = DocumentFile.query.get_or_404(doc_id)
    abs_dir = app.config['UPLOAD_FOLDER']
    dirname, filename = os.path.split(doc.docx_path)
    return send_from_directory(os.path.join(abs_dir, dirname), filename, as_attachment=True)



# ----- AI Action Suggestions -----
@app.route('/suggest_actions_ai', methods=['POST'])
@login_required
def suggest_actions_ai():
    """
    Returns JSON: { success, actions: [...], next_up: [...] }

    Now seeded with curated template items (category + severity) as a soft prior.
    """
    data = request.get_json() or {}
    concerns  = (data.get('concerns')  or '').strip()
    severity  = (data.get('severity')  or '').strip()
    frequency = (data.get('frequency') or '').strip()
    tags      = (data.get('tags')      or '').strip()
    category  = (data.get('category')  or '').strip()

    # ---- Curated template "soft prior" (category + severity)
    try:
        prior_actions = _pick_actions_from_templates(category, severity)  # from ACTION_TEMPLATES helper
    except Exception:
        prior_actions = []

    # small utility
    def _dedupe_clean(items, cap=None):
        out, seen = [], set()
        for x in (items or []):
            s = (x or "").strip()
            if not s:
                continue
            k = s.lower()
            if k not in seen:
                out.append(s)
                seen.add(k)
            if cap and len(out) >= cap:
                break
        return out

    # ----- Build a strict-JSON prompt seeded with template items -----
    sys_msg = (
        "You are an HR advisor in the UK.\n"
        "Return ONLY valid JSON with two arrays:\n"
        '{"actions": ["short concrete manager actions"], "next_up": ["quick follow-ups or escalations"]}.\n'
        "Actions must be specific, measurable where possible, supportive, and suitable for a PIP context.\n"
        "No prose, no markdown, JSON only."
    )

    # Include curated 'prior' suggestions to bias the model without forcing it
    # Model should consider them first, adapt or add more based on inputs.
    prior_block = ""
    if prior_actions:
        import json
        prior_block = "Seed actions (consider and adapt as appropriate): " + json.dumps(prior_actions, ensure_ascii=False)

    user_msg = f"""
Concern Category: {category or "[unspecified]"}
Concerns: {concerns or "[none]"}
Tags: {tags or "[none]"}
Severity: {severity or "[unspecified]"}
Frequency: {frequency or "[unspecified]"}

{prior_block}

Rules:
- Provide 3–5 'actions' tailored to the inputs.
- Provide 2–4 'next_up' items (e.g., monitoring cadence, policy references, escalation steps).
- Keep each item under 140 characters.
- JSON ONLY.
"""

    actions_llm, next_up_llm = [], []
    raw = ""

    try:
        resp = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": sys_msg},
                {"role": "user", "content": user_msg}
            ],
            temperature=0.5,
            max_tokens=300
        )
        raw = (resp.choices[0].message.content or "").strip()

        # Try strict JSON parse
        import json, re
        m = re.search(r"\{[\s\S]*\}", raw)
        payload = json.loads(m.group(0) if m else raw)

        actions_llm = payload.get("actions", []) or []
        next_up_llm = payload.get("next_up", []) or []

    except Exception:
        # Fallback: very rough parse if model didn't return valid JSON
        lines = [ln.strip("-•* 0123456789.\t") for ln in (raw.splitlines() if raw else [])]
        actions_llm = [ln for ln in lines if ln][:5]

    # ---- Merge with curated prior + server heuristics ----
    # Start with model → then weave in prior to ensure they don't get lost.
    merged_actions = _dedupe_clean(actions_llm, cap=None)
    if prior_actions:
        # place prior near the front while keeping dedupe
        merged_actions = _dedupe_clean(prior_actions + merged_actions, cap=8)

    # Server-side enrichment for 'next_up'
    next_up = _dedupe_clean(next_up_llm, cap=None)
    tag_list = [t.strip().lower() for t in tags.split(",")] if tags else []
    cat = (category or "").lower()
    sev = (severity or "").lower()
    freq = (frequency or "").lower()

    enrich = []
    if 'lateness' in tag_list or 'timekeeping' in cat:
        enrich += ["Daily start-time check-ins for 2 weeks", "Agree punctuality targets; log variances"]
    if 'conduct' in tag_list or cat == 'conduct':
        enrich += ["Reference conduct policy; document conversations", "Book values/behaviour refresher"]
    if 'performance' in cat or ('missed deadlines' in (tags or '').lower()):
        enrich += ["Weekly milestones with due dates", "Stand-up updates Mon/Wed/Fri"]
    if sev == 'high':
        enrich += ["Escalate to formal stage if no progress"]
    if freq in ('frequent', 'persistent'):
        enrich += ["Increase monitoring and assign a buddy/mentor"]

    next_up = _dedupe_clean(next_up + enrich, cap=8)

    # Ensure sensible caps
    merged_actions = merged_actions[:8] if merged_actions else []
    next_up = next_up[:8] if next_up else []

    return jsonify({
        "success": True,
        "actions": merged_actions,
        "next_up": next_up
    }), 200

# ----- Probation module -----

# ------------------------------
# Probation Wizard Routes
# ------------------------------

@app.route("/probation/create-wizard", methods=["GET"])
@login_required
def probation_create_wizard():
    # Try to load existing probation draft
    draft = DraftProbation.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    step = draft.step if draft else 1
    data = draft.payload if draft else {}

    return render_template(
        "probation_create_wizard.html",
        step=step,
        data=data,
        draft=draft
    )


@app.route("/probation/save-draft", methods=["POST"])
@login_required
def probation_save_draft():
    payload = request.json or {}
    step = payload.get("step", 1)

    draft = DraftProbation.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    if not draft:
        draft = DraftProbation(user_id=current_user.id, step=step, payload=payload)
        db.session.add(draft)
    else:
        draft.step = step
        draft.payload = payload
    db.session.commit()

    return jsonify({"success": True, "updated_at": draft.updated_at.strftime("%Y-%m-%d %H:%M:%S")})


@app.route("/probation/resume-draft")
@login_required
def probation_resume_draft():
    draft = DraftProbation.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    if draft:
        return redirect(url_for("probation_create_wizard"))
    flash("No probation draft available.", "info")
    return redirect(url_for("probation_dashboard"))


@app.route('/probation/<int:id>')
@login_required
def view_probation(id):
    probation = ProbationRecord.query.get_or_404(id)
    employee = probation.employee
    return render_template('view_probation.html', probation=probation, employee=employee)

@app.route('/probation/<int:id>/review/add', methods=['GET', 'POST'])
@login_required
def add_probation_review(id):
    probation = ProbationRecord.query.get_or_404(id)
    form = ProbationReviewForm()

    if form.validate_on_submit():
        review = ProbationReview(
            probation_id=probation.id,
            review_date=form.review_date.data,
            reviewer=form.reviewer.data,
            summary=form.summary.data,
            concerns_flag=(form.concerns_flag.data.lower() == 'yes')
        )
        db.session.add(review)

        event = TimelineEvent(
            pip_record_id=None,
            event_type="Probation Review",
            notes=f"Review added by {current_user.username}",
            updated_by=current_user.username
        )
        db.session.add(event)

        db.session.commit()
        flash('Probation review added.', 'success')
        return redirect(url_for('view_probation', id=probation.id))

    return render_template('add_probation_review.html', form=form, probation=probation)

@app.route('/probation/<int:id>/plan/add', methods=['GET', 'POST'])
@login_required
def add_probation_plan(id):
    probation = ProbationRecord.query.get_or_404(id)
    form = ProbationPlanForm()

    if form.validate_on_submit():
        plan = ProbationPlan(
            probation_id=probation.id,
            objectives=form.objectives.data,
            deadline=form.deadline.data,
            outcome=form.outcome.data
        )
        db.session.add(plan)

        event = TimelineEvent(
            pip_record_id=None,
            event_type="Probation Plan Added",
            notes=f"Plan created by {current_user.username}",
            updated_by=current_user.username
        )
        db.session.add(event)

        db.session.commit()
        flash('Development plan added.', 'success')
        return redirect(url_for('view_probation', id=probation.id))

    return render_template('add_probation_plan.html', form=form, probation=probation)

@app.route('/probation/<int:id>/edit', methods=['GET', 'POST'])
@login_required
def edit_probation(id):
    probation = ProbationRecord.query.get_or_404(id)
    form = ProbationRecordForm(obj=probation)

    if form.validate_on_submit():
        probation.start_date = form.start_date.data
        probation.expected_end_date = form.expected_end_date.data
        probation.notes = form.notes.data
        db.session.commit()
        flash('Probation record updated.', 'success')
        return redirect(url_for('view_probation', id=probation.id))

    return render_template('edit_probation.html', form=form, probation=probation)

@app.route('/probation/<int:id>/status/<new_status>', methods=['POST'])
@login_required
def update_probation_status(id, new_status):
    probation = ProbationRecord.query.get_or_404(id)
    valid_statuses = ['Completed', 'Extended', 'Failed']

    if new_status not in valid_statuses:
        flash('Invalid status update.', 'danger')
        return redirect(url_for('view_probation', id=id))

    probation.status = new_status
    db.session.add(probation)

    event = TimelineEvent(
        pip_record_id=None,
        event_type="Probation Status Updated",
        notes=f"Status changed to {new_status} by {current_user.username}",
        updated_by=current_user.username
    )
    db.session.add(event)

    db.session.commit()
    flash(f'Status updated to {new_status}.', 'success')
    return redirect(url_for('view_probation', id=id))

@app.route("/probation/create-wizard")
@login_required
def probation_create_wizard():
    employees = Employee.query.order_by(Employee.last_name).all()
    draft = DraftProbation.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    return render_template("create_probation_wizard.html", employees=employees, draft=draft)


@app.route("/probation/save_draft", methods=["POST"])
@login_required
def probation_save_draft():
    data = request.get_json() or {}
    draft_id = data.get("draft_id")
    step = data.get("step")
    payload = data.get("payload") or {}

    # ensure one draft per user
    draft = DraftProbation.query.filter_by(user_id=current_user.id).first()
    if not draft:
        draft = DraftProbation(user_id=current_user.id, step=step, payload={})
        db.session.add(draft)

    draft.step = step
    draft.payload[str(step)] = payload
    draft.updated_at = datetime.utcnow()
    db.session.commit()

    return jsonify({"ok": True, "draft_id": draft.id, "saved_at": draft.updated_at.isoformat()})


@app.route("/probation/validate_wizard_step", methods=["POST"])
@login_required
def probation_validate_wizard_step():
    data = request.get_json() or {}
    step = int(data.get("step", 0))
    payload = data.get("payload") or {}
    errors = {}

    if step == 1:
        if not payload.get("employee_id"):
            errors["employee_id"] = "Please select an employee."
        if not payload.get("probation_start"):
            errors["probation_start"] = "Start date required."
        if not payload.get("probation_end"):
            errors["probation_end"] = "End date required."

    elif step == 2:
        if not payload.get("concern_category"):
            errors["concern_category"] = "Category required."

    elif step == 3:
        if not payload.get("expected_standards"):
            errors["expected_standards"] = "Please define standards."

    elif step == 5:
        if not payload.get("reviewer"):
            errors["reviewer"] = "Reviewer is required."
        if not payload.get("review_dates"):
            errors["review_dates"] = "At least one review date required."

    return jsonify({"ok": not errors, "errors": errors})

@app.route('/probation/create/<int:employee_id>', methods=['GET', 'POST'])
@login_required
def create_probation(employee_id):
    employee = Employee.query.get_or_404(employee_id)
    form = ProbationRecordForm()

    if form.validate_on_submit():
        probation = ProbationRecord(
            employee_id=employee.id,
            start_date=form.start_date.data,
            expected_end_date=form.expected_end_date.data,
            notes=form.notes.data
        )
        db.session.add(probation)
        db.session.commit()
        flash('Probation record created successfully.', 'success')
        return redirect(url_for('view_probation', id=probation.id))

    return render_template('create_probation.html', form=form, employee=employee)

@app.route('/probation/dashboard')
@login_required
def probation_dashboard():
    # Global counters (keep these org-wide like PIP)
    global_active = ProbationRecord.query.filter_by(status='Active').count()
    global_completed = ProbationRecord.query.filter_by(status='Completed').count()
    global_extended = ProbationRecord.query.filter_by(status='Extended').count()

    # Dates
    today = datetime.now(timezone.utc).date()
    soon = today + timedelta(days=14)

    # Base query (scoped by team for line managers)
    q_records = ProbationRecord.query.join(Employee)
    q_reviews = ProbationReview.query.join(ProbationRecord, ProbationReview.probation_id == ProbationRecord.id).join(Employee)

    if current_user.admin_level == 0:
        # Line manager: only their team’s employees
        q_records = q_records.filter(Employee.team_id == current_user.team_id)
        q_reviews = q_reviews.filter(Employee.team_id == current_user.team_id)

    # Lists (scoped as above)
    active_probations = (
        q_records.filter(ProbationRecord.status == 'Active')
        .order_by(ProbationRecord.expected_end_date.asc().nullslast())
        .all()
    )

    upcoming_reviews = (
        q_reviews.filter(
            ProbationRecord.status == 'Active',
            ProbationReview.review_date >= today,
            ProbationReview.review_date <= soon
        )
        .order_by(ProbationReview.review_date.asc())
        .all()
    )

    overdue_reviews = (
        q_reviews.filter(
            ProbationRecord.status == 'Active',
            ProbationReview.review_date < today
        ).count()
    )

    due_soon_count = (
        q_records.filter(
            ProbationRecord.status == 'Active',
            ProbationRecord.expected_end_date >= today,
            ProbationRecord.expected_end_date <= soon
        ).count()
    )

    # Draft banner (probation wizard)
    probation_draft = get_active_probation_draft_for_user(current_user.id)

    return render_template(
        'probation_dashboard.html',
        active_module='Probation',
        # global counters
        global_active=global_active,
        global_completed=global_completed,
        global_extended=global_extended,
        # scoped lists/counters
        active_probations=active_probations,
        upcoming_reviews=upcoming_reviews,
        overdue_reviews=overdue_reviews,
        due_soon_count=due_soon_count,
        # draft
        draft=probation_draft
    )


# ================================
# Dismiss probation draft (for banner button)
# ================================
@app.route('/probation/dismiss-draft', methods=['POST'])
@login_required
def dismiss_probation_draft():
    draft = DraftProbation.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    if not draft:
        return jsonify({'success': False, 'message': 'No active draft'}), 400
    draft.is_dismissed = True
    db.session.commit()
    return jsonify({'success': True})

@app.route('/probation/employees')
@login_required
def probation_employee_list():
    session['active_module'] = 'Probation'
    employees = Employee.query.all()
    return render_template('probation_employee_list.html', employees=employees)

# ----- Draft helpers -----
def get_active_draft_for_user(user_id):
    return DraftPIP.query.filter_by(user_id=user_id, is_dismissed=False)\
                         .order_by(DraftPIP.updated_at.desc()).first()

@app.route('/dismiss_draft', methods=['POST'])
@login_required
@csrf.exempt
def dismiss_draft():
    draft = DraftPIP.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    if draft:
        draft.is_dismissed = True
        draft.updated_at = datetime.now(timezone.utc)
        db.session.commit()
        return jsonify({"success": True})
    return jsonify({"success": False, "message": "No active draft found"}), 404

@app.route('/save_pip_draft', methods=['POST'])
@login_required
@csrf.exempt
def save_pip_draft():
    try:
        data = {
            'employee_id': request.form.get('employee_id'),
            'draft_name': request.form.get('draft_name', '').strip(),
            'concerns': request.form.get('concerns', '').strip(),
            'concern_category': request.form.get('concern_category', '').strip(),
            'severity': request.form.get('severity', '').strip(),
            'frequency': request.form.get('frequency', '').strip(),
            'concern_tags': request.form.get('concern_tags', '').strip(),
            'start_date': request.form.get('start_date'),
            'review_date': request.form.get('review_date'),
            'capability_meeting_date': request.form.get('capability_meeting_date'),
            'capability_meeting_time': request.form.get('capability_meeting_time'),
            'capability_meeting_venue': request.form.get('capability_meeting_venue'),
            'action_plan_items': request.form.getlist('action_plan_items[]')
        }

        cleaned_data = {k: v for k, v in data.items() if v not in [None, '', []]}

        existing_draft = DraftPIP.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
        if existing_draft:
            db.session.delete(existing_draft)
            db.session.commit()

        new_draft = DraftPIP(
            user_id=current_user.id,
            data=cleaned_data,
            step=session.get('wizard_step', 1),
            is_dismissed=False,
            updated_at=datetime.now(timezone.utc)
        )
        db.session.add(new_draft)
        db.session.commit()

        return jsonify({"success": True, "message": "Draft saved."})

    except Exception as e:
        print(f"[ERROR] Failed to save draft: {e}")
        return jsonify({"success": False, "message": "Failed to save draft."}), 500

@app.route('/pip/wizard/resume', methods=['GET'])
@login_required
def pip_wizard_resume():
    """Hydrate the wizard session from the latest active DraftPIP and jump to the furthest valid step."""
    draft = get_active_draft_for_user(current_user.id)
    if not draft or not draft.data:
        flash("No active draft to resume.", "warning")
        return redirect(url_for('dashboard'))

    session['pip_data'] = dict(draft.data)
    session['wizard_step'] = _max_wizard_step(session['pip_data'])
    return redirect(url_for('create_pip_wizard'))

# ----- Dashboard -----
from flask_wtf.csrf import generate_csrf
@app.context_processor
def inject_csrf_token():
    return dict(csrf_token=generate_csrf)

@app.route('/dashboard')
@login_required
def dashboard():
    # Top-level counters (global, not team-filtered)
    total_employees = Employee.query.count()
    active_pips = PIPRecord.query.filter_by(status='Open').count()
    completed_pips = PIPRecord.query.filter_by(status='Completed').count()

    # Dates
    today = datetime.now(timezone.utc).date()
    upcoming_deadline = today + timedelta(days=7)

    if current_user.admin_level == 0:
        # Line manager: only see PIPs assigned to them in their team
        q_base = (
            PIPRecord.query
            .join(Employee)
            .filter(Employee.team_id == current_user.team_id)
            .filter(PIPRecord.assigned_to == current_user.id)
        )

        # Overdue reviews: assigned only
        overdue_reviews = (
            q_base.filter(
                PIPRecord.status == 'Open',
                PIPRecord.review_date < today
            ).count()
        )

        # Open PIPs
        open_pips = (
            q_base.filter(PIPRecord.status == 'Open')
            .order_by(PIPRecord.review_date.asc().nullslast())
            .all()
        )

        # Upcoming PIPs
        q_upcoming = q_base.filter(
            PIPRecord.status == 'Open',
            PIPRecord.review_date >= today,
            PIPRecord.review_date <= upcoming_deadline
        )
        upcoming_pips = q_upcoming.order_by(PIPRecord.review_date.asc()).all()
        due_soon_count = q_upcoming.count()

        # Recent activity: only events for their assigned PIPs
        recent_activity = (
            TimelineEvent.query
            .join(PIPRecord, TimelineEvent.pip_record_id == PIPRecord.id)
            .filter(PIPRecord.assigned_to == current_user.id)
            .order_by(TimelineEvent.timestamp.desc())
            .limit(10)
            .all()
        )

    else:
        # Admins/superusers: org-wide
        overdue_reviews = (
            PIPRecord.query
            .filter(
                PIPRecord.status == 'Open',
                PIPRecord.review_date < today
            ).count()
        )

        open_pips = (
            PIPRecord.query
            .filter(PIPRecord.status == 'Open')
            .join(Employee)
            .order_by(PIPRecord.review_date.asc().nullslast())
            .all()
        )

        upcoming_pips = (
            PIPRecord.query
            .filter(
                PIPRecord.status == 'Open',
                PIPRecord.review_date >= today,
                PIPRecord.review_date <= upcoming_deadline
            )
            .order_by(PIPRecord.review_date.asc())
            .all()
        )

        due_soon_count = (
            PIPRecord.query
            .filter(
                PIPRecord.status == 'Open',
                PIPRecord.review_date >= today,
                PIPRecord.review_date <= upcoming_deadline
            ).count()
        )

        recent_activity = (
            TimelineEvent.query
            .order_by(TimelineEvent.timestamp.desc())
            .limit(10)
            .all()
        )

    # Active draft banner
    draft = get_active_draft_for_user(current_user.id)

    return render_template(
        'dashboard.html',
        total_employees=total_employees,
        active_pips=active_pips,
        completed_pips=completed_pips,
        overdue_reviews=overdue_reviews,
        recent_activity=recent_activity,
        upcoming_pips=upcoming_pips,
        open_pips=open_pips,
        due_soon_count=due_soon_count,
        draft=draft
    )



@app.route('/dashboard/stats.json')
@login_required
def dashboard_stats_json():
    """
    Tiny JSON payload for dashboard charts.
    - by_category: counts of Open PIPs by concern_category
    - by_severity: counts of Open PIPs by severity
    - totals: quick summary used by UI (open, due_soon, overdue)
    """
    today = datetime.now(timezone.utc).date()
    upcoming_deadline = today + timedelta(days=7)

    # Base scoped query (open PIPs)
    base = _open_pips_scoped_query()

    # Counts
    by_category = _counts_by_field(PIPRecord.concern_category)
    by_severity = _counts_by_field(PIPRecord.severity)

    # Totals for header chips (scoped the same way)
    open_total = base.count()
    due_soon = base.filter(
        PIPRecord.review_date >= today,
        PIPRecord.review_date <= upcoming_deadline
    ).count()
    overdue = base.filter(PIPRecord.review_date < today).count()

    return jsonify({
        "by_category": by_category,
        "by_severity": by_severity,
        "totals": {
            "open": open_total,
            "due_soon": due_soon,
            "overdue": overdue
        }
    })

# ----- Employee add/list/quick-add -----
@app.route('/employee/add', methods=['GET', 'POST'])
@login_required
def add_employee():
    if current_user.admin_level < 1:
        flash('Access denied.')
        return redirect(url_for('home'))
    form = EmployeeForm()
    if form.validate_on_submit():
        emp = Employee(
            first_name=form.first_name.data,
            last_name=form.last_name.data,
            job_title=form.job_title.data,
            line_manager=form.line_manager.data,
            service=form.service.data,
            start_date=form.start_date.data,
            team_id=form.team_id.data,
            email=form.email.data
        )
        db.session.add(emp)
        db.session.commit()
        flash('New employee added.')
        return redirect(url_for('employee_list'))
    return render_template('add_employee.html', form=form)

@app.route('/employee/quick-add', methods=['POST'])
@login_required
@csrf.exempt
def quick_add_employee():
    """
    Minimal quick-add endpoint for use in wizard Step 1.
    Expects JSON: { first_name, last_name, role, service }
    """
    data = request.get_json(force=True, silent=True) or {}
    first = (data.get("first_name") or "").strip()
    last = (data.get("last_name") or "").strip()
    role = (data.get("role") or "").strip()
    service = (data.get("service") or "").strip()

    if not first or not last:
        return jsonify({"success": False, "error": "First and last name are required"}), 400

    emp = Employee(
        first_name=first,
        last_name=last,
        role=role,  # NOTE: kept as-is to avoid changing your models/routes
        service=service,
        manager_id=getattr(current_user, "id", None)
    )
    db.session.add(emp)
    db.session.commit()

    try:
        evt = TimelineEvent(
            event_type="Employee Created",
            notes="Employee created via Quick-Add in wizard",
            updated_by=current_user.username
        )
        db.session.add(evt)
        db.session.commit()
    except Exception:
        pass

    return jsonify({"success": True, "id": emp.id, "display_name": f"{emp.first_name} {emp.last_name}"})

@app.route('/employee/list')
@login_required
def employee_list():
    template = 'probation_employee_list.html' if session.get('active_module') == 'Probation' else 'employee_list.html'
    q = Employee.query
    if current_user.admin_level == 0:
        if current_user.team_id:
            q = q.filter(Employee.team_id == current_user.team_id)
        else:
            q = q.filter(False)
    employees = q.order_by(Employee.last_name.asc(), Employee.first_name.asc()).all()
    return render_template(template, employees=employees)

# ----- Employee Import: upload/validate/commit -----
def _suggest_mapping(headers):
    mapping = {}
    for h in headers or []:
        n = _normalize_header(h)
        if n in EMPLOYEE_FIELDS:
            mapping[h] = n
        elif n in ("firstname", "first"):
            mapping[h] = "first_name"
        elif n in ("lastname", "last", "surname"):
            mapping[h] = "last_name"
        elif n in ("mail", "email_address"):
            mapping[h] = "email"
        elif n in ("role", "position", "title", "job", "jobrole", "job_role", "jobtitle"):
            mapping[h] = "job_title"
        elif n in ("manager", "line_manager", "linemanager", "manager_name"):
            mapping[h] = "line_manager"
        elif n in ("team", "teamid", "team_id", "dept", "department"):
            mapping[h] = "team_id"
        else:
            mapping[h] = ""  # not mapped yet
    return mapping

@app.route("/employee/import", methods=["GET", "POST"])
@login_required
@superuser_required
def employee_import():
    if request.method == "GET":
        return render_template("employee_import.html")

    file = request.files.get("file")
    if not file or file.filename == "":
        abort(400, "No file uploaded")

    filename = secure_filename(file.filename)
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_EXTS:
        abort(400, f"Unsupported file type: .{ext}")

    file_bytes = file.read()

    if ext == "csv":
        headers, rows = _read_csv_bytes(file_bytes)
    else:
        headers, rows = _read_xlsx_bytes(file_bytes)

    normalised_headers = [{ "raw": h, "norm": _normalize_header(h) } for h in (headers or [])]

    tmp = tempfile.NamedTemporaryFile(prefix="emp_import_", suffix=f".{ext}", delete=False)
    tmp.write(file_bytes)
    tmp.flush()
    temp_id = tmp.name

    preview = rows[:10] if rows else []

    return jsonify({
        "temp_id": temp_id,
        "headers": headers,
        "headers_norm": normalised_headers,
        "suggested_mapping": _suggest_mapping(headers),
        "preview_rows": preview,
        "xlsx_enabled": XLSX_ENABLED
    }), 200

@app.route("/employee/import/validate", methods=["POST"])
@login_required
@superuser_required
def employee_import_validate():
    data = request.get_json(force=True, silent=True) or {}
    temp_id = data.get("temp_id")
    mapping = data.get("mapping") or {}
    unique_key = (data.get("unique_key") or "email").strip()

    if not temp_id:
        abort(400, "Missing temp_id")
    try:
        with open(temp_id, "rb") as f:
            file_bytes = f.read()
    except Exception:
        abort(400, "Invalid temp_id or temporary file expired")

    ext = temp_id.rsplit(".", 1)[-1].lower()
    headers, rows = (_read_csv_bytes(file_bytes) if ext == "csv" else _read_xlsx_bytes(file_bytes))

    mapped_rows = []
    unmapped_headers = []
    for h in headers or []:
        if not mapping.get(h):
            unmapped_headers.append(h)

    for r in rows:
        out = {}
        for h, v in r.items():
            field = mapping.get(h)
            if not field:
                continue
            if field == "start_date":
                out[field] = _try_parse_date(v)
            else:
                out[field] = (str(v).strip() if v is not None else None)
        mapped_rows.append(out)

    # Required fields check
    missing_required = []
    for idx, r in enumerate(mapped_rows, start=1):
        missing = [f for f in REQUIRED_FIELDS if not r.get(f)]
        if missing:
            missing_required.append({"row": idx, "missing": missing})

    # Duplicate checks (in-file)
    duplicates_in_file = []
    if unique_key:
        keys = unique_key.split(",")
        seen = set()
        for idx, r in enumerate(mapped_rows, start=1):
            key_tuple = tuple((r.get(k.strip()) or "").lower() for k in keys)
            if all(key_tuple):
                if key_tuple in seen:
                    duplicates_in_file.append({"row": idx, "key": key_tuple})
                else:
                    seen.add(key_tuple)

    # Duplicate checks (in DB)
    duplicates_in_db = []
    try:
        if unique_key == "email" and any(r.get("email") for r in mapped_rows):
            emails = list({r.get("email") for r in mapped_rows if r.get("email")})
            existing = set(e[0].lower() for e in db.session.query(Employee.email).filter(Employee.email.in_(emails)).all())
            for idx, r in enumerate(mapped_rows, start=1):
                em = (r.get("email") or "").lower()
                if em and em in existing:
                    duplicates_in_db.append({"row": idx, "email": r.get("email")})
        elif unique_key == "first_name,last_name":
            pairs = {( (r.get("first_name") or "").lower(), (r.get("last_name") or "").lower() )
                     for r in mapped_rows if r.get("first_name") and r.get("last_name")}
            if pairs:
                q = db.session.query(Employee.first_name, Employee.last_name).all()
                existing_pairs = {(fn.lower(), ln.lower()) for fn, ln in q}
                for idx, r in enumerate(mapped_rows, start=1):
                    t = ((r.get("first_name") or "").lower(), (r.get("last_name") or "").lower())
                    if all(t) and t in existing_pairs:
                        duplicates_in_db.append({"row": idx, "name": f"{r.get('first_name')} {r.get('last_name')}"})
    except Exception as e:
        duplicates_in_db = [{"error": f"DB duplicate check skipped: {e}"}]

    report = {
        "unmapped_headers": unmapped_headers,
        "missing_required": missing_required,
        "duplicates_in_file": duplicates_in_file,
        "duplicates_in_db": duplicates_in_db,
        "rows_ready": len(mapped_rows) - len(missing_required) - len(duplicates_in_file) - len(duplicates_in_db),
        "total_rows": len(mapped_rows)
    }
    return jsonify({"temp_id": temp_id, "report": report}), 200

@app.route("/employee/import/commit", methods=["POST"])
@login_required
@superuser_required
def employee_import_commit():
    data = request.get_json(force=True, silent=True) or {}
    temp_id = data.get("temp_id")
    mapping = data.get("mapping") or {}
    confirm = bool(data.get("confirm"))
    unique_key = (data.get("unique_key") or "email").strip()
    if not (temp_id and confirm and mapping):
        abort(400, "Missing temp_id, mapping, or confirm flag")

    try:
        with open(temp_id, "rb") as f:
            file_bytes = f.read()
    except Exception:
        abort(400, "Invalid temp_id or temporary file expired")

    ext = temp_id.rsplit(".", 1)[-1].lower()
    headers, rows = (_read_csv_bytes(file_bytes) if ext == "csv" else _read_xlsx_bytes(file_bytes))

    created, skipped, errors = 0, 0, []

    for idx, src in enumerate(rows, start=1):
        payload = {}
        for h, v in src.items():
            field = mapping.get(h)
            if not field:
                continue
            if field == "start_date":
                payload[field] = _try_parse_date(v)
            else:
                payload[field] = (str(v).strip() if v is not None else None)

        # Required fields
        if any(not payload.get(f) for f in REQUIRED_FIELDS):
            skipped += 1
            continue

        # Duplicate checks (quick)
        try:
            exists = False
            if unique_key == "email" and payload.get("email"):
                exists = db.session.query(Employee.id).filter_by(email=payload["email"]).first() is not None
            elif unique_key == "first_name,last_name" and payload.get("first_name") and payload.get("last_name"):
                exists = db.session.query(Employee.id).filter_by(
                    first_name=payload["first_name"], last_name=payload["last_name"]
                ).first() is not None
            if exists:
                skipped += 1
                continue
        except Exception as e:
            errors.append({"row": idx, "error": f"Duplicate check failed: {e}"})
            skipped += 1
            continue

        try:
            emp = Employee(**{k: v for k, v in payload.items() if k in EMPLOYEE_FIELDS})
            db.session.add(emp)
            created += 1
        except Exception as e:
            errors.append({"row": idx, "error": str(e)})
            skipped += 1

    db.session.commit()

    try:
        username = getattr(current_user, "username", "system")
        notes = f"Employee Import: created={created}, skipped={skipped}, errors={len(errors)}"
        evt = TimelineEvent(event_type="Import", notes=notes, updated_by=username)
        db.session.add(evt)
        db.session.commit()
    except Exception:
        pass

    return jsonify({"created": created, "skipped": skipped, "errors": errors}), 200

# ----- New DOCX-based document generation (Invite / Plan / Outcome) -----
@app.route('/pip/<int:id>/generate/invite')
@login_required
def generate_invite_letter(id):
    pip = PIPRecord.query.get_or_404(id)
    emp = pip.employee

    context = {
        "[[LETTER_DATE]]": datetime.now(ZoneInfo("Europe/London")).strftime("%d %B %Y"),
        "[[EMPLOYEE_NAME]]": f"{emp.first_name} {emp.last_name}",
        "[[EMPLOYEE_FIRST_NAME]]": emp.first_name or "",
        "[[JOB_TITLE]]": emp.job_title or "",
        "[[DEPARTMENT]]": emp.service or "",
        "[[MANAGER_NAME]]": pip.created_by or getattr(current_user, "username", ""),
        "[[MANAGER_TITLE]]": getattr(emp, "line_manager", "") or "",
        "[[MEETING_LOCATION]]": pip.capability_meeting_venue or "",
        "[[MEETING_DATETIME]]": (
            f"{(pip.capability_meeting_date.strftime('%d %B %Y') if pip.capability_meeting_date else '')} "
            f"{(pip.capability_meeting_time or '')}"
        ).strip(),
        "[[CONCERNS_SUMMARY]]": pip.concerns or "",
        "[[EVIDENCE_SUMMARY]]": "",
        "[[SUPPORT_LIST]]": "",
        "[[PIP_START_DATE]]": pip.start_date.strftime("%d %B %Y") if pip.start_date else "",
        "[[REVIEW_PERIOD_WEEKS]]": "",
        "[[HR_CONTACT_NAME]]": "",
        "[[HR_CONTACT_EMAIL]]": "",
    }

    buf = render_docx("PIP_Invite_Letter_Template_v2025-08-28.docx", context)
    filename = f"Invite_Letter_{emp.last_name}_{pip.id}.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/pip/<int:id>/generate/plan')
@login_required
def generate_plan_document(id):
    pip = PIPRecord.query.get_or_404(id)
    emp = pip.employee

    # Build per-objective values 1..5 (or loop dynamically if you support more)
    obj_ctx = {}
    for i, obj in enumerate(pip.objectives[:5] if hasattr(pip, "objectives") else pip.action_items[:5], start=1):
        # If you don’t have a separate Objectives model, map from action_items
        text = getattr(obj, "text", None) or getattr(obj, "description", "") or ""
        obj_ctx.update({
            f"[[OBJ{i}_TEXT]]": text,
            f"[[OBJ{i}_SUCCESS]]": getattr(obj, "success_criteria", "") or "",
            f"[[OBJ{i}_ACTIONS_EMPLOYEE]]": getattr(obj, "actions_employee", "") or "",
            f"[[OBJ{i}_ACTIONS_MANAGER]]": getattr(obj, "actions_manager", "") or "",
            f"[[OBJ{i}_SUPPORT]]": getattr(obj, "support", "") or "",
            f"[[OBJ{i}_DEADLINE]]": (obj.deadline.strftime("%d %B %Y") if getattr(obj, "deadline", None) else ""),
            f"[[OBJ{i}_METRICS]]": getattr(obj, "metrics", "") or "",
        })

    context = {
        "[[EMPLOYEE_NAME]]": f"{emp.first_name} {emp.last_name}",
        "[[JOB_TITLE]]": emp.job_title or "",
        "[[DEPARTMENT]]": emp.service or "",
        "[[MANAGER_NAME]]": pip.created_by or getattr(current_user, "username", ""),
        "[[REVIEWER_NAME]]": "",
        "[[HR_CONTACT_NAME]]": "",
        "[[PIP_START_DATE]]": pip.start_date.strftime("%d %B %Y") if pip.start_date else "",
        "[[PIP_END_DATE]]": pip.review_date.strftime("%d %B %Y") if pip.review_date else "",
        "[[REVIEW_PERIOD_WEEKS]]": "",
        "[[POLICY_REFERENCE]]": getattr(pip, "policy_reference", "") or "",
        "[[CONCERN_CATEGORIES]]": getattr(pip, "concern_category", "") or "",
        "[[CONCERN_TAGS]]": (pip.tags or ""),
        "[[SEVERITY_RATING]]": getattr(pip, "severity", "") or "",
        "[[FREQUENCY_RATING]]": getattr(pip, "frequency", "") or "",
        "[[CONFIDENTIAL_NOTES]]": getattr(pip, "confidential_notes", "") or "",
        "[[REVIEW_DATES_LIST]]": getattr(pip, "review_dates_str", "") or "",
        "[[REPORTING_METHOD]]": getattr(pip, "reporting_method", "") or "",
        "[[ADJUSTMENTS]]": getattr(pip, "adjustments", "") or "",
        "[[TRAINING_PLAN]]": getattr(pip, "training_plan", "") or "",
        **obj_ctx,
    }

    buf = render_docx("PIP_Action_Plan_Template_v2025-08-28.docx", context)
    filename = f"PIP_Plan_{emp.last_name}_{pip.id}.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/pip/<int:id>/generate/outcome')
@login_required
def generate_outcome_letter(id):
    pip = PIPRecord.query.get_or_404(id)
    emp = pip.employee

    # Decide which outcome to keep via query param: ?outcome=successful|extension|unsuccessful
    choice = (request.args.get("outcome", "SUCCESSFUL") or "SUCCESSFUL").upper()

    # Build per-objective outcomes (1..5)
    obj_ctx = {}
    for i, obj in enumerate(pip.objectives[:5] if hasattr(pip, "objectives") else pip.action_items[:5], start=1):
        text = getattr(obj, "text", None) or getattr(obj, "description", "") or ""
        obj_ctx.update({
            f"[[OBJ{i}_TEXT]]": text,
            f"[[OBJ{i}_OUTCOME]]": getattr(obj, "outcome_text", "") or "",
            f"[[OBJ{i}_EVIDENCE]]": getattr(obj, "outcome_evidence", "") or "",
        })

    context = {
        "[[LETTER_DATE]]": datetime.now(ZoneInfo("Europe/London")).strftime("%d %B %Y"),
        "[[EMPLOYEE_NAME]]": f"{emp.first_name} {emp.last_name}",
        "[[EMPLOYEE_FIRST_NAME]]": emp.first_name or "",
        "[[JOB_TITLE]]": emp.job_title or "",
        "[[DEPARTMENT]]": emp.service or "",
        "[[MANAGER_NAME]]": pip.created_by or getattr(current_user, "username", ""),
        "[[MANAGER_TITLE]]": getattr(emp, "line_manager", "") or "",
        "[[PIP_START_DATE]]": pip.start_date.strftime("%d %B %Y") if pip.start_date else "",
        "[[PIP_END_DATE]]": pip.review_date.strftime("%d %B %Y") if pip.review_date else "",
        "[[ONGOING_SUPPORT]]": getattr(pip, "ongoing_support", "") or "",
        "[[EXTENSION_WEEKS]]": str(getattr(pip, "extension_weeks", "") or ""),
        "[[HR_CONTACT_NAME]]": "",
        "[[HR_CONTACT_EMAIL]]": "",
        "[[APPENDIX_DETAIL]]": getattr(pip, "outcome_appendix", "") or "",
        **obj_ctx,
    }

    buf = render_docx("PIP_Outcome_Letter_Template_v2025-08-28.docx", context, outcome_choice=choice)
    filename = f"Outcome_Letter_{emp.last_name}_{pip.id}.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# ----- Create DB if missing -----
with app.app_context():
    if not os.path.exists(os.path.join(BASE_DIR, 'pip_crm.db')):
        db.create_all()
        print('✅ Database created')

@app.route('/ping')
def ping():
    return 'Pong!'

print('✅ Flask app initialized and ready.')

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
