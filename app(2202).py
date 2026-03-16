import os
import io
import re
import csv
import zipfile
import tempfile
import bleach
import mammoth

from io import BytesIO
from functools import wraps
from pathlib import Path

from sqlalchemy.sql import func
from sqlalchemy.exc import SQLAlchemyError

from flask import (
    Flask, session, render_template, redirect, url_for,
    request, flash, send_file, jsonify, abort, send_from_directory
)
from flask_migrate import Migrate
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, login_user, login_required,
    logout_user, current_user
)
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect, generate_csrf

from html2docx import html2docx
from dotenv import load_dotenv
load_dotenv()

from datetime import datetime, date, timedelta, timezone

# -------------------------------------------
# Timezone helpers (robust on Windows)
# -------------------------------------------
try:
    from zoneinfo import ZoneInfo, ZoneInfoNotFoundError
except ImportError:
    ZoneInfo = None
    ZoneInfoNotFoundError = Exception  # fallback type

if ZoneInfo is not None:
    try:
        LONDON_TZ = ZoneInfo("Europe/London")
    except Exception:
        # tzdata missing or zone not found – fall back to UTC
        LONDON_TZ = timezone.utc
else:
    # very old Python without zoneinfo – fall back to UTC
    LONDON_TZ = timezone.utc


def now_utc() -> datetime:
    """Timezone-aware 'now' in UTC."""
    return datetime.now(timezone.utc)


def now_local() -> datetime:
    """Timezone-aware 'now' in UK time (or UTC fallback)."""
    return datetime.now(LONDON_TZ)


def today_local() -> date:
    """Local 'today' date in UK time (or UTC fallback)."""
    return now_local().date()

def _clamp_date_range(start: date | None, end: date | None, window_start: date, window_end: date):
    """Return (clamped_start, clamped_end) or (None, None) if no overlap / invalid."""
    if not start:
        return None, None
    real_end = end or window_end

    # If the case ends before window starts OR starts after window end => no overlap
    if real_end < window_start or start > window_end:
        return None, None

    clamped_start = max(start, window_start)
    clamped_end = min(real_end, window_end)
    if clamped_end < clamped_start:
        return None, None
    return clamped_start, clamped_end

def _clamp_date_range(start: date | None, end: date | None, window_start: date, window_end: date):
    """
    Return (clamped_start, clamped_end) for overlap of [start, end] with [window_start, window_end].
    Returns (None, None) if there is no overlap or invalid inputs.
    """
    if not start:
        return None, None

    real_end = end or window_end

    # No overlap
    if real_end < window_start or start > window_end:
        return None, None

    clamped_start = max(start, window_start)
    clamped_end = min(real_end, window_end)

    if clamped_end < clamped_start:
        return None, None

    return clamped_start, clamped_end


def compute_sickness_trigger_metrics(
    q_cases,
    *,
    today: date,
    window_days: int = 365,
    bradford_medium: int = 200,
    bradford_high: int = 400,
    episodes_threshold: int = 3,
    total_days_threshold: int = 14,
    long_term_days: int = 28,
):
    """
    Phase 3 rolling trigger metrics per employee.

    Episode definition (Phase 3 default): each SicknessCase == 1 episode.
    Days: overlap days within the rolling window (clamped).
    Bradford: episodes^2 * total_days
    """
    window_start = today - timedelta(days=window_days)
    window_end = today

    # Pull all cases with start_date (we clamp overlap in Python)
    cases = (
        q_cases.filter(
            SicknessCase.start_date.isnot(None),
            SicknessCase.start_date <= window_end,
        )
        .all()
    )

    metrics_by_employee = {}

    for sc in cases:
        if not sc.employee or not sc.start_date:
            continue

        clamped_start, clamped_end = _clamp_date_range(sc.start_date, sc.end_date, window_start, window_end)
        if not clamped_start:
            continue

        days = (clamped_end - clamped_start).days + 1
        if days <= 0:
            continue

        data = metrics_by_employee.setdefault(
            sc.employee.id,
            {
                "employee": sc.employee,
                "episodes": 0,
                "total_days": 0,
                "has_long_term": False,
                "longest_spell_days": 0,
            },
        )

        data["episodes"] += 1
        data["total_days"] += days
        data["longest_spell_days"] = max(data["longest_spell_days"], days)

        if days >= long_term_days:
            data["has_long_term"] = True

    potential_triggers = []

    for data in metrics_by_employee.values():
        employee = data["employee"]
        episodes = data["episodes"]
        total_days = data["total_days"]
        has_long_term = data["has_long_term"]
        longest_spell_days = data["longest_spell_days"]

        bradford = (episodes * episodes * total_days) if total_days > 0 else 0

        flags = []
        actions = []

        if episodes >= episodes_threshold:
            flags.append(f"Episodes ≥ {episodes_threshold} in 12 months")
            actions.append("Review absence pattern and agree next steps (informal stage).")

        if total_days >= total_days_threshold:
            flags.append(f"≥ {total_days_threshold} days total in 12 months")
            actions.append("Check fit note / evidence and update absence records.")

        if has_long_term:
            flags.append(f"Long-term case (≥ {long_term_days} days)")
            actions.append("Consider OH referral and a welfare meeting plan.")

        severity = "none"
        if bradford >= bradford_high:
            flags.append(f"Bradford ≥ {bradford_high}")
            severity = "high"
            actions.append("Consider formal sickness stage (policy dependent) and document rationale.")
        elif bradford >= bradford_medium:
            flags.append(f"Bradford ≥ {bradford_medium}")
            severity = "medium"
            actions.append("Book an absence review and set review checkpoints.")
        elif flags:
            severity = "low"
            actions.append("Keep monitoring; ensure RTW notes are complete.")

        if not flags:
            continue

        # de-dupe actions (keep order)
        seen = set()
        actions_unique = []
        for a in actions:
            if a not in seen:
                seen.add(a)
                actions_unique.append(a)

        potential_triggers.append(
            {
                "employee": employee,
                "episodes": episodes,
                "total_days": total_days,
                "bradford": bradford,
                "has_long_term": has_long_term,
                "longest_spell_days": longest_spell_days,
                "flags_label": ", ".join(flags),
                "severity": severity,
                "actions": actions_unique,
            }
        )

    # Sort: high -> medium -> low, then bradford desc, then episodes desc
    severity_rank = {"high": 3, "medium": 2, "low": 1, "none": 0}
    potential_triggers.sort(
        key=lambda t: (severity_rank.get(t["severity"], 0), t["bradford"], t["episodes"]),
        reverse=True,
    )

    return potential_triggers

def compute_sickness_trigger_metrics(
    q_cases,
    today: date,
    window_days: int = 365,
    *,
    bradford_medium: int = 200,
    bradford_high: int = 400,
    episodes_threshold: int = 3,
    total_days_threshold: int = 14,
    long_term_days: int = 28,
):
    """
    Compute rolling-window sickness trigger metrics per employee.

    Episode definition (Phase 3 default): each SicknessCase == 1 episode.
    Days: overlap days within [today-window_days, today] window.
    """
    window_start = today - timedelta(days=window_days)
    window_end = today

    # Pull cases that overlap the window (not just those that started within it)
    recent_cases = (
        q_cases.filter(
            SicknessCase.start_date.isnot(None),
            SicknessCase.start_date <= window_end,  # started on/before today
        )
        .all()
    )

    metrics_by_employee = {}

    for case in recent_cases:
        employee = case.employee
        if not employee or not case.start_date:
            continue

        clamped_start, clamped_end = _clamp_date_range(case.start_date, case.end_date, window_start, window_end)
        if not clamped_start:
            continue

        days = (clamped_end - clamped_start).days + 1
        if days < 0:
            continue

        data = metrics_by_employee.setdefault(
            employee.id,
            {
                "employee": employee,
                "episodes": 0,
                "total_days": 0,
                "has_long_term": False,
                "longest_spell_days": 0,
            },
        )

        data["episodes"] += 1
        data["total_days"] += days
        data["longest_spell_days"] = max(data["longest_spell_days"], days)
        if days >= long_term_days:
            data["has_long_term"] = True

    potential_triggers = []

    for data in metrics_by_employee.values():
        employee = data["employee"]
        episodes = data["episodes"]
        total_days = data["total_days"]
        has_long_term = data["has_long_term"]
        longest_spell_days = data["longest_spell_days"]

        bradford = (episodes * episodes * total_days) if total_days > 0 else 0

        flags = []
        actions = []

        if episodes >= episodes_threshold:
            flags.append(f"Episodes ≥ {episodes_threshold} in 12 months")
            actions.append("Review absence pattern and agree next steps (informal stage).")

        if total_days >= total_days_threshold:
            flags.append(f"≥ {total_days_threshold} days total in 12 months")
            actions.append("Check fit note / medical evidence and update absence records.")

        if has_long_term:
            flags.append(f"Long-term case (≥ {long_term_days} days)")
            actions.append("Consider Occupational Health referral and a welfare meeting plan.")

        severity = "none"
        if bradford >= bradford_high:
            flags.append(f"Bradford ≥ {bradford_high}")
            severity = "high"
            actions.append("Consider formal sickness stage (policy dependent) and document rationale.")
        elif bradford >= bradford_medium:
            flags.append(f"Bradford ≥ {bradford_medium}")
            severity = "medium"
            actions.append("Book an absence review meeting and set review checkpoints.")
        elif flags:
            severity = "low"
            actions.append("Keep monitoring; ensure RTW notes are complete.")

        if not flags:
            continue

        # De-dupe action suggestions but keep order
        seen = set()
        actions_unique = []
        for a in actions:
            if a not in seen:
                seen.add(a)
                actions_unique.append(a)

        potential_triggers.append(
            {
                "employee": employee,
                "episodes": episodes,
                "total_days": total_days,
                "bradford": bradford,
                "has_long_term": has_long_term,
                "longest_spell_days": longest_spell_days,
                "flags_label": ", ".join(flags),
                "severity": severity,
                "actions": actions_unique,  # Phase 3: ready for UI
            }
        )

    potential_triggers.sort(key=lambda t: (t["severity"] == "high", t["severity"] == "medium", t["bradford"], t["episodes"]), reverse=True)

    return potential_triggers


# Use python-docx for templates
from docx import Document

# Forms (used across routes)
from forms import (
    LoginForm,
    EmployeeForm,
    PIPForm,
    ProbationRecordForm,
    ProbationReviewForm,
    ProbationPlanForm,
    UserForm,
    SicknessCaseForm,
    SicknessMeetingForm,
)

# Models (including sickness module models)
from models import (
    db,
    User,
    Employee,
    PIPRecord,
    PIPActionItem,
    TimelineEvent,
    ProbationRecord,
    ProbationReview,
    ProbationPlan,
    DraftPIP,
    DraftProbation,
    ImportJob,
    DocumentFile,
    SicknessCase,
    SicknessMeeting,
)

# Optional: .xlsx support
try:
    import openpyxl  # noqa: F401
    XLSX_ENABLED = True
except Exception:
    XLSX_ENABLED = False

ALLOWED_EXTS = {"csv", "xlsx"} if XLSX_ENABLED else {"csv"}

# ----- Employee import mapping -----
EMPLOYEE_FIELDS = [
    "first_name", "last_name", "email", "job_title", "line_manager",
    "service", "team_id", "start_date"
]
REQUIRED_FIELDS = ["first_name", "last_name"]

# ----- Curated tags & action templates -----
CURATED_TAGS = {
    "Timekeeping": ["lateness", "missed clock-in", "extended breaks", "early finish", "timekeeping policy"],
    "Attendance": ["unauthorised absence", "short notice absence", "patterns of absence", "fit note", "return to work"],
    "Quality of Work": ["accuracy", "attention to detail", "rework", "documentation", "SOP adherence"],
    "Productivity": ["missed deadlines", "slow throughput", "low output", "prioritisation", "time management"],
    "Conduct": ["inappropriate language", "unprofessional behaviour", "non-compliance", "policy breach", "conflict"],
    "Communication": ["tone", "late replies", "stakeholder updates", "handover quality", "listening"],
    "Teamwork/Collaboration": ["handover gaps", "knowledge sharing", "supporting peers", "collaboration tools"],
    "Compliance/Process": ["data entry errors", "checklist missed", "audit finding", "process deviation"],
    "Customer Service": ["response times", "complaint handling", "service standards", "follow-up"],
    "Health & Safety": ["PPE", "risk assessment", "manual handling", "incident reporting"]
}

def _merge_curated_and_recent(category: str, recent_tags: list[str], cap: int = 30):
    out, seen = [], set()
    for t in CURATED_TAGS.get(category, []):
        k = t.lower()
        if k not in seen:
            out.append(t)
            seen.add(k)
    for t in recent_tags:
        if not t:
            continue
        k = t.lower().strip()
        if k and k not in seen:
            out.append(t.strip())
            seen.add(k)
        if len(out) >= cap:
            break
    return out

ACTION_TEMPLATES = {
    "Timekeeping": {
        "Low": [
            "Agree start-time target and grace window",
            "Daily check-in for 2 weeks at start of shift",
            "Keep punctuality log; review weekly",
        ],
        "Moderate": [
            "Formal punctuality target with variance log",
            "Escalate if 2+ breaches in a week",
            "Buddy assigned for morning routine",
        ],
        "High": [
            "Issue written reminder citing policy",
            "Daily manager sign-off for 3 weeks",
            "Escalation to formal stage if breaches continue",
        ],
        "_default": [
            "Agree punctuality expectations",
            "Daily check-in for first 2 weeks",
            "Weekly review of log",
        ],
    },
    "Performance": {
        "Low": [
            "Break down tasks into weekly milestones",
            "Mid-week check-in with progress update",
            "Share example of quality standard",
        ],
        "Moderate": [
            "Set SMART targets per task",
            "Stand-up updates Mon/Wed/Fri",
            "Peer review before handoff",
        ],
        "High": [
            "Written performance targets with deadlines",
            "Daily status update for 10 working days",
            "Escalate to formal PIP stage if no improvement",
        ],
        "_default": [
            "Agree 2–3 SMART targets",
            "Weekly progress review",
            "Identify training/module to close gap",
        ],
    },
    "Conduct": {
        "_default": [
            "Reference conduct policy and expectations",
            "Agree behaviour standards; confirm by email",
            "Book values refresher session",
        ]
    },
    "Attendance": {
        "_default": [
            "Follow reporting procedure for absence",
            "Return-to-work meeting after each absence",
            "Pattern review after 4 weeks",
        ]
    },
    "Communication": {
        "_default": [
            "Acknowledge messages within agreed SLA",
            "Use agreed update template for stakeholders",
            "Add handover note at end of shift",
        ]
    },
    "Quality of Work": {
        "_default": [
            "Introduce checklist for critical steps",
            "Peer review for first 4 weeks",
            "Log defects and agree prevention steps",
        ]
    },
    "Productivity": {
        "_default": [
            "Time-block key tasks; share plan daily",
            "Weekly throughput targets",
            "Remove low-value tasks with manager",
        ]
    },
}

def _pick_actions_from_templates(category: str, severity: str) -> list[str]:
    cat = (category or "").strip()
    sev = (severity or "").strip()
    block = ACTION_TEMPLATES.get(cat) or {}
    if not block:
        block = {"_default": ["Agree clear targets", "Weekly review", "Training / buddy support as needed"]}
    if sev in block and block[sev]:
        return block[sev]
    if block.get("_default"):
        return block["_default"]
    for v in block.values():
        if isinstance(v, list) and v:
            return v
    return []

# ----- File readers -----
def _read_csv_bytes(file_bytes: bytes):
    text = file_bytes.decode("utf-8-sig", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    rows = [dict(r) for r in reader]
    headers = reader.fieldnames or []
    return headers, rows

def _read_xlsx_bytes(file_bytes: bytes):
    workbook = openpyxl.load_workbook(io.BytesIO(file_bytes))
    sheet = workbook.active
    headers = [c.value for c in next(sheet.iter_rows(min_row=1, max_row=1))]
    rows = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        rows.append({headers[i]: (row[i] if i < len(headers) else None) for i in range(len(headers))})
    return headers, rows

def _normalize_header(h):
    return (h or "").strip().lower().replace(" ", "_")

def _try_parse_date(value):
    if not value:
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(str(value).strip(), fmt).date()
        except Exception:
            continue
    return None

def _parse_iso_date(s):
    try:
        return datetime.strptime((s or "").strip(), "%Y-%m-%d").date()
    except Exception:
        return None

# ----- Flask init -----
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-only-secret')
# Make timezone helpers available inside Jinja templates
app.jinja_env.globals["now_utc"] = now_utc
app.jinja_env.globals["now_local"] = now_local
app.jinja_env.globals["today_local"] = today_local
BASE_DIR = os.path.abspath(os.path.dirname(__file__))

DB_PATH = os.path.join(BASE_DIR, 'pip_crm.db')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + DB_PATH
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads', 'documents')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

app.config['WTF_CSRF_ENABLED'] = True
app.config['WTF_CSRF_TIME_LIMIT'] = None
csrf = CSRFProtect(app)

from models import (
    db, User, Employee, PIPRecord, PIPActionItem, TimelineEvent,
    ProbationRecord, ProbationReview, ProbationPlan, DraftPIP, DraftProbation, DocumentFile
)
db.init_app(app)
migrate = Migrate(app, db)

# Import blueprints
from pip_app.blueprints.auth import auth_bp
from pip_app.blueprints.main import main_bp

# ----- Login manager -----
login_manager = LoginManager()
login_manager.login_view = 'auth.login'
login_manager.login_message_category = 'info'
login_manager.init_app(app)

# Register blueprints
app.register_blueprint(auth_bp)
app.register_blueprint(main_bp)

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))

# ----- OpenAI -----
from openai import OpenAI
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

# ----- Context processors -----
@app.context_processor
def inject_module():
    return dict(active_module=session.get('active_module'))

@app.context_processor
def inject_csrf_token():
    # expose a callable for templates {{ csrf_token() }}
    return dict(csrf_token=generate_csrf)

@app.context_processor
def inject_time_helpers():
    """
    Make timezone-aware helpers available inside Jinja templates.

    This prevents template errors like:
      jinja2.exceptions.UndefinedError: 'now_local' is undefined
    """
    return {
        "now_local": now_local,
        "today_local": today_local,
        "now_utc": now_utc,
    }
# ----- Superuser decorator -----
def superuser_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_superuser():
            abort(403)
        return f(*args, **kwargs)
    return decorated_function

# ----- Active module switcher -----
@app.before_request
def set_active_module():
    path = (request.path or "").lower()
    if path.startswith('/pip/') or path.startswith('/employee/') or path in (
        '/dashboard', '/pip_list', '/employee/list', '/employee/add', '/pip/select-employee'
    ):
        session['active_module'] = 'PIP'
    elif path.startswith('/probation/'):
        session['active_module'] = 'Probation'
    elif path == '/':
        session.pop('active_module', None)

# ===== Document Helpers (placeholders + conditionals) =====
TEMPLATE_DIR = Path(__file__).resolve().parent / "templates" / "docx"

LEGACY_TO_NEW_KEYS = {
    # Keep here if you still support legacy keys in context
    # "employee_name": "[[EMPLOYEE_NAME]]",
}

def _iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p

def _replace_in_runs(paragraph, mapping: dict):
    for run in paragraph.runs:
        text = run.text
        for k, v in mapping.items():
            if k in text:
                text = text.replace(k, v)
        run.text = text

def replace_placeholders_docx(doc: Document, context: dict):
    """
    Replace [[PLACEHOLDER]] tokens throughout the document, including
    tables, headers and footers, in a way that still works when Word
    splits placeholders across runs.

    - Uses a single mapping dict (including legacy keys).
    - Adds default [[GENERATED_DATE]] and [[DOC_VERSION]] if not supplied.
    - Strips any leftover [[UNUSED_PLACEHOLDERS]] so letters don't show raw tokens.
    """
    # Build mapping: normalise everything to strings
    base_ctx = context or {}
    mapping = {k: ("" if v is None else str(v)) for k, v in base_ctx.items()}

    # Support legacy context keys if you ever reintroduce them
    for legacy_key, new_key in LEGACY_TO_NEW_KEYS.items():
        if legacy_key in base_ctx and new_key not in mapping:
            v = base_ctx.get(legacy_key)
            mapping[new_key] = "" if v is None else str(v)

    # Sensible defaults if the caller didn't provide them
    now_uk = datetime.now()
    mapping.setdefault("[[GENERATED_DATE]]", now_uk.strftime("%d %B %Y"))
    mapping.setdefault("[[DOC_VERSION]]",  now_uk.strftime("v%Y.%m.%d"))

    # Precompile a regex to strip any remaining [[PLACEHOLDER_LIKE]] tokens
    placeholder_pattern = re.compile(r"\[\[[A-Z0-9_]+\]\]")

    for p in _iter_all_paragraphs(doc):
        if not p.runs:
            continue

        # Join all runs so we can replace across run boundaries
        original_text = "".join(run.text for run in p.runs)
        if not original_text:
            continue

        new_text = original_text

        # Apply all placeholder replacements
        for k, v in mapping.items():
            if k in new_text:
                new_text = new_text.replace(k, v)

        # Remove any leftover placeholder tokens so they don't bleed into the letter
        if "[[" in new_text:
            new_text = placeholder_pattern.sub("", new_text)

        # If nothing changed, skip rewriting runs
        if new_text == original_text:
            continue

        # Write updated text into the first run, blank the rest
        # (we sacrifice some per-word formatting in the paragraph, but
        # for template letters that's usually fine and MUCH more reliable
        # for placeholder replacement across runs).
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""


def strip_outcome_conditionals(doc: Document, keep: str):
    valid = {"SUCCESSFUL", "EXTENSION", "UNSUCCESSFUL"}
    choice = (keep or "").upper().strip()
    if choice not in valid:
        raise ValueError(f"Invalid outcome choice: {keep}")

    start_tokens = {f"[[IF_{tag}]]": tag for tag in valid}
    in_block = None
    keep_block = False
    to_delete = []

    paragraphs = list(doc.paragraphs)

    def contains_token(p, token):
        return any(token in r.text for r in p.runs) or token in p.text

    for p in paragraphs:
        for token, tag in start_tokens.items():
            if contains_token(p, token):
                in_block = tag
                keep_block = (tag == choice)
                for r in p.runs:
                    r.text = r.text.replace(token, "")
                if not keep_block:
                    to_delete.append(p)
                break

        if in_block:
            if not keep_block and p not in to_delete:
                to_delete.append(p)
            end_token = f"[[/IF_{in_block}]]"
            if contains_token(p, end_token):
                for r in p.runs:
                    r.text = r.text.replace(end_token, "")
                if not keep_block and p not in to_delete:
                    to_delete.append(p)
                in_block = None
                keep_block = False

    for p in to_delete:
        p._element.getparent().remove(p._element)

def render_docx(template_filename: str, context: dict, outcome_choice: str | None = None) -> BytesIO:
    template_path = TEMPLATE_DIR / template_filename
    if not template_path.exists():
        abort(404, f"Template not found: {template_path.name}")
    doc = Document(str(template_path))
    replace_placeholders_docx(doc, context or {})
    if outcome_choice:
        strip_outcome_conditionals(doc, outcome_choice)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ----- HTML sanitization for editor -----
ALLOWED_TAGS = bleach.sanitizer.ALLOWED_TAGS.union({
    "p","h1","h2","h3","h4","h5","h6",
    "strong","em","u","span","div","br","hr",
    "ul","ol","li",
    "table","thead","tbody","tr","th","td",
    "blockquote"
})
ALLOWED_ATTRS = {"*": ["class"]}

def sanitize_html(html: str) -> str:
    return bleach.clean(html, tags=ALLOWED_TAGS, attributes=ALLOWED_ATTRS, strip=True)

# ----- Timeline helper -----
def log_timeline_event(pip_id: int, event_type: str, notes: str):
    try:
        ev = TimelineEvent(
            pip_record_id=pip_id,
            event_type=event_type,
            notes=notes,
            updated_by=getattr(current_user, "username", None) or "system",
        )
        db.session.add(ev)
        db.session.commit()
    except Exception as e:
        app.logger.exception(f"TimelineEvent failed: {e}")

# ----- Probation draft banner helper -----
def get_active_probation_draft_for_user(user_id: int):
    return DraftProbation.query.filter_by(user_id=user_id, is_dismissed=False).first()

# ----- PIP stats helpers -----
def _open_pips_scoped_query():
    base = PIPRecord.query.filter(PIPRecord.status == 'Open')
    if current_user.admin_level == 0:
        base = base.join(Employee).filter(Employee.team_id == current_user.team_id)
    return base

def _counts_by_field(field_expr):
    q = _open_pips_scoped_query().with_entities(field_expr, func.count(PIPRecord.id)).group_by(field_expr)
    rows = q.all()
    out = {}
    for label, cnt in rows:
        out[(label or "Unspecified")] = int(cnt or 0)
    return out

# ----- Document store helpers -----
def _next_version_for(pip_id, doc_type):
    last = (DocumentFile.query
            .filter_by(pip_id=pip_id, doc_type=doc_type)
            .order_by(DocumentFile.version.desc())
            .first())
    return 1 if not last else last.version + 1

def _save_file(bytes_data: bytes, rel_dir: str, filename: str) -> str:
    dir_path = os.path.join(app.config['UPLOAD_FOLDER'], rel_dir)
    os.makedirs(dir_path, exist_ok=True)
    fpath = os.path.join(dir_path, secure_filename(filename))
    with open(fpath, 'wb') as f:
        f.write(bytes_data)
    return os.path.join(rel_dir, secure_filename(filename))

def generate_docx_bytes(template_path: str, mapping: dict, outcome_choice: str = None) -> bytes:
    doc = Document(template_path)
    replace_placeholders_docx(doc, mapping)
    if outcome_choice:
        strip_outcome_conditionals(doc, outcome_choice)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()

def docx_to_html(docx_bytes: bytes) -> str:
    with BytesIO(docx_bytes) as f:
        return mammoth.convert_to_html(f).value

def html_to_docx_bytes(html: str) -> bytes:
    out = BytesIO()
    html2docx(html, out)
    return out.getvalue()

# --- Wizard helper ---
def _auto_review_date(start_date_str: str | None) -> str | None:
    if not start_date_str:
        return None
    try:
        dt = datetime.strptime(start_date_str.strip(), "%Y-%m-%d").date()
        return (dt + timedelta(days=28)).isoformat()
    except Exception:
        return None
# =========================
# app.py — PART 2 / 4
# (Root/auth, taxonomy, export, user admin, employee & PIP basics)
# =========================

# ----- Root / Auth -----
@app.route('/admin/dashboard')
@login_required
def admin_dashboard():
    if not current_user.is_superuser():
        flash('You do not have permission to access the admin dashboard.', 'danger')
        return redirect(url_for('main.home'))
    return render_template('admin_dashboard.html')

# ----- Taxonomy (curated + categories + suggestions) -----
@app.route('/taxonomy/predefined_tags', methods=['GET'])
@login_required
def taxonomy_predefined_tags():
    cat = (request.args.get('category') or '').strip()
    tags = CURATED_TAGS.get(cat, [])
    return jsonify({"category": cat, "tags": tags})

@app.route('/taxonomy/categories', methods=['GET'])
@login_required
def taxonomy_categories():
    return jsonify({"categories": list(CURATED_TAGS.keys())})

@app.route('/taxonomy/tags_suggest', methods=['GET'])
@login_required
def taxonomy_tags_suggest():
    q = (request.args.get('q') or '').strip().lower()
    category = (request.args.get('category') or '').strip()
    try:
        recent_rows = db.session.query(PIPRecord.tags).order_by(PIPRecord.id.desc()).limit(200).all()
    except Exception:
        recent_rows = []
    recent = []
    for (tag_str,) in recent_rows:
        if not tag_str:
            continue
        for t in (tag_str.split(',') if isinstance(tag_str, str) else []):
            t = (t or '').strip()
            if t:
                recent.append(t)
    merged = _merge_curated_and_recent(category, recent, cap=40)
    if q:
        merged = [t for t in merged if q in t.lower()]
    return jsonify({"tags": merged[:30]})

# ----- Export Data -----
@app.route('/admin/export')
@login_required
@superuser_required
def export_data():
    zip_buffer = BytesIO()

    with tempfile.TemporaryDirectory() as tmpdir:
        def write_csv(filename, fieldnames, rows):
            filepath = os.path.join(tmpdir, filename)
            with open(filepath, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(rows)
            export_zip.write(filepath, arcname=filename)

        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as export_zip:
            # Employees
            employees = Employee.query.all()
            write_csv(
                'employees.csv',
                ['id', 'first_name', 'last_name', 'job_title', 'line_manager', 'service', 'start_date', 'team_id', 'email'],
                [{
                    'id': e.id,
                    'first_name': getattr(e, 'first_name', ''),
                    'last_name': getattr(e, 'last_name', ''),
                    'job_title': e.job_title,
                    'line_manager': e.line_manager,
                    'service': e.service,
                    'start_date': e.start_date.strftime('%Y-%m-%d') if e.start_date else '',
                    'team_id': e.team_id,
                    'email': e.email
                } for e in employees]
            )

            # PIP Records
            pips = PIPRecord.query.all()
            write_csv(
                'pip_records.csv',
                ['id', 'employee_id', 'concerns', 'concern_category', 'severity', 'frequency', 'tags',
                 'start_date', 'review_date', 'status', 'created_by'],
                [{
                    'id': p.id,
                    'employee_id': p.employee_id,
                    'concerns': p.concerns,
                    'concern_category': getattr(p, 'concern_category', ''),
                    'severity': getattr(p, 'severity', ''),
                    'frequency': getattr(p, 'frequency', ''),
                    'tags': getattr(p, 'tags', ''),
                    'start_date': p.start_date.strftime('%Y-%m-%d') if p.start_date else '',
                    'review_date': p.review_date.strftime('%Y-%m-%d') if p.review_date else '',
                    'status': p.status,
                    'created_by': getattr(p, 'created_by', '')
                } for p in pips]
            )

            # Timeline Events
            events = TimelineEvent.query.all()
            write_csv(
                'timeline_events.csv',
                ['id', 'pip_record_id', 'employee_id', 'event_type', 'notes', 'updated_by', 'timestamp'],
                [{
                    'id': t.id,
                    'pip_record_id': t.pip_record_id if t.pip_record_id else '',
                    'employee_id': t.pip_record.employee_id if t.pip_record else '',
                    'event_type': getattr(t, 'event_type', ''),
                    'notes': getattr(t, 'notes', ''),
                    'updated_by': getattr(t, 'updated_by', ''),
                    'timestamp': t.timestamp.strftime('%Y-%m-%d %H:%M:%S') if t.timestamp else ''
                } for t in events]
            )

            # Users
            users = User.query.all()
            write_csv(
                'users.csv',
                ['id', 'username', 'email', 'admin_level', 'team_id'],
                [{
                    'id': u.id,
                    'username': u.username,
                    'email': u.email,
                    'admin_level': u.admin_level,
                    'team_id': u.team_id
                } for u in users]
            )

            # Probation Records
            probations = ProbationRecord.query.all()
            write_csv(
                'probation_records.csv',
                ['id', 'employee_id', 'status', 'start_date', 'expected_end_date', 'notes'],
                [{
                    'id': p.id,
                    'employee_id': p.employee_id,
                    'status': p.status,
                    'start_date': p.start_date.strftime('%Y-%m-%d') if p.start_date else '',
                    'expected_end_date': p.expected_end_date.strftime('%Y-%m-%d') if p.expected_end_date else '',
                    'notes': p.notes
                } for p in probations]
            )

            # Probation Reviews
            reviews = ProbationReview.query.all()
            write_csv(
                'probation_reviews.csv',
                ['id', 'probation_id', 'review_date', 'reviewer', 'summary', 'concerns_flag'],
                [{
                    'id': r.id,
                    'probation_id': r.probation_id,
                    'review_date': r.review_date.strftime('%Y-%m-%d') if r.review_date else '',
                    'reviewer': r.reviewer,
                    'summary': r.summary,
                    'concerns_flag': r.concerns_flag
                } for r in reviews]
            )

            # Probation Plans
            plans = ProbationPlan.query.all()
            write_csv(
                'probation_plans.csv',
                ['id', 'probation_id', 'objectives', 'outcome', 'deadline'],
                [{
                    'id': p.id,
                    'probation_id': p.probation_id,
                    'objectives': getattr(p, 'objectives', ''),
                    'outcome': getattr(p, 'outcome', ''),
                    'deadline': p.deadline.strftime('%Y-%m-%d') if getattr(p, 'deadline', None) else ''
                } for p in plans]
            )

    zip_buffer.seek(0)
    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M')
    return send_file(zip_buffer, mimetype='application/zip', as_attachment=True, download_name=f'export_{timestamp}.zip')

# ----- User Management -----
@app.route('/admin/users')
@login_required
def manage_users():
    if not current_user.is_superuser():
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    users = User.query.all()
    return render_template('admin_users.html', users=users)

@app.route('/admin/users/edit/<int:user_id>', methods=['GET', 'POST'])
@login_required
def edit_user(user_id):
    if not current_user.is_superuser():
        flash('Access denied.', 'danger')
        return redirect(url_for('dashboard'))
    user = User.query.get_or_404(user_id)
    form = UserForm(obj=user)
    if form.validate_on_submit():
        user.username = form.username.data
        user.email = form.email.data
        user.admin_level = form.admin_level.data
        user.team_id = form.team_id.data
        db.session.commit()
        flash('User updated successfully.', 'success')
        return redirect(url_for('manage_users'))
    return render_template('edit_user.html', form=form, user=user)

@app.route('/admin/users/create', methods=['GET', 'POST'])
@login_required
def create_user():
    if not current_user.is_superuser():
        flash("Access denied: Superuser only.", "danger")
        return redirect(url_for('dashboard'))
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        admin_level = int(request.form.get('admin_level', 0))
        team_id = request.form.get('team_id') or None
        if not username or not email or not password:
            flash("All fields except team ID are required.", "danger")
            return redirect(request.url)
        if User.query.filter_by(username=username).first():
            flash("Username already exists.", "danger")
            return redirect(request.url)
        if User.query.filter_by(email=email).first():
            flash("Email already in use.", "danger")
            return redirect(request.url)
        hashed_pw = generate_password_hash(password)
        new_user = User(username=username, email=email, password_hash=hashed_pw, admin_level=admin_level, team_id=team_id)
        db.session.add(new_user)
        db.session.commit()
        flash("User created successfully.", "success")
        return redirect(url_for('manage_users'))
    return render_template('admin_create_user.html')

@app.route('/admin/users/delete/<int:user_id>', methods=['POST'])
@login_required
def delete_user(user_id):
    if not current_user.is_superuser():
        flash("Access denied: Superuser only.", "danger")
        return redirect(url_for('dashboard'))
    user = User.query.get_or_404(user_id)
    if user.id == current_user.id:
        flash("You cannot delete your own account while logged in.", "warning")
        return redirect(url_for('manage_users'))
    db.session.delete(user)
    db.session.commit()
    flash("User deleted successfully.", "success")
    return redirect(url_for('manage_users'))

@app.route('/admin/backup')
@login_required
def backup_database():
    if not current_user.is_superuser():
        flash('Access denied: Superuser only.', 'danger')
        return redirect(url_for('main.home'))
    db_path = os.path.join(os.getcwd(), 'pip_crm.db')
    if os.path.exists(db_path):
        return send_file(db_path, as_attachment=True)
    else:
        flash('Database file not found.', 'danger')
        return redirect(url_for('admin_dashboard'))

# ----- Employee & PIP basics -----
@app.route('/employee/<int:employee_id>')
@login_required
def employee_detail(employee_id):
    # Pull employee plus related records up-front to avoid lazy-load loops in templates
    employee = (
        Employee.query.options(
            db.joinedload(Employee.pips),
            db.joinedload(Employee.probation_records),
            db.joinedload(Employee.sickness_cases).joinedload(SicknessCase.meetings),
        )
        .get_or_404(employee_id)
    )

    if current_user.admin_level == 0 and employee.team_id != current_user.team_id:
        flash('Access denied')
        return redirect(url_for('dashboard'))

    # ---------- Phase 3: employee-level sickness trigger summary (rolling 12 months) ----------
    today = today_local()
    q_cases = SicknessCase.query.join(Employee).filter(SicknessCase.employee_id == employee.id)

    # Keep consistent with dashboard thresholds
    BRADFORD_MEDIUM = 200
    BRADFORD_HIGH = 400
    EPISODES_THRESHOLD = 3
    TOTAL_DAYS_THRESHOLD = 14
    LONG_TERM_DAYS = 28

    triggers = compute_sickness_trigger_metrics(
        q_cases,
        today=today,
        window_days=365,
        bradford_medium=BRADFORD_MEDIUM,
        bradford_high=BRADFORD_HIGH,
        episodes_threshold=EPISODES_THRESHOLD,
        total_days_threshold=TOTAL_DAYS_THRESHOLD,
        long_term_days=LONG_TERM_DAYS,
    )

    # triggers is a list; for a single employee we either get 0 or 1 entry
    sickness_trigger = triggers[0] if triggers else None

    return render_template(
        'employee_detail.html',
        employee=employee,
        sickness_trigger=sickness_trigger,
    )

@app.route('/pip/<int:id>')
@login_required
def pip_detail(id):
    pip = PIPRecord.query.get_or_404(id)
    employee = pip.employee
    return render_template('pip_detail.html', pip=pip, employee=employee)

@app.route('/pip/edit/<int:id>,', methods=['GET', 'POST'])
@app.route('/pip/edit/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_pip(id):
    pip = PIPRecord.query.get_or_404(id)
    employee = pip.employee

    # Build form
    if request.method == 'POST':
        form = PIPForm()
        form.process(request.form)
    else:
        form = PIPForm(obj=pip)
        # ensure action entries for existing items
        for _ in range(len(pip.action_items) - len(form.actions.entries)):
            form.actions.append_entry()
        for idx, ai in enumerate(pip.action_items):
            form.actions.entries[idx].form.description.data = ai.description
            form.actions.entries[idx].form.status.data = ai.status

    advice_text = None

    # AI advice branch
    if request.method == 'POST' and 'generate_advice' in request.form:
        prompt = (
            f"You are a performance coach.\n"
            f"Employee: {employee.first_name} {employee.last_name}\n"
            f"Job Title: {employee.job_title}\n"
            f"Concerns: {form.concerns.data or '[none]'}\n"
            "Action Items:\n"
        )
        for ai_field in form.actions.entries:
            desc = ai_field.form.description.data or '[no description]'
            stat = ai_field.form.status.data or '[no status]'
            prompt += f"- {desc} [{stat}]\n"
        prompt += f"Meeting Notes: {form.meeting_notes.data or '[none]'}\n"
        prompt += "Provide 3 bulleted actionable tips for the manager to support this employee."

        resp = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        advice_text = resp.choices[0].message.content.strip()
        return render_template('edit_pip.html', form=form, pip=pip, employee=employee, advice_text=advice_text)

    # Save branch
    if form.validate_on_submit():
        pip.concerns = form.concerns.data
        pip.start_date = form.start_date.data
        pip.review_date = form.review_date.data
        pip.status = form.status.data
        pip.meeting_notes = form.meeting_notes.data
        pip.capability_meeting_date = form.capability_meeting_date.data
        pip.capability_meeting_time = form.capability_meeting_time.data
        pip.capability_meeting_venue = form.capability_meeting_venue.data

        pip.action_items.clear()
        for ai_field in form.actions.entries:
            pip.action_items.append(
                PIPActionItem(
                    description=ai_field.form.description.data,
                    status=ai_field.form.status.data
                )
            )

        db.session.commit()
        flash('PIP updated successfully.', 'success')
        return redirect(url_for('pip_detail', id=pip.id))

    return render_template('edit_pip.html', form=form, pip=pip, employee=employee, advice_text=advice_text)

@app.route('/pip/<int:id>/generate/advice', methods=['POST'])
@login_required
def generate_ai_advice(id):
    pip = PIPRecord.query.get_or_404(id)
    employee = pip.employee

    prompt = (
        "You are an experienced HR Line Manager based in the UK. "
        "Using the information below, provide three clear and practical suggestions for how a line manager "
        "can support the employee's performance improvement. Your advice should reflect HR best practice in the UK.\n\n"
        f"Employee Name: {employee.first_name} {employee.last_name}\n"
        f"Job Title: {employee.job_title}\n"
        f"Concerns: {pip.concerns or '[none]'}\n"
        "Action Items:\n"
    )
    for item in pip.action_items:
        prompt += f"- {item.description or '[no description]'} [{item.status or '[no status]'}]\n"
    prompt += f"Meeting Notes: {pip.meeting_notes or '[none]'}\n\n"
    prompt += "Provide your advice as a bullet-pointed list."

    resp = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
    )
    pip.ai_advice = resp.choices[0].message.content.strip()
    pip.ai_advice_generated_at = datetime.now(timezone.utc)

    event = TimelineEvent(
        pip_record_id=pip.id,
        event_type="AI Advice Generated",
        notes="Advice generated using OpenAI",
        updated_by=current_user.username
    )
    db.session.add(event)

    db.session.commit()
    flash('AI advice generated.', 'success')
    return redirect(url_for('pip_detail', id=pip.id))

@app.route('/pip/create/<int:employee_id>', methods=['GET', 'POST'])
@login_required
def create_pip(employee_id):
    employee = Employee.query.get_or_404(employee_id)
    if current_user.admin_level == 0 and employee.team_id != current_user.team_id:
        flash('Access denied.')
        return redirect(url_for('dashboard'))

    form = PIPForm()
    if request.method == 'GET' and len(form.actions.entries) == 0:
        form.actions.append_entry()

    if request.method == 'POST':
        action_fields = [k for k in request.form if 'actions-' in k and '-description' in k]
        form.actions.min_entries = len(set(k.split('-')[1] for k in action_fields))

    if form.validate_on_submit():
        pip = PIPRecord(
            employee_id=employee.id,
            concerns=form.concerns.data,
            start_date=form.start_date.data,
            review_date=form.review_date.data,
            meeting_notes=form.meeting_notes.data
        )
        db.session.add(pip)
        db.session.flush()

        for action_form in form.actions.entries:
            item = PIPActionItem(
                pip_record_id=pip.id,
                description=action_form.form.description.data,
                status=action_form.form.status.data
            )
            db.session.add(item)

        db.session.commit()
        flash('New PIP created.')
        return redirect(url_for('employee_detail', employee_id=employee.id))

    return render_template('create_pip.html', form=form, employee=employee)

@app.route('/employee/edit/<int:employee_id>', methods=['GET', 'POST'])
@login_required
def edit_employee(employee_id):
    employee = Employee.query.get_or_404(employee_id)
    if current_user.admin_level == 0 and employee.team_id != current_user.team_id:
        flash('Access denied.')
        return redirect(url_for('dashboard'))
    form = EmployeeForm(obj=employee)
    if form.validate_on_submit():
        employee.first_name = form.first_name.data
        employee.last_name = form.last_name.data
        employee.job_title = form.job_title.data
        employee.line_manager = form.line_manager.data
        employee.service = form.service.data
        employee.start_date = form.start_date.data
        employee.team_id = form.team_id.data
        employee.email = form.email.data
        db.session.commit()
        flash('Employee details updated.', 'success')
        return redirect(url_for('employee_detail', employee_id=employee.id))
    return render_template('edit_employee.html', form=form, employee=employee)

@app.route('/pip/list')
@login_required
def pip_list():
    pips = PIPRecord.query.join(Employee).all()
    return render_template('pip_list.html', pips=pips)

@app.route('/pip/select-employee', methods=['GET', 'POST'])
@login_required
def select_employee_for_pip():
    employees = Employee.query.order_by(Employee.last_name).all()
    if request.method == 'POST':
        return redirect(url_for('create_pip', employee_id=request.form.get('employee_id')))
    return render_template('pip_select_employee.html', employees=employees)

@app.route('/taxonomy/action_templates', methods=['GET'])
@login_required
def taxonomy_action_templates():
    category = (request.args.get('category') or '').strip()
    severity = (request.args.get('severity') or '').strip()
    items = _pick_actions_from_templates(category, severity)
    return jsonify({"category": category, "severity": severity, "items": items})

# =========================
# app.py — PART 3 / 4
# (PIP wizard, document drafting/editing, AI suggestions)
# =========================

# -------- Wizard entry ----------
class DummyForm(FlaskForm):
    pass

def _max_wizard_step(data: dict) -> int:
    s = 1
    if data.get('employee_id'): s = 2
    if all(data.get(k) for k in ('concerns', 'concern_category', 'severity', 'frequency')): s = 3
    if all(data.get(k) for k in ('start_date', 'review_date')): s = 4
    if all(data.get(k) for k in ('capability_meeting_date', 'capability_meeting_time', 'capability_meeting_venue')): s = 5
    items = data.get('action_plan_items') or []
    if s == 5 and isinstance(items, list) and any((x or '').strip() for x in items):
        s = 6
    return s

@app.route('/pip/create-wizard', methods=['GET', 'POST'])
@login_required
def create_pip_wizard():
    if 'wizard_step' not in session:
        session['wizard_step'] = 1
        session['pip_data'] = {}

    step = session['wizard_step']
    data = session.get('pip_data', {}) or {}
    wizard_errors = {}
    draft = None

    if request.method == 'GET':
        try:
            goto = int(request.args.get('goto', 0))
        except (TypeError, ValueError):
            goto = 0
        if goto:
            max_allowed = _max_wizard_step(data)
            if 1 <= goto <= max_allowed:
                session['wizard_step'] = goto
                step = goto

        if step == 3 and data.get('start_date') and not data.get('review_date'):
            auto_val = _auto_review_date(data.get('start_date'))
            if auto_val:
                data['review_date'] = auto_val
                data['auto_review_populated'] = True
                data['auto_review_date'] = auto_val
                session['pip_data'] = data

    if request.method == 'POST':
        if step == 1:
            employee_id = request.form.get('employee_id')
            draft_name = request.form.get('draft_name', '').strip()
            if not employee_id:
                wizard_errors['employee_id'] = "Please select an employee."
            else:
                data['employee_id'] = int(employee_id)
            data['draft_name'] = draft_name

        elif step == 2:
            concerns = request.form.get('concerns', '').strip()
            category = request.form.get('concern_category', '').strip()
            severity = request.form.get('severity', '').strip()
            frequency = request.form.get('frequency', '').strip()
            tags = request.form.get('concern_tags', '').strip()
            draft_name = request.form.get('draft_name', '').strip()

            if not concerns:
                wizard_errors['concerns'] = "Concerns cannot be empty."
            if not category:
                wizard_errors['concern_category'] = "Please choose a concern category."
            if not severity:
                wizard_errors['severity'] = "Please select severity."
            if not frequency:
                wizard_errors['frequency'] = "Please select frequency."

            data.update({
                'concerns': concerns,
                'concern_category': category,
                'severity': severity,
                'frequency': frequency,
                'concern_tags': tags,
                'draft_name': draft_name,
            })

        elif step == 3:
            start_date = (request.form.get('start_date') or '').strip()
            review_date = (request.form.get('review_date') or '').strip()
            draft_name = request.form.get('draft_name', '').strip()
            review_weeks = (request.form.get('review_weeks') or '').strip()

            if not start_date:
                wizard_errors['start_date'] = "Start date is required."

            auto_flag = False
            if start_date and not review_date:
                auto_val = _auto_review_date(start_date)
                if auto_val:
                    review_date = auto_val
                    auto_flag = True

            if not wizard_errors:
                data['start_date'] = start_date
                data['review_date'] = review_date
                if review_weeks.isdigit():
                    data['review_weeks'] = int(review_weeks)
                elif 'review_weeks' not in data:
                    data['review_weeks'] = 4
                data['auto_review_populated'] = bool(auto_flag)
                data['auto_review_date'] = review_date if auto_flag else None

            data['draft_name'] = draft_name

        elif step == 4:
            data['capability_meeting_date'] = request.form.get('capability_meeting_date')
            data['capability_meeting_time'] = request.form.get('capability_meeting_time')
            data['capability_meeting_venue'] = request.form.get('capability_meeting_venue')
            data['draft_name'] = request.form.get('draft_name', '').strip()

        elif step == 5:
            action_items = request.form.getlist('action_plan_items[]')
            valid_items = [item.strip() for item in action_items if item.strip()]
            if not valid_items:
                wizard_errors['action_plan_items'] = "Add at least one action plan item."
            else:
                data['action_plan_items'] = valid_items
                session['pip_data'] = data
                session['wizard_step'] = 6
                return redirect(url_for('create_pip_wizard'))

        elif step == 6:
            try:
                items = data.get('action_plan_items') or []
                if not any((x or '').strip() for x in items):
                    flash("Please add at least one action plan item.", "warning")
                    session['wizard_step'] = 5
                    return redirect(url_for('create_pip_wizard'))

                pip = PIPRecord(
                    employee_id=int(data['employee_id']),
                    concerns=data['concerns'],
                    concern_category=data.get('concern_category'),
                    severity=data.get('severity'),
                    frequency=data.get('frequency'),
                    tags=data.get('concern_tags'),
                    start_date=datetime.strptime(data['start_date'], '%Y-%m-%d').date(),
                    review_date=datetime.strptime(data['review_date'], '%Y-%m-%d').date(),
                    capability_meeting_date=datetime.strptime(data['capability_meeting_date'], '%Y-%m-%d')
                        if data.get('capability_meeting_date') else None,
                    capability_meeting_time=data.get('capability_meeting_time'),
                    capability_meeting_venue=data.get('capability_meeting_venue'),
                    created_by=current_user.username
                )

                db.session.add(pip)
                db.session.commit()

                for item_text in items:
                    action = PIPActionItem(pip_record_id=pip.id, description=item_text)
                    db.session.add(action)

                timeline = TimelineEvent(
                    pip_record_id=pip.id,
                    event_type="PIP Created",
                    notes="PIP created via multi-step wizard",
                    updated_by=current_user.username
                )
                db.session.add(timeline)

                db.session.commit()

                session.pop('wizard_step', None)
                session.pop('pip_data', None)

                flash("PIP created successfully!", "success")
                return redirect(url_for('dashboard'))

            except Exception as e:
                print(f"[ERROR] Failed to save PIP: {e}")
                flash(f"Error creating PIP: {str(e)}", "danger")

        session['pip_data'] = data
        if not wizard_errors and step < 5:
            session['wizard_step'] = step + 1
            return redirect(url_for('create_pip_wizard'))

    employees = Employee.query.all() if step == 1 else []
    max_allowed = _max_wizard_step(data)

    return render_template(
        'create_pip_wizard.html',
        step=step,
        draft=draft,
        data=data,
        wizard_errors=wizard_errors,
        employees=employees,
        max_allowed_step=max_allowed,
        auto_review_populated=bool(data.get('auto_review_populated')),
        auto_review_date=data.get('auto_review_date')
    )

# ----- AJAX validation for wizard -----
@app.route('/validate_wizard_step', methods=['POST'])
@login_required
def validate_wizard_step():
    step = int(request.form.get("step", 1))
    errors = {}

    if step == 1:
        if not request.form.get("employee_id"):
            errors["employee_id"] = "Please select an employee."

    elif step == 2:
        if not (request.form.get("concerns", "").strip()):
            errors["concerns"] = "Please describe the concern."
        if not (request.form.get("concern_category", "").strip()):
            errors["concern_category"] = "Please select a concern category."
        if not (request.form.get("severity", "").strip()):
            errors["severity"] = "Please select severity."
        if not (request.form.get("frequency", "").strip()):
            errors["frequency"] = "Please select frequency."

    elif step == 3:
        if not (request.form.get("start_date", "").strip()):
            errors["start_date"] = "Please enter a start date."
        # review_date can be auto-populated

    elif step == 4:
        meeting_date = request.form.get("meeting_date", "").strip()
        meeting_time = request.form.get("meeting_time", "").strip()
        if not meeting_date:
            errors["meeting_date"] = "Please enter a meeting date."
        if not meeting_time:
            errors["meeting_time"] = "Please enter a meeting time."

    elif step == 5:
        actions = [a.strip() for a in request.form.getlist("action_plan_items[]") if a.strip()]
        if not actions:
            errors["action_plan_items"] = "Please add at least one action item."

    return jsonify({"success": not errors, "errors": errors})

# ----- Document Drafts -----
def build_doc_rel_dir(pip_id: int, doc_type: str, version: int) -> str:
    return os.path.join(f"PIP_{pip_id}", doc_type, f"v{version}")

def build_placeholder_mapping(pip_rec: PIPRecord) -> dict:
    """Build a placeholder → value mapping for all PIP document templates."""
    emp = getattr(pip_rec, "employee", None)

    def _attr(obj, name: str, default: str = "") -> str:
        """Safe getattr that always returns a string (or default)."""
        if obj is None:
            return default
        value = getattr(obj, name, default)
        return "" if value is None else str(value)

    def _clean_str(value: str) -> str:
        """Normalise whitespace and ensure we never return 'None'."""
        if not value:
            return ""
        # Convert to string and squish whitespace
        text = str(value).replace("None", "").strip()
        return " ".join(text.split())

    # --- Employee basics -----------------------------------------------------
    first = _clean_str(_attr(emp, "first_name"))
    last = _clean_str(_attr(emp, "last_name"))
    full = (first + " " + last).strip()

    job_candidates = [
        _attr(emp, "job_title"),
        _attr(emp, "role"),
        _attr(emp, "position_title"),
    ]
    job = _clean_str(next((j for j in job_candidates if j), ""))

    dept_candidates = [
        _attr(emp, "service"),
        _attr(emp, "department"),
        _attr(emp, "team_name"),
    ]
    dept = _clean_str(next((d for d in dept_candidates if d), ""))

    # --- Manager / HR details ------------------------------------------------
    manager_candidates = [
        getattr(pip_rec, "manager_name", ""),
        getattr(pip_rec, "created_by", ""),
        _attr(emp, "line_manager"),
        getattr(current_user, "full_name", ""),
        getattr(current_user, "username", ""),
    ]
    mgr = _clean_str(next((m for m in manager_candidates if m), ""))

    manager_title = _clean_str(_attr(emp, "line_manager_title"))

    hr_name = (
        getattr(pip_rec, "hr_contact_name", "")
        or getattr(current_user, "full_name", "")
        or getattr(current_user, "username", "")
        or ""
    )
    hr_name = _clean_str(hr_name)

    hr_email = (
        getattr(pip_rec, "hr_contact_email", "")
        or getattr(current_user, "email", "")
        or ""
    )
    hr_email = _clean_str(hr_email)

    policy_ref = getattr(pip_rec, "policy_reference", "") or "Capability and Performance Policy"
    policy_ref = _clean_str(policy_ref)

    # --- Meeting info --------------------------------------------------------
    capability_date = getattr(pip_rec, "capability_meeting_date", None)
    meeting_date_str = capability_date.strftime("%d %B %Y") if capability_date else ""
    meeting_time_str = _clean_str(getattr(pip_rec, "capability_meeting_time", "") or "")
    meeting_dt = _clean_str(f"{meeting_date_str} {meeting_time_str}")
    meeting_location = _clean_str(getattr(pip_rec, "capability_meeting_venue", "") or "")

    # --- PIP dates -----------------------------------------------------------
    start_date = getattr(pip_rec, "start_date", None)
    review_date = getattr(pip_rec, "review_date", None)
    start_str = start_date.strftime("%d %B %Y") if start_date else ""
    review_str = review_date.strftime("%d %B %Y") if review_date else ""

    review_weeks_val = getattr(pip_rec, "review_weeks", "") or ""
    extension_weeks_val = getattr(pip_rec, "extension_weeks", "") or ""

    # --- Concern / ratings / notes ------------------------------------------
    concern_categories = _clean_str(getattr(pip_rec, "concern_category", "") or "")
    concern_tags = _clean_str(getattr(pip_rec, "tags", "") or "")
    severity = _clean_str(getattr(pip_rec, "severity", "") or "")
    frequency = _clean_str(getattr(pip_rec, "frequency", "") or "")

    concerns_summary = _clean_str(getattr(pip_rec, "concerns", "") or "")
    evidence_summary = _clean_str(getattr(pip_rec, "evidence", "") or "")
    support_list = _clean_str(getattr(pip_rec, "support_list", "") or "")
    confidential_notes = _clean_str(getattr(pip_rec, "confidential_notes", "") or "")
    review_dates_list = _clean_str(getattr(pip_rec, "review_dates_str", "") or "")
    reporting_method = _clean_str(getattr(pip_rec, "reporting_method", "") or "")
    adjustments = _clean_str(getattr(pip_rec, "adjustments", "") or "")
    training_plan = _clean_str(getattr(pip_rec, "training_plan", "") or "")
    ongoing_support = _clean_str(getattr(pip_rec, "ongoing_support", "") or "")
    appendix_detail = _clean_str(getattr(pip_rec, "outcome_appendix", "") or "")

    reviewer_name = _clean_str(getattr(pip_rec, "reviewer_name", "") or "")

    # --- Generated meta ------------------------------------------------------
    now = datetime.now()
    today_str = now.strftime("%d %B %Y")

    mapping: dict[str, str] = {
        # Dates / meta
        "[[GENERATED_DATE]]": today_str,
        "[[LETTER_DATE]]":    today_str,
        "[[TODAY_DATE]]":     today_str,
        "[[DOC_VERSION]]":    now.strftime("v%Y.%m.%d"),
        "[[PIP_ID]]":         str(getattr(pip_rec, "id", "") or ""),

        # Employee
        "[[EMPLOYEE_NAME]]":       full,
        "[[EMPLOYEE_FIRST_NAME]]": first,
        "[[EMPLOYEE_LAST_NAME]]":  last,
        "[[JOB_TITLE]]":           job,
        "[[DEPARTMENT]]":          dept,
        "[[SERVICE_NAME]]":        dept,

        # Manager / HR
        "[[MANAGER_NAME]]":    mgr,
        "[[MANAGER_TITLE]]":   manager_title,
        "[[HR_CONTACT_NAME]]": hr_name,
        "[[HR_CONTACT_EMAIL]]": hr_email,
        "[[REVIEWER_NAME]]":   reviewer_name,
        "[[POLICY_REFERENCE]]": policy_ref,

        # Meeting
        "[[MEETING_LOCATION]]": meeting_location,
        "[[MEETING_DATETIME]]": meeting_dt,
        "[[MEETING_DATE]]":     meeting_date_str,
        "[[MEETING_TIME]]":     meeting_time_str,

        # PIP dates
        "[[PIP_START_DATE]]": start_str,
        "[[PIP_END_DATE]]":   review_str,

        # Content / concerns
        "[[CONCERNS_SUMMARY]]":   concerns_summary,
        "[[EVIDENCE_SUMMARY]]":   evidence_summary,
        "[[SUPPORT_LIST]]":       support_list,
        "[[CONFIDENTIAL_NOTES]]": confidential_notes,
        "[[REVIEW_DATES_LIST]]":  review_dates_list,
        "[[REPORTING_METHOD]]":   reporting_method,
        "[[ADJUSTMENTS]]":        adjustments,
        "[[TRAINING_PLAN]]":      training_plan,
        "[[ONGOING_SUPPORT]]":    ongoing_support,
        "[[APPENDIX_DETAIL]]":    appendix_detail,

        # Ratings / categories
        "[[CONCERN_CATEGORIES]]": concern_categories,
        "[[CONCERN_TAGS]]":       concern_tags,
        "[[SEVERITY_RATING]]":    severity,
        "[[FREQUENCY_RATING]]":   frequency,

        # Numbers
        "[[REVIEW_PERIOD_WEEKS]]": str(review_weeks_val) if review_weeks_val != "" else "",
        "[[EXTENSION_WEEKS]]":     str(extension_weeks_val) if extension_weeks_val != "" else "",
    }

    # --- Objective placeholders: initialise OBJ1–OBJ5 to blanks -------------
    for i in range(1, 6):
        idx = str(i)
        mapping.setdefault(f"[[OBJ{idx}_TEXT]]", "")
        mapping.setdefault(f"[[OBJ{idx}_SUCCESS]]", "")
        mapping.setdefault(f"[[OBJ{idx}_ACTIONS_EMPLOYEE]]", "")
        mapping.setdefault(f"[[OBJ{idx}_ACTIONS_MANAGER]]", "")
        mapping.setdefault(f"[[OBJ{idx}_SUPPORT]]", "")
        mapping.setdefault(f"[[OBJ{idx}_DEADLINE]]", "")
        mapping.setdefault(f"[[OBJ{idx}_METRICS]]", "")
        mapping.setdefault(f"[[OBJ{idx}_OUTCOME]]", "")
        mapping.setdefault(f"[[OBJ{idx}_EVIDENCE]]", "")

    # --- Phase 1.5: hydrate objectives from PIPActionItem -------------------
    # We assume a relationship like pip_rec.action_items; if your relationship
    # name or field names differ, tweak this section.
    actions = getattr(pip_rec, "action_items", []) or []

    # Sort by created_at if present, else by id so order is stable
    try:
        actions_sorted = sorted(
            actions,
            key=lambda a: (
                getattr(a, "created_at", None) or getattr(a, "id", 0) or 0
            ),
        )
    except Exception:
        actions_sorted = list(actions)

    for i, action in enumerate(actions_sorted[:5], start=1):
        idx = str(i)

        text = _clean_str(
            getattr(action, "description", "")
            or getattr(action, "summary", "")
            or getattr(action, "objective_text", "")
            or ""
        )

        success = _clean_str(
            getattr(action, "success_criteria", "")
            or getattr(action, "success_measure", "")
            or getattr(action, "expected_outcome", "")
            or ""
        )

        actions_employee = _clean_str(
            getattr(action, "employee_actions", "")
            or getattr(action, "employee_steps", "")
            or getattr(action, "employee_responsibilities", "")
            or ""
        )

        actions_manager = _clean_str(
            getattr(action, "manager_actions", "")
            or getattr(action, "manager_steps", "")
            or getattr(action, "manager_support_actions", "")
            or ""
        )

        support = _clean_str(
            getattr(action, "support", "")
            or getattr(action, "resources", "")
            or getattr(action, "support_required", "")
            or ""
        )

        deadline_date = getattr(action, "due_date", None) or getattr(action, "target_date", None)
        deadline = deadline_date.strftime("%d %B %Y") if deadline_date else ""

        metrics = _clean_str(
            getattr(action, "metrics", "")
            or getattr(action, "measurement", "")
            or getattr(action, "performance_measure", "")
            or ""
        )

        outcome = _clean_str(
            getattr(action, "outcome", "")
            or getattr(action, "status", "")
            or getattr(action, "result", "")
            or ""
        )

        evidence = _clean_str(
            getattr(action, "evidence", "")
            or getattr(action, "evidence_summary", "")
            or ""
        )

        if text:
            mapping[f"[[OBJ{idx}_TEXT]]"] = text
        if success:
            mapping[f"[[OBJ{idx}_SUCCESS]]"] = success
        if actions_employee:
            mapping[f"[[OBJ{idx}_ACTIONS_EMPLOYEE]]"] = actions_employee
        if actions_manager:
            mapping[f"[[OBJ{idx}_ACTIONS_MANAGER]]"] = actions_manager
        if support:
            mapping[f"[[OBJ{idx}_SUPPORT]]"] = support
        if deadline:
            mapping[f"[[OBJ{idx}_DEADLINE]]"] = deadline
        if metrics:
            mapping[f"[[OBJ{idx}_METRICS]]"] = metrics
        if outcome:
            mapping[f"[[OBJ{idx}_OUTCOME]]"] = outcome
        if evidence:
            mapping[f"[[OBJ{idx}_EVIDENCE]]"] = evidence

    return mapping



@app.route("/pip/<int:pip_id>/doc/create/<string:doc_type>", methods=["POST"])
@login_required
def create_pip_doc_draft(pip_id, doc_type):
    pip_rec = PIPRecord.query.get_or_404(pip_id)
    mapping = build_placeholder_mapping(pip_rec)

    template_map = {
        "invite": os.path.join(BASE_DIR, "templates", "docx", "PIP_Invite_Letter_Template_v2025-08-28.docx"),
        "plan":   os.path.join(BASE_DIR, "templates", "docx", "PIP_Action_Plan_Template_v2025-08-28.docx"),
        "outcome":os.path.join(BASE_DIR, "templates", "docx", "PIP_Outcome_Letter_Template_v2025-08-28.docx"),
    }
    if doc_type not in template_map:
        abort(404)

    outcome_choice = getattr(pip_rec, "outcome_choice", None)
    docx_bytes = generate_docx_bytes(template_map[doc_type], mapping, outcome_choice=outcome_choice)

    # DOCX → HTML → sanitize before storing
    html_raw = docx_to_html(docx_bytes)
    html_clean = sanitize_html(html_raw)

    version = _next_version_for(pip_id, doc_type)
    rel_dir = build_doc_rel_dir(pip_id, doc_type, version)
    rel_docx_path = _save_file(docx_bytes, rel_dir, f"{doc_type}_v{version}.docx")

    doc = DocumentFile(
        pip_id=pip_id,
        doc_type=doc_type,
        version=version,
        status="draft",
        docx_path=rel_docx_path,
        html_snapshot=html_clean,
        created_by=current_user.username,
    )

    db.session.add(doc)
    db.session.commit()

    log_timeline_event(
        pip_id=pip_id,
        event_type="Document Draft Created",
        notes=f"{doc.doc_type.capitalize()} v{doc.version} created from template.",
    )

    flash(f"{doc_type.capitalize()} draft v{version} created.", "success")
    return redirect(url_for("edit_pip_doc", pip_id=pip_id, doc_id=doc.id))

@app.route("/pip/<int:pip_id>/doc/<int:doc_id>/edit", methods=["GET", "POST"])
@login_required
def edit_pip_doc(pip_id, doc_id):
    pip_rec = PIPRecord.query.get_or_404(pip_id)
    doc = DocumentFile.query.filter_by(id=doc_id, pip_id=pip_id).first_or_404()

    if request.method == "POST":
        html = request.form.get("html", "")
        if not html:
            flash("No content received.", "warning")
            return redirect(request.url)

        # Sanitize the HTML before we store it or convert it back to DOCX
        clean_html = sanitize_html(html)

        # Build DOCX from the cleaned HTML so the round-trip stays tidy/safe
        new_docx = html_to_docx_bytes(clean_html)
        rel_dir = build_doc_rel_dir(pip_id, doc.doc_type, doc.version)
        rel_docx_path = _save_file(
            new_docx,
            rel_dir,
            f"{doc.doc_type}_v{doc.version}_edited.docx"
        )

        # Store the cleaned HTML as our canonical snapshot
        doc.html_snapshot = clean_html
        doc.docx_path = rel_docx_path
        db.session.commit()

        log_timeline_event(
            pip_id=pip_id,
            event_type="Document Draft Updated",
            notes=f"{doc.doc_type.capitalize()} v{doc.version} draft updated.",
        )

        flash("Draft updated.", "success")
        return redirect(request.url)

    # GET: initial load
    return render_template(
        "doc_editor.html",
        pip_rec=pip_rec,
        doc=doc,
        html_content=doc.html_snapshot,
    )


@app.route("/pip/<int:pip_id>/doc/<int:doc_id>/finalise", methods=["POST"])
@login_required
def finalise_pip_doc(pip_id, doc_id):
    doc = DocumentFile.query.filter_by(id=doc_id, pip_id=pip_id).first_or_404()
    if doc.status == "final":
        flash("Document already final.", "info")
        return redirect(url_for("pip_documents", pip_id=pip_id))
    doc.status = "final"
    db.session.commit()
    log_timeline_event(
        pip_id=pip_id,
        event_type="Document Finalised",
        notes=f"{doc.doc_type.capitalize()} v{doc.version} finalised.",
    )
    flash(f"{doc.doc_type.capitalize()} v{doc.version} finalised.", "success")
    return redirect(url_for("pip_documents", pip_id=pip_id))

@app.route("/pip/<int:pip_id>/documents", methods=["GET"])
@login_required
def pip_documents(pip_id):
    pip_rec = PIPRecord.query.get_or_404(pip_id)
    docs = (DocumentFile.query
            .filter_by(pip_id=pip_id)
            .order_by(DocumentFile.created_at.desc())
            .all())
    return render_template("pip_documents.html", pip_rec=pip_rec, docs=docs)

@app.route("/download/doc/<int:doc_id>")
@login_required
def download_doc(doc_id):
    doc = DocumentFile.query.get_or_404(doc_id)
    abs_dir = app.config['UPLOAD_FOLDER']
    dirname, filename = os.path.split(doc.docx_path)
    return send_from_directory(os.path.join(abs_dir, dirname), filename, as_attachment=True)

# ----- AI Action Suggestions -----
@app.route('/suggest_actions_ai', methods=['POST'])
@login_required
@csrf.exempt
def suggest_actions_ai():
    data = request.get_json(silent=True) or {}
    concerns  = (data.get('concerns')  or '').strip()
    severity  = (data.get('severity')  or '').strip()
    frequency = (data.get('frequency') or '').strip()
    tags      = (data.get('tags')      or '').strip()
    category  = (data.get('category')  or '').strip()

    try:
        prior_actions = _pick_actions_from_templates(category, severity)
    except Exception:
        prior_actions = []

    def _dedupe_clean(items, cap=None):
        out, seen = [], set()
        for x in (items or []):
            s = (x or "").strip()
            if not s:
                continue
            k = s.lower()
            if k not in seen:
                out.append(s)
                seen.add(k)
            if cap and len(out) >= cap:
                break
        return out

    sys_msg = (
        "You are an HR advisor in the UK.\n"
        "Return ONLY valid JSON with two arrays:\n"
        '{"actions": ["short concrete manager actions"], "next_up": ["quick follow-ups or escalations"]}.\n'
        "Actions must be specific, measurable where possible, supportive, and suitable for a PIP context.\n"
        "No prose, no markdown, JSON only."
    )

    prior_block = ""
    if prior_actions:
        import json as _json
        prior_block = "Seed actions (consider and adapt as appropriate): " + _json.dumps(
            prior_actions, ensure_ascii=False
        )

    user_msg = f"""
Concern Category: {category or "[unspecified]"}
Concerns: {concerns or "[none]"}
Tags: {tags or "[none]"}
Severity: {severity or "[unspecified]"}
Frequency: {frequency or "[unspecified]"}

{prior_block}

Rules:
- Provide 3–5 'actions' tailored to the inputs.
- Provide 2–4 'next_up' items (e.g., monitoring cadence, policy references, escalation steps).
- Keep each item under 140 characters.
- JSON ONLY.
"""

    actions_llm, next_up_llm, raw = [], [], ""

    def _extract_text_from_choice(choice):
        """
        Normalise OpenAI response content across SDK versions:
        - Older: choice.message.content is a string
        - Newer: choice.message.content may be a list of parts
        """
        content = getattr(choice.message, "content", "")
        if isinstance(content, str):
            return content
        if isinstance(content, list):
            parts = []
            for part in content:
                # content part could be a dict with nested text
                if isinstance(part, dict):
                    text = part.get("text")
                    if isinstance(text, dict):
                        parts.append(str(text.get("value", "")))
                    elif text is not None:
                        parts.append(str(text))
                else:
                    parts.append(str(part))
            return "".join(parts)
        return str(content or "")

    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",  # or "gpt-4" if you prefer
            messages=[
                {"role": "system", "content": sys_msg},
                {"role": "user", "content": user_msg},
            ],
            temperature=0.5,
            max_tokens=300,
        )

        raw = _extract_text_from_choice(resp.choices[0]).strip()

        import json as _json, re as _re
        m = _re.search(r"\{[\s\S]*\}", raw)
        json_str = m.group(0) if m else raw
        payload = _json.loads(json_str)

        actions_llm = payload.get("actions", []) or []
        next_up_llm = payload.get("next_up", []) or []
    except Exception:
        # Fallback: try to salvage bullet-style lines out of whatever we got back
        lines = [ln.strip("-•* 0123456789.\t") for ln in (raw.splitlines() if raw else [])]
        actions_llm = [ln for ln in lines if ln][:5]
        next_up_llm = []

    # Merge curated + LLM suggestions
    merged_actions = _dedupe_clean(actions_llm, cap=None)
    if prior_actions:
        merged_actions = _dedupe_clean(prior_actions + merged_actions, cap=8)

    next_up = _dedupe_clean(next_up_llm, cap=None)

    # Heuristic enrichments
    tag_list = [t.strip().lower() for t in tags.split(",")] if tags else []
    cat = (category or "").lower()
    sev = (severity or "").lower()
    freq = (frequency or "").lower()

    enrich = []
    if 'lateness' in tag_list or 'timekeeping' in cat:
        enrich += [
            "Daily start-time check-ins for 2 weeks",
            "Agree punctuality targets; log variances",
        ]
    if 'conduct' in tag_list or cat == 'conduct':
        enrich += [
            "Reference conduct policy; document conversations",
            "Book values/behaviour refresher",
        ]
    if 'performance' in cat or ('missed deadlines' in (tags or '').lower()):
        enrich += [
            "Weekly milestones with due dates",
            "Stand-up updates Mon/Wed/Fri",
        ]
    if sev == 'high':
        enrich += ["Escalate to formal stage if no progress"]
    if freq in ('frequent', 'persistent'):
        enrich += ["Increase monitoring and assign a buddy/mentor"]

    next_up = _dedupe_clean(next_up + enrich, cap=8)

    merged_actions = merged_actions[:8] if merged_actions else []
    next_up = next_up[:8] if next_up else []

    return jsonify({"success": True, "actions": merged_actions, "next_up": next_up}), 200

# =========================
# app.py — PART 4 / 4
# (Probation module, drafts, dashboards, employee import, doc generation, ping/main)
# =========================

# ------------------------------
# Probation Wizard (single, de-duplicated versions)
# ------------------------------
@app.route("/probation/create-wizard", methods=["GET"])
@login_required
def probation_create_wizard():
    # Load existing probation draft
    draft = DraftProbation.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    step = draft.step if draft else 1
    data = draft.payload if draft else {}
    return render_template("probation_create_wizard.html", step=step, data=data, draft=draft)

@app.route("/probation/save-draft", methods=["POST"])
@login_required
def probation_save_draft():
    payload = request.json or {}
    step = payload.get("step", 1)

    draft = DraftProbation.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    if not draft:
        draft = DraftProbation(user_id=current_user.id, step=step, payload=payload)
        db.session.add(draft)
    else:
        draft.step = step
        draft.payload = payload
    db.session.commit()
    return jsonify({"success": True, "updated_at": draft.updated_at.strftime("%Y-%m-%d %H:%M:%S")})

@app.route("/probation/resume-draft")
@login_required
def probation_resume_draft():
    draft = DraftProbation.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    if draft:
        return redirect(url_for("probation_create_wizard"))
    flash("No probation draft available.", "info")
    return redirect(url_for("probation_dashboard"))

@app.route('/probation/<int:id>')
@login_required
def view_probation(id):
    probation = ProbationRecord.query.get_or_404(id)
    employee = probation.employee
    return render_template('view_probation.html', probation=probation, employee=employee)

@app.route('/probation/<int:id>/review/add', methods=['GET', 'POST'])
@login_required
def add_probation_review(id):
    probation = ProbationRecord.query.get_or_404(id)
    form = ProbationReviewForm()
    if form.validate_on_submit():
        review = ProbationReview(
            probation_id=probation.id,
            review_date=form.review_date.data,
            reviewer=form.reviewer.data,
            summary=form.summary.data,
            concerns_flag=(form.concerns_flag.data.lower() == 'yes')
        )
        db.session.add(review)
        event = TimelineEvent(
            pip_record_id=None,
            event_type="Probation Review",
            notes=f"Review added by {current_user.username}",
            updated_by=current_user.username
        )
        db.session.add(event)
        db.session.commit()
        flash('Probation review added.', 'success')
        return redirect(url_for('view_probation', id=probation.id))
    return render_template('add_probation_review.html', form=form, probation=probation)

@app.route('/probation/<int:id>/plan/add', methods=['GET', 'POST'])
@login_required
def add_probation_plan(id):
    probation = ProbationRecord.query.get_or_404(id)
    form = ProbationPlanForm()
    if form.validate_on_submit():
        plan = ProbationPlan(
            probation_id=probation.id,
            objectives=form.objectives.data,
            deadline=form.deadline.data,
            outcome=form.outcome.data
        )
        db.session.add(plan)
        event = TimelineEvent(
            pip_record_id=None,
            event_type="Probation Plan Added",
            notes=f"Plan created by {current_user.username}",
            updated_by=current_user.username
        )
        db.session.add(event)
        db.session.commit()
        flash('Development plan added.', 'success')
        return redirect(url_for('view_probation', id=probation.id))
    return render_template('add_probation_plan.html', form=form, probation=probation)

@app.route('/probation/<int:id>/edit', methods=['GET', 'POST'])
@login_required
def edit_probation(id):
    probation = ProbationRecord.query.get_or_404(id)
    form = ProbationRecordForm(obj=probation)
    if form.validate_on_submit():
        probation.start_date = form.start_date.data
        probation.expected_end_date = form.expected_end_date.data
        probation.notes = form.notes.data
        db.session.commit()
        flash('Probation record updated.', 'success')
        return redirect(url_for('view_probation', id=probation.id))
    return render_template('edit_probation.html', form=form, probation=probation)

@app.route('/probation/<int:id>/status/<new_status>', methods=['POST'])
@login_required
def update_probation_status(id, new_status):
    probation = ProbationRecord.query.get_or_404(id)
    valid_statuses = ['Completed', 'Extended', 'Failed']
    if new_status not in valid_statuses:
        flash('Invalid status update.', 'danger')
        return redirect(url_for('view_probation', id=id))
    probation.status = new_status
    db.session.add(probation)
    event = TimelineEvent(
        pip_record_id=None,
        event_type="Probation Status Updated",
        notes=f"Status changed to {new_status} by {current_user.username}",
        updated_by=current_user.username
    )
    db.session.add(event)
    db.session.commit()
    flash(f'Status updated to {new_status}.', 'success')
    return redirect(url_for('view_probation', id=id))

@app.route('/probation/create/<int:employee_id>', methods=['GET', 'POST'])
@login_required
def create_probation(employee_id):
    employee = Employee.query.get_or_404(employee_id)
    form = ProbationRecordForm()
    if form.validate_on_submit():
        probation = ProbationRecord(
            employee_id=employee.id,
            start_date=form.start_date.data,
            expected_end_date=form.expected_end_date.data,
            notes=form.notes.data
        )
        db.session.add(probation)
        db.session.commit()
        flash('Probation record created successfully.', 'success')
        return redirect(url_for('view_probation', id=probation.id))
    return render_template('create_probation.html', form=form, employee=employee)

@app.route('/probation/dashboard')
@login_required
def probation_dashboard():
    global_active = ProbationRecord.query.filter_by(status='Active').count()
    global_completed = ProbationRecord.query.filter_by(status='Completed').count()
    global_extended = ProbationRecord.query.filter_by(status='Extended').count()

    today = today_local()
    soon = today + timedelta(days=14)

    q_records = ProbationRecord.query.join(Employee)
    q_reviews = ProbationReview.query.join(ProbationRecord, ProbationReview.probation_id == ProbationRecord.id).join(Employee)

    if current_user.admin_level == 0:
        q_records = q_records.filter(Employee.team_id == current_user.team_id)
        q_reviews = q_reviews.filter(Employee.team_id == current_user.team_id)

    active_probations = (
        q_records.filter(ProbationRecord.status == 'Active')
        .order_by(ProbationRecord.expected_end_date.asc().nullslast())
        .all()
    )

    upcoming_reviews = (
        q_reviews.filter(
            ProbationRecord.status == 'Active',
            ProbationReview.review_date >= today,
            ProbationReview.review_date <= soon
        )
        .order_by(ProbationReview.review_date.asc())
        .all()
    )

    overdue_reviews = (
        q_reviews.filter(
            ProbationRecord.status == 'Active',
            ProbationReview.review_date < today
        ).count()
    )

    due_soon_count = (
        q_records.filter(
            ProbationRecord.status == 'Active',
            ProbationRecord.expected_end_date >= today,
            ProbationRecord.expected_end_date <= soon
        ).count()
    )

    probation_draft = get_active_probation_draft_for_user(current_user.id)

    return render_template(
        'probation_dashboard.html',
        active_module='Probation',
        global_active=global_active,
        global_completed=global_completed,
        global_extended=global_extended,
        active_probations=active_probations,
        upcoming_reviews=upcoming_reviews,
        overdue_reviews=overdue_reviews,
        due_soon_count=due_soon_count,
        draft=probation_draft
    )

@app.route('/probation/dismiss-draft', methods=['POST'])
@login_required
def dismiss_probation_draft():
    draft = DraftProbation.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    if not draft:
        return jsonify({'success': False, 'message': 'No active draft'}), 400
    draft.is_dismissed = True
    db.session.commit()
    return jsonify({'success': True})

@app.route('/probation/employees')
@login_required
def probation_employee_list():
    session['active_module'] = 'Probation'
    employees = Employee.query.all()
    return render_template('probation_employee_list.html', employees=employees)

# ------------------------------
# Sickness Management Module
# ------------------------------

@app.route("/sickness/dashboard")
@login_required
def sickness_dashboard():
    """High-level view of sickness cases and RTW meetings, plus trigger flags."""
    today = today_local()
    long_term_threshold = today - timedelta(days=28)
    one_year_ago = today - timedelta(days=365)
    upcoming_from = today
    upcoming_to = today + timedelta(days=14)

    # Optional Phase 3 filters (no template changes required; safe defaults)
    # Examples:
    #   /sickness/dashboard?severity=high
    #   /sickness/dashboard?service=Riverside
    severity_filter = (request.args.get("severity") or "").strip().lower()
    service_filter = (request.args.get("service") or "").strip()

    q_cases = SicknessCase.query.join(Employee)
    q_meetings = SicknessMeeting.query.join(SicknessCase).join(Employee)

    # Line managers see only their team
    if current_user.admin_level == 0:
        q_cases = q_cases.filter(Employee.team_id == current_user.team_id)
        q_meetings = q_meetings.filter(Employee.team_id == current_user.team_id)

    # Optional service filter (admin/superuser use-cases, but safe for everyone)
    if service_filter:
        q_cases = q_cases.filter(Employee.service == service_filter)
        q_meetings = q_meetings.filter(Employee.service == service_filter)

    # Open cases (strictly status == "Open")
    open_cases = (
        q_cases.filter(SicknessCase.status == "Open")
        .order_by(SicknessCase.start_date.desc())
        .all()
    )
    open_count = len(open_cases)

    # Closed in last 12 months (with an end_date)
    closed_last_12m = (
        q_cases.filter(
            SicknessCase.status == "Closed",
            SicknessCase.end_date.isnot(None),
            SicknessCase.end_date >= one_year_ago,
        ).count()
    )

    # Upcoming RTW / review meetings in next 14 days (inclusive)
    upcoming_meetings = (
        q_meetings.filter(
            SicknessMeeting.meeting_date.isnot(None),
            SicknessMeeting.meeting_date >= upcoming_from,
            SicknessMeeting.meeting_date <= upcoming_to,
        )
        .order_by(SicknessMeeting.meeting_date.asc())
        .all()
    )

    # Long-term open cases (> 28 days open)
    long_term_cases = (
        q_cases.filter(
            SicknessCase.status == "Open",
            SicknessCase.start_date.isnot(None),
            SicknessCase.start_date <= long_term_threshold,
        ).count()
    )

    # ---------- Phase 3: Trigger metrics (last 12 months) ----------
    # Overlap-aware and reusable for employee view later.
    potential_triggers = compute_sickness_trigger_metrics(
        q_cases,
        today=today,
        window_days=365,
        bradford_medium=200,
        bradford_high=400,
        episodes_threshold=3,
        total_days_threshold=14,
        long_term_days=28,
    )

    # Optional: filter triggers by severity via querystring
    if severity_filter in {"high", "medium", "low"}:
        potential_triggers = [
            t for t in (potential_triggers or [])
            if (t.get("severity") or "").lower() == severity_filter
        ]

    # Phase 3: add recommended actions (does not change the template unless you choose to display it)
    # Kept as simple heuristics for now: predictable, explainable, HR-friendly.
    for t in potential_triggers or []:
        sev = (t.get("severity") or "").lower()
        flags_label = (t.get("flags_label") or "").lower()

        actions = []

        # Common hygiene actions
        actions.append("Check last RTW note and update case notes")

        if "long-term" in flags_label:
            actions.append("Consider Occupational Health referral / welfare review")
            actions.append("Confirm fit note status and expected return date")

        if sev == "high":
            actions.insert(0, "Escalate: formal absence review meeting (HR involved)")
            actions.append("Agree formal plan: review dates, adjustments, expectations")
        elif sev == "medium":
            actions.insert(0, "Schedule absence review meeting within 7 days")
            actions.append("Discuss patterns, support, adjustments and next steps")
        elif sev == "low":
            actions.insert(0, "Monitor and prompt manager to complete RTW process")
            actions.append("Watch for repeat episodes over next 4–8 weeks")

        # Add to trigger row for later UI use
        t["recommended_actions"] = actions

    return render_template(
        "sickness_dashboard.html",
        active_module="Sickness",
        open_cases=open_cases,
        open_count=open_count,
        closed_last_12m=closed_last_12m,
        upcoming_meetings=upcoming_meetings,
        long_term_cases=long_term_cases,
        potential_triggers=potential_triggers,
        # Optional extras for future UI controls (safe to ignore in template)
        severity_filter=severity_filter,
        service_filter=service_filter,
    )

@app.route("/sickness/list")
@login_required
def sickness_list():
    """List sickness cases with simple filters."""
    today = today_local()

    status = request.args.get("status", "open")  # open | closed | all
    service = request.args.get("service") or None
    trigger = request.args.get("trigger") or None

    q = SicknessCase.query.join(Employee)

    # Line managers see only their team
    if current_user.admin_level == 0:
        q = q.filter(Employee.team_id == current_user.team_id)

    # Status filter
    if status == "open":
        q = q.filter(SicknessCase.status == "Open")
    elif status == "closed":
        q = q.filter(SicknessCase.status == "Closed")
    # "all" = no extra filter

    # Service filter
    if service:
        q = q.filter(Employee.service == service)

    # Trigger type filter
    if trigger:
        q = q.filter(SicknessCase.trigger_type == trigger)

    cases = q.order_by(SicknessCase.start_date.desc()).all()

    # Build list of services that actually have sickness cases
    services_query = (
        db.session.query(Employee.service)
        .join(SicknessCase, SicknessCase.employee_id == Employee.id)
        .filter(Employee.service.isnot(None))
    )
    if current_user.admin_level == 0:
        services_query = services_query.filter(Employee.team_id == current_user.team_id)

    services = sorted({row[0] for row in services_query})

    trigger_types = [
        ("short_term", "Short-term trigger"),
        ("long_term", "Long-term"),
        ("pattern", "Pattern concern"),
        ("other", "Other"),
    ]

    return render_template(
        "sickness_list.html",
        active_module="Sickness",
        cases=cases,
        status=status,
        service=service,
        trigger=trigger,
        services=services,
        trigger_types=trigger_types,
    )


from flask_wtf import FlaskForm
from wtforms import DateField, TextAreaField, SelectField, SubmitField
from wtforms.validators import DataRequired

class SicknessCaseForm(FlaskForm):
    start_date = DateField("First day of absence", validators=[DataRequired()], format="%Y-%m-%d")
    status = SelectField(
        "Status",
        choices=[("Open", "Open"), ("Closed", "Closed")],
        default="Open",
    )
    notes = TextAreaField("Reason / notes")
    submit = SubmitField("Save sickness record")

@app.route("/employee/<int:employee_id>/sickness/new", methods=["GET", "POST"])
@login_required
def sickness_create_for_employee(employee_id):
    session["active_module"] = "Sickness"

    employee = Employee.query.get_or_404(employee_id)
    form = SicknessCaseForm()

    if request.method == "GET" and not form.start_date.data:
        form.start_date.data = datetime.now(timezone.utc).date()

    if form.validate_on_submit():
        case = SicknessCase(
            employee_id=employee.id,
            start_date=form.start_date.data,
            status=form.status.data,
            # If your model has a notes/reason field, wire it here:
            # notes=form.notes.data,
        )
        db.session.add(case)
        db.session.commit()
        flash("Sickness case created.", "success")
        return redirect(url_for("sickness_dashboard"))
    
        # --- Sickness summary for this employee ---
    today = today_local()
    one_year_ago = today - timedelta(days=365)

    sickness_q = SicknessCase.query.filter_by(employee_id=employee.id)

    cases_last_12m = sickness_q.filter(SicknessCase.start_date >= one_year_ago).all()
    sickness_episodes_12m = len(cases_last_12m)

    sickness_days_lost_12m = 0
    for c in cases_last_12m:
        if c.start_date:
            end = c.end_date or today
            sickness_days_lost_12m += (end - c.start_date).days + 1

    open_sickness_case = (
        sickness_q.filter(SicknessCase.status == "Open")
        .order_by(SicknessCase.start_date.desc())
        .first()
    )

    recent_sickness_cases = (
        sickness_q.order_by(SicknessCase.start_date.desc())
        .limit(3)
        .all()
    )


    return render_template(
        "sickness_case_form.html",
        form=form,
        employee=employee,
        active_module="Sickness",
        sickness_episodes_12m=sickness_episodes_12m,
        sickness_days_lost_12m=sickness_days_lost_12m,
        open_sickness_case=open_sickness_case,
        recent_sickness_cases=recent_sickness_cases,

    )



@app.route("/sickness/create/<int:employee_id>", methods=["GET", "POST"])
@login_required
def create_sickness_case(employee_id):
    """Create a sickness case for a specific employee."""
    employee = Employee.query.get_or_404(employee_id)
    form = SicknessCaseForm()

    if form.validate_on_submit():
        case = SicknessCase(
            employee_id=employee.id,
            start_date=form.start_date.data,
            end_date=form.end_date.data,
            reason=form.reason.data.strip() if form.reason.data else None,
            trigger_type=form.trigger_type.data or None,
            notes=form.notes.data,
            status="Open",
            created_by=current_user.username,
        )
        db.session.add(case)

        # Log a timeline event (generic, not linked via FK yet)
        event = TimelineEvent(
            pip_record_id=None,
            event_type="Sickness Case Created",
            notes=f"Sickness case {case.id} created for {employee.first_name} {employee.last_name} by {current_user.username}",
            updated_by=current_user.username,
        )
        db.session.add(event)

        db.session.commit()
        flash("Sickness case created.", "success")
        return redirect(url_for("view_sickness_case", case_id=case.id))

    # Default start date: today
    if request.method == "GET" and not form.start_date.data:
        form.start_date.data = datetime.utcnow().date()

    return render_template(
        "create_sickness_case.html",
        form=form,
        employee=employee,
    )


@app.route("/sickness/<int:case_id>")
@login_required
def view_sickness_case(case_id):
    """View a sickness case and its meetings."""
    case = SicknessCase.query.get_or_404(case_id)
    employee = case.employee
    meetings = (
        SicknessMeeting.query.filter_by(sickness_case_id=case.id)
        .order_by(SicknessMeeting.meeting_date.asc())
        .all()
    )

    today = today_local()  # UK-local "today" (date)

    return render_template(
        "view_sickness_case.html",
        case=case,
        employee=employee,
        meetings=meetings,
        today=today,
    )


@app.route("/sickness/<int:case_id>/meeting/add", methods=["GET", "POST"])
@login_required
def add_sickness_meeting(case_id):
    """Add a RTW / absence review / welfare meeting to a sickness case."""
    case = SicknessCase.query.get_or_404(case_id)
    employee = case.employee
    form = SicknessMeetingForm()

    if form.validate_on_submit():
        meeting = SicknessMeeting(
            sickness_case_id=case.id,
            meeting_date=form.meeting_date.data,
            meeting_type=form.meeting_type.data,
            chair=form.chair.data.strip() if form.chair.data else None,
            notes=form.notes.data,
            outcome=form.outcome.data,
        )
        db.session.add(meeting)

        event = TimelineEvent(
            pip_record_id=None,
            event_type="Sickness Meeting Logged",
            notes=f"{form.meeting_type.data} for sickness case {case.id} ({employee.first_name} {employee.last_name}) logged by {current_user.username}",
            updated_by=current_user.username,
        )
        db.session.add(event)

        db.session.commit()
        flash("Sickness meeting saved.", "success")
        return redirect(url_for("view_sickness_case", case_id=case.id))

    # Default meeting date: today
    if request.method == "GET" and not form.meeting_date.data:
        form.meeting_date.data = datetime.utcnow().date()

    return render_template(
        "add_sickness_meeting.html",
        form=form,
        case=case,
        employee=employee,
    )


@app.route("/sickness/<int:case_id>/status/<new_status>", methods=["POST"])
@login_required
def update_sickness_status(case_id, new_status):
    """Update sickness case status (Open / Closed / Under Review)."""
    case = SicknessCase.query.get_or_404(case_id)
    valid_statuses = ["Open", "Closed", "Under Review"]
    if new_status not in valid_statuses:
        flash("Invalid status update.", "danger")
        return redirect(url_for("view_sickness_case", case_id=case.id))

    case.status = new_status
    db.session.add(case)

    event = TimelineEvent(
        pip_record_id=None,
        event_type="Sickness Status Updated",
        notes=f"Sickness case {case.id} status changed to {new_status} by {current_user.username}",
        updated_by=current_user.username,
    )
    db.session.add(event)

    db.session.commit()
    flash(f"Status updated to {new_status}.", "success")
    return redirect(url_for("view_sickness_case", case_id=case.id))







# ----- PIP draft helpers -----
def get_active_draft_for_user(user_id):
    return DraftPIP.query.filter_by(user_id=user_id, is_dismissed=False)\
                         .order_by(DraftPIP.updated_at.desc()).first()

@app.route('/dismiss_draft', methods=['POST'])
@login_required
@csrf.exempt
def dismiss_draft():
    draft = DraftPIP.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
    if draft:
        draft.is_dismissed = True
        draft.updated_at = datetime.now(timezone.utc)
        db.session.commit()
        return jsonify({"success": True})
    return jsonify({"success": False, "message": "No active draft found"}), 404

@app.route('/save_pip_draft', methods=['POST'])
@login_required
@csrf.exempt
def save_pip_draft():
    try:
        data = {
            'employee_id': request.form.get('employee_id'),
            'draft_name': request.form.get('draft_name', '').strip(),
            'concerns': request.form.get('concerns', '').strip(),
            'concern_category': request.form.get('concern_category', '').strip(),
            'severity': request.form.get('severity', '').strip(),
            'frequency': request.form.get('frequency', '').strip(),
            'concern_tags': request.form.get('concern_tags', '').strip(),
            'start_date': request.form.get('start_date'),
            'review_date': request.form.get('review_date'),
            'capability_meeting_date': request.form.get('capability_meeting_date'),
            'capability_meeting_time': request.form.get('capability_meeting_time'),
            'capability_meeting_venue': request.form.get('capability_meeting_venue'),
            'action_plan_items': request.form.getlist('action_plan_items[]')
        }
        cleaned_data = {k: v for k, v in data.items() if v not in [None, '', []]}

        existing_draft = DraftPIP.query.filter_by(user_id=current_user.id, is_dismissed=False).first()
        if existing_draft:
            db.session.delete(existing_draft)
            db.session.commit()

        new_draft = DraftPIP(
            user_id=current_user.id,
            data=cleaned_data,
            step=session.get('wizard_step', 1),
            is_dismissed=False,
            updated_at=datetime.now(timezone.utc)
        )
        db.session.add(new_draft)
        db.session.commit()

        return jsonify({"success": True, "message": "Draft saved."})
    except Exception as e:
        print(f"[ERROR] Failed to save draft: {e}")
        return jsonify({"success": False, "message": "Failed to save draft."}), 500

@app.route('/pip/wizard/resume', methods=['GET'])
@login_required
def pip_wizard_resume():
    draft = get_active_draft_for_user(current_user.id)
    if not draft or not draft.data:
        flash("No active draft to resume.", "warning")
        return redirect(url_for('dashboard'))
    session['pip_data'] = dict(draft.data)
    session['wizard_step'] = _max_wizard_step(session['pip_data'])
    return redirect(url_for('create_pip_wizard'))

# ----- Dashboard -----
@app.route('/dashboard')
@login_required
def dashboard():
    total_employees = Employee.query.count()
    active_pips = PIPRecord.query.filter_by(status='Open').count()
    completed_pips = PIPRecord.query.filter_by(status='Completed').count()

    today = today_local()
    upcoming_deadline = today + timedelta(days=7)

    if current_user.admin_level == 0:
        q_base = (
            PIPRecord.query
            .join(Employee)
            .filter(Employee.team_id == current_user.team_id)
            .filter(PIPRecord.assigned_to == current_user.id)
        )

        overdue_reviews = (
            q_base.filter(
                PIPRecord.status == 'Open',
                PIPRecord.review_date < today
            ).count()
        )

        open_pips = (
            q_base.filter(PIPRecord.status == 'Open')
            .order_by(PIPRecord.review_date.asc().nullslast())
            .all()
        )

        q_upcoming = q_base.filter(
            PIPRecord.status == 'Open',
            PIPRecord.review_date >= today,
            PIPRecord.review_date <= upcoming_deadline
        )
        upcoming_pips = q_upcoming.order_by(PIPRecord.review_date.asc()).all()
        due_soon_count = q_upcoming.count()

        recent_activity = (
            TimelineEvent.query
            .join(PIPRecord, TimelineEvent.pip_record_id == PIPRecord.id)
            .filter(PIPRecord.assigned_to == current_user.id)
            .order_by(TimelineEvent.timestamp.desc())
            .limit(10)
            .all()
        )
    else:
        overdue_reviews = (
            PIPRecord.query
            .filter(
                PIPRecord.status == 'Open',
                PIPRecord.review_date < today
            ).count()
        )

        open_pips = (
            PIPRecord.query
            .filter(PIPRecord.status == 'Open')
            .join(Employee)
            .order_by(PIPRecord.review_date.asc().nullslast())
            .all()
        )

        upcoming_pips = (
            PIPRecord.query
            .filter(
                PIPRecord.status == 'Open',
                PIPRecord.review_date >= today,
                PIPRecord.review_date <= upcoming_deadline
            )
            .order_by(PIPRecord.review_date.asc())
            .all()
        )

        due_soon_count = (
            PIPRecord.query
            .filter(
                PIPRecord.status == 'Open',
                PIPRecord.review_date >= today,
                PIPRecord.review_date <= upcoming_deadline
            ).count()
        )

        recent_activity = (
            TimelineEvent.query
            .order_by(TimelineEvent.timestamp.desc())
            .limit(10)
            .all()
        )

    draft = get_active_draft_for_user(current_user.id)

    return render_template(
        'dashboard.html',
        total_employees=total_employees,
        active_pips=active_pips,
        completed_pips=completed_pips,
        overdue_reviews=overdue_reviews,
        recent_activity=recent_activity,
        upcoming_pips=upcoming_pips,
        open_pips=open_pips,
        due_soon_count=due_soon_count,
        draft=draft
    )

@app.route('/dashboard/stats.json')
@login_required
def dashboard_stats_json():
    today = today_local()
    upcoming_deadline = today + timedelta(days=7)

    base = _open_pips_scoped_query()
    by_category = _counts_by_field(PIPRecord.concern_category)
    by_severity = _counts_by_field(PIPRecord.severity)

    open_total = base.count()
    due_soon = base.filter(
        PIPRecord.review_date >= today,
        PIPRecord.review_date <= upcoming_deadline
    ).count()
    overdue = base.filter(PIPRecord.review_date < today).count()

    return jsonify({
        "by_category": by_category,
        "by_severity": by_severity,
        "totals": {"open": open_total, "due_soon": due_soon, "overdue": overdue}
    })

# ----- Employee add/list/quick-add -----
@app.route('/employee/add', methods=['GET', 'POST'])
@login_required
def add_employee():
    if current_user.admin_level < 1:
        flash('Access denied.')
        return redirect(url_for('main.home'))
    form = EmployeeForm()
    if form.validate_on_submit():
        emp = Employee(
            first_name=form.first_name.data,
            last_name=form.last_name.data,
            job_title=form.job_title.data,
            line_manager=form.line_manager.data,
            service=form.service.data,
            start_date=form.start_date.data,
            team_id=form.team_id.data,
            email=form.email.data
        )
        db.session.add(emp)
        db.session.commit()
        flash('New employee added.')
        return redirect(url_for('employee_list'))
    return render_template('add_employee.html', form=form)

@app.route('/employee/quick-add', methods=['POST'])
@login_required
@csrf.exempt
def quick_add_employee():
    data = request.get_json(force=True, silent=True) or {}
    first = (data.get("first_name") or "").strip()
    last = (data.get("last_name") or "").strip()
    role = (data.get("role") or "").strip()
    service = (data.get("service") or "").strip()

    if not first or not last:
        return jsonify({"success": False, "error": "First and last name are required"}), 400

    emp = Employee(
        first_name=first,
        last_name=last,
        role=role,  # kept as-is (model may accept 'role')
        service=service,
        manager_id=getattr(current_user, "id", None)
    )
    db.session.add(emp)
    db.session.commit()

    try:
        evt = TimelineEvent(
            event_type="Employee Created",
            notes="Employee created via Quick-Add in wizard",
            updated_by=current_user.username
        )
        db.session.add(evt)
        db.session.commit()
    except Exception:
        pass

    return jsonify({"success": True, "id": emp.id, "display_name": f"{emp.first_name} {emp.last_name}"})

@app.route('/employee/list')
@login_required
def employee_list():
    template = 'probation_employee_list.html' if session.get('active_module') == 'Probation' else 'employee_list.html'
    q = Employee.query
    if current_user.admin_level == 0:
        if current_user.team_id:
            q = q.filter(Employee.team_id == current_user.team_id)
        else:
            q = q.filter(False)
    employees = q.order_by(Employee.last_name.asc(), Employee.first_name.asc()).all()
    return render_template(template, employees=employees)

# ----- Employee Import: upload/validate/commit -----
def _suggest_mapping(headers):
    mapping = {}
    for h in headers or []:
        n = _normalize_header(h)
        if n in EMPLOYEE_FIELDS:
            mapping[h] = n
        elif n in ("firstname", "first"):
            mapping[h] = "first_name"
        elif n in ("lastname", "last", "surname"):
            mapping[h] = "last_name"
        elif n in ("mail", "email_address"):
            mapping[h] = "email"
        elif n in ("role", "position", "title", "job", "jobrole", "job_role", "jobtitle"):
            mapping[h] = "job_title"
        elif n in ("manager", "line_manager", "linemanager", "manager_name"):
            mapping[h] = "line_manager"
        elif n in ("team", "teamid", "team_id", "dept", "department"):
            mapping[h] = "team_id"
        else:
            mapping[h] = ""
    return mapping

@app.route("/employee/import", methods=["GET", "POST"])
@login_required
@superuser_required
def employee_import():
    if request.method == "GET":
        return render_template("employee_import.html")

    file = request.files.get("file")
    if not file or file.filename == "":
        abort(400, "No file uploaded")

    filename = secure_filename(file.filename)
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_EXTS:
        abort(400, f"Unsupported file type: .{ext}")

    file_bytes = file.read()

    if ext == "csv":
        headers, rows = _read_csv_bytes(file_bytes)
    else:
        headers, rows = _read_xlsx_bytes(file_bytes)

    normalised_headers = [{"raw": h, "norm": _normalize_header(h)} for h in (headers or [])]

    tmp = tempfile.NamedTemporaryFile(prefix="emp_import_", suffix=f".{ext}", delete=False)
    tmp.write(file_bytes)
    tmp.flush()
    temp_id = tmp.name

    preview = rows[:10] if rows else []

    return jsonify({
        "temp_id": temp_id,
        "headers": headers,
        "headers_norm": normalised_headers,
        "suggested_mapping": _suggest_mapping(headers),
        "preview_rows": preview,
        "xlsx_enabled": XLSX_ENABLED
    }), 200

@app.route("/employee/import/validate", methods=["POST"])
@login_required
@superuser_required
def employee_import_validate():
    data = request.get_json(force=True, silent=True) or {}
    temp_id = data.get("temp_id")
    mapping = data.get("mapping") or {}
    unique_key = (data.get("unique_key") or "email").strip()

    if not temp_id:
        abort(400, "Missing temp_id")
    try:
        with open(temp_id, "rb") as f:
            file_bytes = f.read()
    except Exception:
        abort(400, "Invalid temp_id or temporary file expired")

    ext = temp_id.rsplit(".", 1)[-1].lower()
    headers, rows = (_read_csv_bytes(file_bytes) if ext == "csv" else _read_xlsx_bytes(file_bytes))

    mapped_rows = []
    unmapped_headers = []
    for h in headers or []:
        if not mapping.get(h):
            unmapped_headers.append(h)

    for r in rows:
        out = {}
        for h, v in r.items():
            field = mapping.get(h)
            if not field:
                continue
            if field == "start_date":
                out[field] = _try_parse_date(v)
            else:
                out[field] = (str(v).strip() if v is not None else None)
        mapped_rows.append(out)

    missing_required = []
    for idx, r in enumerate(mapped_rows, start=1):
        missing = [f for f in REQUIRED_FIELDS if not r.get(f)]
        if missing:
            missing_required.append({"row": idx, "missing": missing})

    duplicates_in_file = []
    if unique_key:
        keys = unique_key.split(",")
        seen = set()
        for idx, r in enumerate(mapped_rows, start=1):
            key_tuple = tuple((r.get(k.strip()) or "").lower() for k in keys)
            if all(key_tuple):
                if key_tuple in seen:
                    duplicates_in_file.append({"row": idx, "key": key_tuple})
                else:
                    seen.add(key_tuple)

    duplicates_in_db = []
    try:
        if unique_key == "email" and any(r.get("email") for r in mapped_rows):
            emails = list({r.get("email") for r in mapped_rows if r.get("email")})
            existing = set(e[0].lower() for e in db.session.query(Employee.email).filter(Employee.email.in_(emails)).all())
            for idx, r in enumerate(mapped_rows, start=1):
                em = (r.get("email") or "").lower()
                if em and em in existing:
                    duplicates_in_db.append({"row": idx, "email": r.get("email")})
        elif unique_key == "first_name,last_name":
            pairs = {((r.get("first_name") or "").lower(), (r.get("last_name") or "").lower())
                     for r in mapped_rows if r.get("first_name") and r.get("last_name")}
            if pairs:
                q = db.session.query(Employee.first_name, Employee.last_name).all()
                existing_pairs = {(fn.lower(), ln.lower()) for fn, ln in q}
                for idx, r in enumerate(mapped_rows, start=1):
                    t = ((r.get("first_name") or "").lower(), (r.get("last_name") or "").lower())
                    if all(t) and t in existing_pairs:
                        duplicates_in_db.append({"row": idx, "name": f"{r.get('first_name')} {r.get('last_name')}"})
    except Exception as e:
        duplicates_in_db = [{"error": f"DB duplicate check skipped: {e}"}]

    report = {
        "unmapped_headers": unmapped_headers,
        "missing_required": missing_required,
        "duplicates_in_file": duplicates_in_file,
        "duplicates_in_db": duplicates_in_db,
        "rows_ready": len(mapped_rows) - len(missing_required) - len(duplicates_in_file) - len(duplicates_in_db),
        "total_rows": len(mapped_rows)
    }
    return jsonify({"temp_id": temp_id, "report": report}), 200

@app.route("/employee/import/commit", methods=["POST"])
@login_required
@superuser_required
def employee_import_commit():
    data = request.get_json(force=True, silent=True) or {}
    temp_id = data.get("temp_id")
    mapping = data.get("mapping") or {}
    confirm = bool(data.get("confirm"))
    unique_key = (data.get("unique_key") or "email").strip()
    if not (temp_id and confirm and mapping):
        abort(400, "Missing temp_id, mapping, or confirm flag")

    try:
        with open(temp_id, "rb") as f:
            file_bytes = f.read()
    except Exception:
        abort(400, "Invalid temp_id or temporary file expired")

    ext = temp_id.rsplit(".", 1)[-1].lower()
    headers, rows = (_read_csv_bytes(file_bytes) if ext == "csv" else _read_xlsx_bytes(file_bytes))

    created, skipped, errors = 0, 0, []

    for idx, src in enumerate(rows, start=1):
        payload = {}
        for h, v in src.items():
            field = mapping.get(h)
            if not field:
                continue
            if field == "start_date":
                payload[field] = _try_parse_date(v)
            else:
                payload[field] = (str(v).strip() if v is not None else None)

        if any(not payload.get(f) for f in REQUIRED_FIELDS):
            skipped += 1
            continue

        try:
            exists = False
            if unique_key == "email" and payload.get("email"):
                exists = db.session.query(Employee.id).filter_by(email=payload["email"]).first() is not None
            elif unique_key == "first_name,last_name" and payload.get("first_name") and payload.get("last_name"):
                exists = db.session.query(Employee.id).filter_by(
                    first_name=payload["first_name"], last_name=payload["last_name"]
                ).first() is not None
            if exists:
                skipped += 1
                continue
        except Exception as e:
            errors.append({"row": idx, "error": f"Duplicate check failed: {e}"})
            skipped += 1
            continue

        try:
            emp = Employee(**{k: v for k, v in payload.items() if k in EMPLOYEE_FIELDS})
            db.session.add(emp)
            created += 1
        except Exception as e:
            errors.append({"row": idx, "error": str(e)})
            skipped += 1

    db.session.commit()

    try:
        username = getattr(current_user, "username", "system")
        notes = f"Employee Import: created={created}, skipped={skipped}, errors={len(errors)}"
        evt = TimelineEvent(event_type="Import", notes=notes, updated_by=username)
        db.session.add(evt)
        db.session.commit()
    except Exception:
        pass

    return jsonify({"created": created, "skipped": skipped, "errors": errors}), 200

# ----- DOCX generation endpoints -----
@app.route('/pip/<int:id>/generate/invite')
@login_required
def generate_invite_letter(id):
    pip = PIPRecord.query.get_or_404(id)
    emp = pip.employee
    context = {
        "[[LETTER_DATE]]": datetime.now(ZoneInfo("Europe/London")).strftime("%d %B %Y"),
        "[[EMPLOYEE_NAME]]": f"{emp.first_name} {emp.last_name}",
        "[[EMPLOYEE_FIRST_NAME]]": emp.first_name or "",
        "[[JOB_TITLE]]": emp.job_title or "",
        "[[DEPARTMENT]]": emp.service or "",
        "[[MANAGER_NAME]]": pip.created_by or getattr(current_user, "username", ""),
        "[[MANAGER_TITLE]]": getattr(emp, "line_manager", "") or "",
        "[[MEETING_LOCATION]]": pip.capability_meeting_venue or "",
        "[[MEETING_DATETIME]]": (
            f"{(pip.capability_meeting_date.strftime('%d %B %Y') if pip.capability_meeting_date else '')} "
            f"{(pip.capability_meeting_time or '')}"
        ).strip(),
        "[[CONCERNS_SUMMARY]]": pip.concerns or "",
        "[[EVIDENCE_SUMMARY]]": "",
        "[[SUPPORT_LIST]]": "",
        "[[PIP_START_DATE]]": pip.start_date.strftime("%d %B %Y") if pip.start_date else "",
        "[[REVIEW_PERIOD_WEEKS]]": "",
        "[[HR_CONTACT_NAME]]": "",
        "[[HR_CONTACT_EMAIL]]": "",
    }
    buf = render_docx("PIP_Invite_Letter_Template_v2025-08-28.docx", context)
    filename = f"Invite_Letter_{emp.last_name}_{pip.id}.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/pip/<int:id>/generate/plan')
@login_required
def generate_plan_document(id):
    pip = PIPRecord.query.get_or_404(id)
    emp = pip.employee

    obj_ctx = {}
    objects = (pip.objectives[:5] if hasattr(pip, "objectives") and pip.objectives else pip.action_items[:5])
    for i, obj in enumerate(objects, start=1):
        text = getattr(obj, "text", None) or getattr(obj, "description", "") or ""
        obj_ctx.update({
            f"[[OBJ{i}_TEXT]]": text,
            f"[[OBJ{i}_SUCCESS]]": getattr(obj, "success_criteria", "") or "",
            f"[[OBJ{i}_ACTIONS_EMPLOYEE]]": getattr(obj, "actions_employee", "") or "",
            f"[[OBJ{i}_ACTIONS_MANAGER]]": getattr(obj, "actions_manager", "") or "",
            f"[[OBJ{i}_SUPPORT]]": getattr(obj, "support", "") or "",
            f"[[OBJ{i}_DEADLINE]]": (obj.deadline.strftime("%d %B %Y") if getattr(obj, "deadline", None) else ""),
            f"[[OBJ{i}_METRICS]]": getattr(obj, "metrics", "") or "",
        })

    context = {
        "[[EMPLOYEE_NAME]]": f"{emp.first_name} {emp.last_name}",
        "[[JOB_TITLE]]": emp.job_title or "",
        "[[DEPARTMENT]]": emp.service or "",
        "[[MANAGER_NAME]]": pip.created_by or getattr(current_user, "username", ""),
        "[[REVIEWER_NAME]]": "",
        "[[HR_CONTACT_NAME]]": "",
        "[[PIP_START_DATE]]": pip.start_date.strftime("%d %B %Y") if pip.start_date else "",
        "[[PIP_END_DATE]]": pip.review_date.strftime("%d %B %Y") if pip.review_date else "",
        "[[REVIEW_PERIOD_WEEKS]]": "",
        "[[POLICY_REFERENCE]]": getattr(pip, "policy_reference", "") or "",
        "[[CONCERN_CATEGORIES]]": getattr(pip, "concern_category", "") or "",
        "[[CONCERN_TAGS]]": (pip.tags or ""),
        "[[SEVERITY_RATING]]": getattr(pip, "severity", "") or "",
        "[[FREQUENCY_RATING]]": getattr(pip, "frequency", "") or "",
        "[[CONFIDENTIAL_NOTES]]": getattr(pip, "confidential_notes", "") or "",
        "[[REVIEW_DATES_LIST]]": getattr(pip, "review_dates_str", "") or "",
        "[[REPORTING_METHOD]]": getattr(pip, "reporting_method", "") or "",
        "[[ADJUSTMENTS]]": getattr(pip, "adjustments", "") or "",
        "[[TRAINING_PLAN]]": getattr(pip, "training_plan", "") or "",
        **obj_ctx,
    }

    buf = render_docx("PIP_Action_Plan_Template_v2025-08-28.docx", context)
    filename = f"PIP_Plan_{emp.last_name}_{pip.id}.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/pip/<int:id>/generate/outcome')
@login_required
def generate_outcome_letter(id):
    pip = PIPRecord.query.get_or_404(id)
    emp = pip.employee
    choice = (request.args.get("outcome", "SUCCESSFUL") or "SUCCESSFUL").upper()

    obj_ctx = {}
    objects = (pip.objectives[:5] if hasattr(pip, "objectives") and pip.objectives else pip.action_items[:5])
    for i, obj in enumerate(objects, start=1):
        text = getattr(obj, "text", None) or getattr(obj, "description", "") or ""
        obj_ctx.update({
            f"[[OBJ{i}_TEXT]]": text,
            f"[[OBJ{i}_OUTCOME]]": getattr(obj, "outcome_text", "") or "",
            f"[[OBJ{i}_EVIDENCE]]": getattr(obj, "outcome_evidence", "") or "",
        })

    context = {
        "[[LETTER_DATE]]": datetime.now(ZoneInfo("Europe/London")).strftime("%d %B %Y"),
        "[[EMPLOYEE_NAME]]": f"{emp.first_name} {emp.last_name}",
        "[[EMPLOYEE_FIRST_NAME]]": emp.first_name or "",
        "[[JOB_TITLE]]": emp.job_title or "",
        "[[DEPARTMENT]]": emp.service or "",
        "[[MANAGER_NAME]]": pip.created_by or getattr(current_user, "username", ""),
        "[[MANAGER_TITLE]]": getattr(emp, "line_manager", "") or "",
        "[[PIP_START_DATE]]": pip.start_date.strftime("%d %B %Y") if pip.start_date else "",
        "[[PIP_END_DATE]]": pip.review_date.strftime("%d %B %Y") if pip.review_date else "",
        "[[ONGOING_SUPPORT]]": getattr(pip, "ongoing_support", "") or "",
        "[[EXTENSION_WEEKS]]": str(getattr(pip, "extension_weeks", "") or ""),
        "[[HR_CONTACT_NAME]]": "",
        "[[HR_CONTACT_EMAIL]]": "",
        "[[APPENDIX_DETAIL]]": getattr(pip, "outcome_appendix", "") or "",
        **obj_ctx,
    }

    buf = render_docx("PIP_Outcome_Letter_Template_v2025-08-28.docx", context, outcome_choice=choice)
    filename = f"Outcome_Letter_{emp.last_name}_{pip.id}.docx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


from datetime import date, timedelta, datetime
import random

def seed_demo_data():
    """
    One-off demo data seeder for local/dev.

    - Creates ~25 employees across a few services/teams
    - Adds PIPs, Probations, and Sickness cases with meetings
    - Does nothing if Employees already exist
    """
    # Protect against accidental double-seed
    if Employee.query.count() > 0:
        print("[seed_demo_data] Employees already exist, skipping seed.")
        return

    today = date.today()

    services = [
        "Riverside",
        "Oak Lodge",
        "Willow Court",
        "Maple House",
        "Cedar View",
    ]
    job_titles = [
        "Support Worker",
        "Senior Support Worker",
        "Team Leader",
        "Deputy Manager",
        "Registered Manager",
    ]
    managers = [
        "Alex Smith",
        "Jamie Lee",
        "Pat Taylor",
        "Jordan Brown",
    ]

    # --- 1. Create Employees ---

    employee_specs = [
        ("Emily", "Clark"),
        ("James", "Wilson"),
        ("Sophie", "Turner"),
        ("Daniel", "Hughes"),
        ("Olivia", "Morris"),
        ("Liam", "Patel"),
        ("Chloe", "Roberts"),
        ("Ethan", "Walker"),
        ("Isla", "Thompson"),
        ("Noah", "White"),
        ("Grace", "Hall"),
        ("Jacob", "Green"),
        ("Amelia", "Harrison"),
        ("Leo", "Wood"),
        ("Mia", "Baker"),
        ("Lucas", "Adams"),
        ("Ella", "Kelly"),
        ("Mason", "Ward"),
        ("Freya", "Watson"),
        ("Henry", "Mitchell"),
        ("Ruby", "Carter"),
        ("Oscar", "Bailey"),
        ("Poppy", "Cooper"),
        ("George", "Richardson"),
        ("Lily", "Wright"),
    ]

    employees = []

    for i, (first, last) in enumerate(employee_specs, start=1):
        start_offset = random.randint(30, 720)  # started sometime in last 2 years
        start_dt = today - timedelta(days=start_offset)

        emp = Employee(
            first_name=first,
            last_name=last,
            job_title=random.choice(job_titles),
            line_manager=random.choice(managers),
            service=random.choice(services),
            start_date=start_dt,
            team_id=random.choice([1, 2, 3]),
            email=f"{first.lower()}.{last.lower()}@example.com",
        )
        db.session.add(emp)
        employees.append(emp)

    db.session.flush()  # Ensure IDs are assigned

    # --- 2. Helper functions for PIP, Probation, Sickness ---

    def create_pip_for_employee(emp: Employee):
        """Create 0–2 PIPs for a given employee."""
        # 50% chance of at least one PIP
        if random.random() < 0.5:
            return

        num_pips = 1
        if random.random() < 0.3:
            num_pips = 2

        for _ in range(num_pips):
            start_offset = random.randint(30, 365)
            pip_start = today - timedelta(days=start_offset)
            review_days = random.choice([30, 60, 90])
            review_date = pip_start + timedelta(days=review_days)

            status = random.choice(["Open", "Completed", "Closed"])
            concern_area = random.choice(
                ["timekeeping", "documentation", "care planning", "communication", "professional conduct"]
            )

            pip = PIPRecord(
                employee_id=emp.id,
                concerns=f"Performance concerns relating to {concern_area}.",
                start_date=pip_start,
                review_date=review_date,
                meeting_notes="Initial performance discussion held.\nObjectives agreed with employee.",
                status=status,
                created_by="seed_demo_data",
                created_at=datetime.utcnow(),
                last_updated=datetime.utcnow(),
                concern_category=random.choice(
                    ["Performance", "Conduct", "Capability", "Quality"]
                ),
                severity=random.choice(["Low", "Medium", "High"]),
                frequency=random.choice(["One-off", "Intermittent", "Persistent"]),
                tags="seed, demo",
            )
            db.session.add(pip)
            db.session.flush()

            # Add 1–3 action items
            num_actions = random.randint(1, 3)
            for i in range(num_actions):
                ai_status = random.choice(["Outstanding", "In Progress", "Completed"])
                action = PIPActionItem(
                    pip_record_id=pip.id,
                    description=f"Action {i+1} for {emp.first_name} {emp.last_name}.",
                    status=ai_status,
                    created_at=datetime.utcnow(),
                )
                db.session.add(action)

            # Add 1–2 timeline events
            num_events = random.randint(1, 2)
            for _ in range(num_events):
                evt = TimelineEvent(
                    pip_record_id=pip.id,
                    timestamp=datetime.utcnow(),
                    event_type=random.choice(["Note", "Meeting", "Document Generated"]),
                    notes="Demo timeline event generated by seed_demo_data.",
                    updated_by="seed_demo_data",
                    created_at=datetime.utcnow(),
                )
                db.session.add(evt)

    def create_probation_for_employee(emp: Employee):
        """Create 0–2 probation records per employee."""
        if random.random() < 0.6:
            return

        num_probs = 1
        if random.random() < 0.25:
            num_probs = 2

        for _ in range(num_probs):
            # Start date near employee start, or within last 18 months
            base_start = emp.start_date or (today - timedelta(days=random.randint(60, 540)))
            # Nudge around that date a bit
            start_dt = base_start + timedelta(days=random.randint(-30, 30))
            expected_end = start_dt + timedelta(days=180)  # 6-month probation

            status = random.choice(["Active", "Extended", "Completed", "Failed"])

            prob = ProbationRecord(
                employee_id=emp.id,
                start_date=start_dt,
                expected_end_date=expected_end,
                status=status,
                notes="Demo probation record for seed data.",
                created_at=datetime.utcnow(),
                last_updated=datetime.utcnow(),
            )
            db.session.add(prob)
            db.session.flush()

            # 1–2 reviews
            num_reviews = random.randint(1, 2)
            for i in range(num_reviews):
                review_dt = start_dt + timedelta(days=(i + 1) * 60)
                review = ProbationReview(
                    probation_id=prob.id,
                    review_date=review_dt,
                    reviewer=random.choice(managers),
                    summary="Progress review completed.\nSeeded for demo.",
                    concerns_flag=random.random() < 0.3,
                    created_at=datetime.utcnow(),
                )
                db.session.add(review)

            # Optional plan
            if random.random() < 0.5:
                plan_deadline = expected_end
                plan = ProbationPlan(
                    probation_id=prob.id,
                    created_at=datetime.utcnow(),
                    objectives="Demonstrate consistent performance in key duties.\nAttend all mandatory training.",
                    deadline=plan_deadline,
                    outcome=random.choice(["Met", "Not Met", "Ongoing", None]),
                )
                db.session.add(plan)

    def create_sickness_for_employee(emp: Employee):
        """Create 0–3 sickness cases with a mix of short/long/pattern."""
        r = random.random()
        if r < 0.4:
            return
        elif r < 0.75:
            num_cases = 1
        elif r < 0.9:
            num_cases = 2
        else:
            num_cases = 3

        for _ in range(num_cases):
            case_type = random.choice(["short_term", "long_term", "pattern"])
            start_offset = random.randint(0, 365)
            start_dt = today - timedelta(days=start_offset)

            end_dt = None
            reason = None
            trigger_type = None
            status = "Open"

            if case_type == "short_term":
                duration = random.randint(1, 5)
                end_dt = start_dt + timedelta(days=duration)
                trigger_type = "short_term"
                reason = random.choice(["Flu", "Stomach bug", "Migraine", "Minor surgery"])
            elif case_type == "long_term":
                duration = random.randint(28, 90)
                end_dt = start_dt + timedelta(days=duration)
                trigger_type = "long_term"
                reason = random.choice(["Back injury", "Stress", "Depression", "Post-op recovery"])
            else:  # pattern
                duration = random.randint(2, 10)
                end_dt = start_dt + timedelta(days=duration)
                trigger_type = "pattern"
                reason = random.choice(["Recurring short-term absences", "Unclear pattern"])

            # If it ends in the future, treat as still open
            if end_dt and end_dt > today:
                end_dt = None  # ongoing
                status = random.choice(["Open", "Monitoring"])
            else:
                status = random.choice(["Open", "Closed", "Monitoring"])

            scase = SicknessCase(
                employee_id=emp.id,
                start_date=start_dt,
                end_date=end_dt,
                reason=reason,
                trigger_type=trigger_type,
                notes="Demo sickness case created by seed_demo_data.",
                status=status,
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow(),
            )
            db.session.add(scase)
            db.session.flush()

            # 0–2 meetings per case
            num_meetings = random.randint(0, 2)
            for _ in range(num_meetings):
                if end_dt:
                    m_date = end_dt + timedelta(days=random.randint(0, 7))
                else:
                    # For ongoing, meeting near today or near start
                    m_date = min(
                        today + timedelta(days=random.randint(0, 14)),
                        today + timedelta(days=21),
                    )

                meeting = SicknessMeeting(
                    sickness_case_id=scase.id,
                    meeting_date=m_date,
                    meeting_type=random.choice(
                        ["RTW", "Absence Review", "Welfare", "Other"]
                    ),
                    chair=random.choice(managers),
                    notes="Seeded sickness meeting for demo.\nUse this to test RTW logging.",
                    outcome=random.choice(
                        [
                            "Employee returned to work.",
                            "Further monitoring required.",
                            "Occupational health referral discussed.",
                            None,
                        ]
                    ),
                    created_at=datetime.utcnow(),
                )
                db.session.add(meeting)

    # --- 3. Generate module data per employee ---

    for emp in employees:
        create_pip_for_employee(emp)
        create_probation_for_employee(emp)
        create_sickness_for_employee(emp)

    db.session.commit()
    print("[seed_demo_data] Demo data created: "
          f"{len(employees)} employees with PIPs, probations, and sickness cases.")


# ----- DB bootstrap -----
with app.app_context():
    if not os.path.exists(os.path.join(BASE_DIR, 'pip_crm.db')):
        db.create_all()
        print('✅ Database created')

# ----- Health -----
@app.route('/ping')
def ping():
    return 'Pong!'

print('✅ Flask app initialized and ready.')

if __name__ == "__main__":
    with app.app_context():
        # Uncomment the next line once, run, then comment it out again
        seed_demo_data()
        app.run(host="0.0.0.0")
